"""
donor_stewardship_sequence/render.py

Generates a 3-touch donor stewardship package as a Word (.docx) document,
tailored for nonprofit Development Directors.

Sections:
  1. Touch 1 — Acknowledgement Letter (within 48 hours)
  2. Touch 2 — Thank-You Call Script (within 1 week)
  3. Touch 3 — Impact Update Email (60-90 days post-gift)

All libraries are pre-installed in Anthropic's code-execution sandbox.
"""

import datetime
import os
import re
import time
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor


# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------

_NAVY = RGBColor(0x1A, 0x47, 0x6B)
_STEEL = RGBColor(0x2E, 0x72, 0x9E)
_AMBER = RGBColor(0xAA, 0x44, 0x00)
_GRAY = RGBColor(0xCC, 0xCC, 0xCC)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _heading1(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=1)
    if h.runs:
        h.runs[0].font.color.rgb = _NAVY


def _heading2(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=2)
    if h.runs:
        h.runs[0].font.color.rgb = _STEEL


def _para(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph(text)
    if italic and p.runs:
        p.runs[0].italic = True


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _bold_label(para, label: str, value: str) -> None:
    r = para.add_run(label + ": ")
    r.bold = True
    para.add_run(value)


def _hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 72)
    run.font.color.rgb = _GRAY
    run.font.size = Pt(7)


def _extract_first_name(donor_name: str) -> str:
    """Return the first token of the donor name as a fallback salutation."""
    parts = donor_name.strip().split()
    return parts[0] if parts else donor_name


# ---------------------------------------------------------------------------
# Touch 1 — Acknowledgement Letter
# ---------------------------------------------------------------------------

def _build_acknowledgement_letter(
    doc: Document,
    donor_name: str,
    donor_salutation: str,
    gift_amount: str,
    gift_date: str,
    gift_purpose: str,
    org_name: str,
    signer_name: str,
    signer_title: str,
) -> None:
    _heading1(doc, "Touch 1 — Acknowledgement Letter")

    timing_p = doc.add_paragraph()
    t_run = timing_p.add_run("Timing: Send within 48 hours of receiving the gift.")
    t_run.italic = True
    t_run.font.color.rgb = _STEEL
    t_run.font.size = Pt(9)

    _hr(doc)
    doc.add_paragraph()

    # Date
    today = datetime.date.today().strftime("%B %d, %Y")
    _para(doc, today)
    doc.add_paragraph()

    # Address block
    addr_p = doc.add_paragraph()
    addr_r = addr_p.add_run(f"{donor_name}\n[Address Line 1]\n[City, State ZIP]")
    addr_r.font.color.rgb = _AMBER
    addr_r.italic = True
    doc.add_paragraph()

    # Salutation
    _para(doc, f"Dear {donor_salutation},")
    doc.add_paragraph()

    # Body paragraph 1 — gratitude + gift specifics
    body1 = (
        f"On behalf of the entire {org_name} team, I want to express our sincere gratitude "
        f"for your generous gift of {gift_amount}, received on {gift_date}. Your support of "
        f"{gift_purpose} makes a direct and meaningful difference for the people we serve."
    )
    _para(doc, body1)

    # Body paragraph 2 — impact bridge
    body2 = (
        f"Because of donors like you, {org_name} can continue to deliver high-quality programming "
        f"that opens doors for those who face the greatest barriers. Your investment is not just "
        f"a donation — it is a statement of belief in the potential of every person we serve. "
        f"[PLACEHOLDER: Insert 1-2 sentences of specific program impact or a brief participant "
        f"success story here to personalize this paragraph.]"
    )
    _para(doc, body2)

    # IRS language
    irs_para = doc.add_paragraph()
    irs_run = irs_para.add_run(
        f"For your records: {org_name} is a tax-exempt organization under Section 501(c)(3) "
        f"of the Internal Revenue Code (EIN: [EIN]). No goods or services were provided to you "
        f"in exchange for this contribution. This letter serves as your official tax receipt for "
        f"a gift of {gift_amount} to {org_name} on {gift_date}."
    )
    irs_run.font.size = Pt(10)
    irs_run.italic = True

    doc.add_paragraph()

    # Close
    _para(doc, "With gratitude,")
    doc.add_paragraph()
    _para(doc, f"{signer_name}\n{signer_title}\n{org_name}")
    doc.add_paragraph()

    ps_para = doc.add_paragraph()
    ps_r = ps_para.add_run("P.S. ")
    ps_r.bold = True
    ps_para.add_run(
        f"[PLACEHOLDER: A brief, personal P.S. dramatically increases read rates for major gift "
        f"letters. Consider adding a sentence referencing something specific about {donor_name}'s "
        f"connection to your mission, an upcoming event, or a program milestone.]"
    )
    ps_para.runs[-1].italic = True


# ---------------------------------------------------------------------------
# Touch 2 — Thank-You Call Script
# ---------------------------------------------------------------------------

def _build_call_script(
    doc: Document,
    donor_name: str,
    donor_salutation: str,
    gift_amount: str,
    gift_purpose: str,
    org_name: str,
    signer_name: str,
    signer_title: str,
    program_impact_data: Optional[str],
    donor_giving_history_summary: Optional[str],
) -> None:
    _heading1(doc, "Touch 2 — Thank-You Call Script")

    timing_p = doc.add_paragraph()
    t_run = timing_p.add_run("Timing: Make this call within 7 days of the gift. Keep it to 5-10 minutes.")
    t_run.italic = True
    t_run.font.color.rgb = _STEEL
    t_run.font.size = Pt(9)

    _hr(doc)

    _heading2(doc, "Before the Call — Caller Prep")

    if donor_giving_history_summary:
        history_p = doc.add_paragraph()
        _bold_label(history_p, "Donor History", donor_giving_history_summary)

    prep_items = [
        f"Have the gift amount ({gift_amount}) and gift date ready to reference naturally",
        "Review any prior notes on this donor before dialing",
        "Identify one specific program update or participant story to share",
        "If leaving a voicemail, keep it under 45 seconds and include your callback number",
    ]
    for item in prep_items:
        _bullet(doc, item)

    doc.add_paragraph()

    _heading2(doc, "Opening")
    opening_script = (
        f'"Hello, {donor_salutation}? This is {signer_name} calling from {org_name}. '
        f'Do you have just a few minutes? I wanted to reach out personally to thank you '
        f'for your generous support."'
    )
    _para(doc, opening_script, italic=True)

    doc.add_paragraph()

    _heading2(doc, "Gratitude Expression")
    gratitude_script = (
        f'"Your gift of {gift_amount} to {gift_purpose} means so much to us. '
        f'We don\'t take that kind of generosity for granted, and I wanted to make sure '
        f'you heard directly from me."'
    )
    _para(doc, gratitude_script, italic=True)

    doc.add_paragraph()

    _heading2(doc, "Impact Talking Points")
    note_p = doc.add_paragraph()
    note_p.add_run("Choose 2-3 of the following, or personalize with your own program data:").italic = True

    if program_impact_data:
        _bullet(doc, f'"{program_impact_data}"')

    default_talking_points = [
        f'"Because of gifts like yours, we were able to [PLACEHOLDER: specific outcome — e.g., \'connect 50 young people with job placements this year\']."',
        f'"The {gift_purpose} fund goes directly to [PLACEHOLDER: describe exactly what this funding supports — participants, staff, materials, etc.]."',
        '"One of the things I\'m most proud of is [PLACEHOLDER: share a brief, genuine story or milestone — keep it human, not programmatic]."',
    ]
    for tp in default_talking_points:
        _bullet(doc, tp)

    doc.add_paragraph()

    _heading2(doc, "Learn More / Deepen Engagement")
    engage_script = (
        f'"I\'d love for you to see the work in person sometime, {donor_salutation}. '
        f'We have [PLACEHOLDER: event name, site visit, or program demo opportunity coming up] — '
        f'would that be something you\'d enjoy? No pressure at all — I just think you\'d love '
        f'seeing the impact of your support firsthand."'
    )
    _para(doc, engage_script, italic=True)

    doc.add_paragraph()

    _heading2(doc, "Close")
    close_script = (
        f'"Thank you again, {donor_salutation}. We are grateful for your partnership. '
        f'Is there anything you\'d like to know about [PLACEHOLDER: org or program name]? '
        f'... [Listen, respond genuinely.] Have a wonderful day — and please don\'t hesitate '
        f'to reach out if you ever have questions."'
    )
    _para(doc, close_script, italic=True)

    doc.add_paragraph()

    _heading2(doc, "After the Call — Log Notes")
    log_items = [
        "Record call date, duration, and key topics in your CRM",
        "Note any interests, connections, or follow-up items the donor mentioned",
        "Flag for Touch 3 (impact email) in 60-90 days",
        "If donor mentioned event interest, add to event RSVP list",
    ]
    for item in log_items:
        _bullet(doc, item)


# ---------------------------------------------------------------------------
# Touch 3 — Impact Update Email
# ---------------------------------------------------------------------------

def _build_impact_email(
    doc: Document,
    donor_name: str,
    donor_salutation: str,
    gift_amount: str,
    gift_purpose: str,
    org_name: str,
    signer_name: str,
    signer_title: str,
    program_impact_data: Optional[str],
) -> None:
    _heading1(doc, "Touch 3 — Impact Update Email")

    timing_p = doc.add_paragraph()
    t_run = timing_p.add_run(
        "Timing: Send 60-90 days after the gift. This is not a fundraising ask — it is a relationship touchpoint."
    )
    t_run.italic = True
    t_run.font.color.rgb = _STEEL
    t_run.font.size = Pt(9)

    _hr(doc)

    _heading2(doc, "Subject Line")
    subject_p = doc.add_paragraph()
    subject_r = subject_p.add_run(
        f"[CHOOSE ONE]\n"
        f'Option A: "What your gift to {org_name} made possible"\n'
        f'Option B: "{donor_salutation}, here\'s the update I promised you"\n'
        f'Option C: "A story I thought you should hear, {donor_salutation}"'
    )
    subject_r.italic = True

    doc.add_paragraph()

    _heading2(doc, "Email Body")

    greeting = doc.add_paragraph()
    greeting.add_run(f"Dear {donor_salutation},")
    doc.add_paragraph()

    opening_para = (
        f"It has been a few months since you made your generous gift of {gift_amount} to "
        f"{gift_purpose} at {org_name}, and I wanted to take a moment to share what has "
        f"happened since."
    )
    _para(doc, opening_para)

    # Impact paragraph
    if program_impact_data:
        impact_para = (
            f"{program_impact_data} "
            f"[PLACEHOLDER: Add 1-2 sentences connecting this outcome directly to the donor's gift. "
            f"Example: 'Your generosity helped make this possible.']"
        )
    else:
        impact_para = (
            f"[PLACEHOLDER: Insert 1-2 sentences of specific program outcomes since the gift date. "
            f"Example: 'Since April, we have served 38 youth in our Pathways to Work program, "
            f"of whom 24 have completed job placements.' Connect those outcomes to this donor's gift.]"
        )
    _para(doc, impact_para)

    # Story bridge
    story_para = (
        f"[PLACEHOLDER: Share a brief anonymized participant story (2-4 sentences max). "
        f"Example: 'One participant — I'll call her M. — came to us uncertain about her future. "
        f"Three months later, she accepted her first job offer. When she called to tell us, "
        f"she said: \"I didn't think anyone believed in me until I came here.\"' "
        f"Stories like this are what your investment makes possible.]"
    )
    _para(doc, story_para, italic=True)

    # Photo placeholder
    photo_note = doc.add_paragraph()
    photo_r = photo_note.add_run(
        "[PHOTO PLACEHOLDER: Insert a program photo, participant moment, or event image here. "
        "Images with real people — with written consent — significantly increase email engagement.]"
    )
    photo_r.italic = True
    photo_r.font.color.rgb = _AMBER

    doc.add_paragraph()

    # Soft next step
    next_step = (
        f"I also wanted to give you a heads-up: we have [PLACEHOLDER: upcoming event, giving "
        f"opportunity, or program milestone] coming up, and I think you'd enjoy being part of it. "
        f"No ask here — just wanted you to have it on your radar."
    )
    _para(doc, next_step)

    # Close
    close_para = (
        f"Thank you again, {donor_salutation}. Your belief in {org_name} continues to matter "
        f"deeply — to our team and, most importantly, to the people we serve together."
    )
    _para(doc, close_para)
    doc.add_paragraph()

    _para(doc, f"With gratitude,\n{signer_name}\n{signer_title}\n{org_name}")
    doc.add_paragraph()

    footer_p = doc.add_paragraph()
    footer_r = footer_p.add_run(
        f"[UNSUBSCRIBE / CONTACT FOOTER — include per your email platform's compliance requirements]"
    )
    footer_r.italic = True
    footer_r.font.size = Pt(8)
    footer_r.font.color.rgb = _AMBER


# ---------------------------------------------------------------------------
# Core render function
# ---------------------------------------------------------------------------

def render(
    donor_name: str,
    gift_amount: str,
    gift_date: str,
    org_name: str,
    signer_name: str,
    signer_title: str,
    donor_salutation: Optional[str] = None,
    gift_purpose: Optional[str] = None,
    program_impact_data: Optional[str] = None,
    donor_giving_history_summary: Optional[str] = None,
    output_dir: str = "/mnt/user-data/outputs",
) -> str:
    """
    Generate a 3-touch donor stewardship sequence Word document.

    Parameters
    ----------
    donor_name : str
        Full name of the donor
    gift_amount : str
        Dollar amount of the gift (e.g. "$5,000")
    gift_date : str
        Date the gift was received (e.g. "April 15, 2026")
    org_name : str
        Full name of the organization
    signer_name : str
        Name of the person signing the acknowledgement letter
    signer_title : str
        Title of the signer
    donor_salutation : str, optional
        How to address the donor (e.g. "Dr. Chen"). Defaults to first name.
    gift_purpose : str, optional
        What the gift supports. Defaults to "general operating".
    program_impact_data : str, optional
        Recent outcome data to reference in the impact email.
    donor_giving_history_summary : str, optional
        Notes on the donor's giving history.
    output_dir : str
        Directory to save the output file

    Returns
    -------
    str
        Absolute path to the generated .docx file
    """
    # Validate required
    required = {
        "donor_name": donor_name,
        "gift_amount": gift_amount,
        "gift_date": gift_date,
        "org_name": org_name,
        "signer_name": signer_name,
        "signer_title": signer_title,
    }
    for field, value in required.items():
        if not value or not str(value).strip():
            raise ValueError(f"{field} is required and cannot be empty")

    # Defaults
    if not donor_salutation or not donor_salutation.strip():
        donor_salutation = _extract_first_name(donor_name)
    gift_purpose = gift_purpose or "general operating"

    doc = Document()

    # Document title
    title = doc.add_heading("Donor Stewardship Sequence", level=0)
    if title.runs:
        title.runs[0].font.color.rgb = _NAVY

    subtitle = doc.add_paragraph()
    sub_r = subtitle.add_run(f"{donor_name}  |  {org_name}  |  Gift: {gift_amount} on {gift_date}")
    sub_r.bold = True
    sub_r.font.size = Pt(12)
    sub_r.font.color.rgb = _STEEL

    _para(doc, "Three-touch stewardship package: acknowledgement letter, thank-you call script, and impact update email.")
    _hr(doc)

    # Touch 1
    _build_acknowledgement_letter(
        doc,
        donor_name=donor_name,
        donor_salutation=donor_salutation,
        gift_amount=gift_amount,
        gift_date=gift_date,
        gift_purpose=gift_purpose,
        org_name=org_name,
        signer_name=signer_name,
        signer_title=signer_title,
    )

    doc.add_page_break()

    # Touch 2
    _build_call_script(
        doc,
        donor_name=donor_name,
        donor_salutation=donor_salutation,
        gift_amount=gift_amount,
        gift_purpose=gift_purpose,
        org_name=org_name,
        signer_name=signer_name,
        signer_title=signer_title,
        program_impact_data=program_impact_data,
        donor_giving_history_summary=donor_giving_history_summary,
    )

    doc.add_page_break()

    # Touch 3
    _build_impact_email(
        doc,
        donor_name=donor_name,
        donor_salutation=donor_salutation,
        gift_amount=gift_amount,
        gift_purpose=gift_purpose,
        org_name=org_name,
        signer_name=signer_name,
        signer_title=signer_title,
        program_impact_data=program_impact_data,
    )

    # Save
    os.makedirs(output_dir, exist_ok=True)
    timestamp = int(time.time())
    safe_donor = re.sub(r"[^\w]+", "_", donor_name)[:30].strip("_")
    filename = f"donor_stewardship_{safe_donor}_{timestamp}.docx"
    out_path = os.path.join(output_dir, filename)
    doc.save(out_path)

    return out_path
