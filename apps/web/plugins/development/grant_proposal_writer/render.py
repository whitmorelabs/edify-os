"""
grant_proposal_writer/render.py

Generates a complete grant proposal or letter of inquiry (LOI) as a Word (.docx)
document, tailored for nonprofit Development Directors.

Sections (full proposal):
  1. Cover Page
  2. Executive Summary
  3. Statement of Need
  4. Program Description
  5. Goals and Objectives
  6. Evaluation Plan
  7. Organizational Capacity
  8. Budget Justification
  9. Sustainability

LOI variant: condensed to cover + executive summary + statement of need + goals.

All libraries are pre-installed in Anthropic's code-execution sandbox.
"""

import datetime
import os
import re
import time
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------

_NAVY = RGBColor(0x1A, 0x47, 0x6B)
_STEEL = RGBColor(0x2E, 0x72, 0x9E)
_GRAY = RGBColor(0xCC, 0xCC, 0xCC)


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _heading1(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=1)
    if h.runs:
        h.runs[0].font.color.rgb = _NAVY


def _heading2(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=2)
    if h.runs:
        h.runs[0].font.color.rgb = _STEEL


def _para(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph(text)
    if italic and p.runs:
        p.runs[0].italic = True


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _bold_label(para, label: str, value: str) -> None:
    r = para.add_run(label + ": ")
    r.bold = True
    para.add_run(value)


def _hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 72)
    run.font.color.rgb = _GRAY
    run.font.size = Pt(7)


def _add_budget_table(doc: Document) -> None:
    """Insert a placeholder budget table with standard nonprofit line items."""
    table = doc.add_table(rows=8, cols=3)
    table.style = "Table Grid"

    headers = ["Line Item", "Description", "Amount Requested"]
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = _NAVY

    line_items = [
        ("Personnel", "Program staff salaries & benefits (pro-rated)", "$[AMOUNT]"),
        ("Consultants / Contractors", "Specialized instructors or job coaches", "$[AMOUNT]"),
        ("Participant Stipends", "Stipends to support participant participation", "$[AMOUNT]"),
        ("Program Supplies", "Materials, curricula, and program supplies", "$[AMOUNT]"),
        ("Indirect / Admin", "Overhead (typically 10-15% of direct costs)", "$[AMOUNT]"),
        ("Other / Travel", "Staff travel, mileage, or transportation support", "$[AMOUNT]"),
        ("TOTAL", "", "$[TOTAL]"),
    ]

    for i, (item, desc, amt) in enumerate(line_items, start=1):
        row = table.rows[i]
        row.cells[0].text = item
        row.cells[1].text = desc
        row.cells[2].text = amt
        if item == "TOTAL":
            for cell in row.cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _build_cover_page(
    doc: Document,
    org_name: str,
    funder_name: str,
    grant_amount: str,
    deadline: str,
    program_name: str,
    proposal_type: str,
) -> None:
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(
        "Letter of Inquiry" if proposal_type == "loi" else "Grant Proposal"
    )
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = _NAVY

    prog_para = doc.add_paragraph()
    prog_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prog_run = prog_para.add_run(program_name)
    prog_run.bold = True
    prog_run.font.size = Pt(14)
    prog_run.font.color.rgb = _STEEL

    doc.add_paragraph()

    fields = [
        ("Submitted by", org_name),
        ("Submitted to", funder_name),
        ("Amount Requested", grant_amount),
        ("Submission Deadline", deadline),
        ("Primary Contact", "[Name, Title, Phone, Email]"),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bold_label(p, label, value)

    doc.add_paragraph()

    compliance = doc.add_paragraph()
    compliance.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = compliance.add_run(
        "GENERATED DRAFT — review with Development Director before submission."
    )
    cr.italic = True
    cr.font.color.rgb = RGBColor(0xAA, 0x44, 0x00)
    cr.font.size = Pt(9)


def _build_executive_summary(
    doc: Document,
    org_name: str,
    org_mission: str,
    program_name: str,
    funding_request_summary: str,
    grant_amount: str,
    funder_name: str,
) -> None:
    _heading1(doc, "Executive Summary")
    summary_text = (
        f"{org_name} respectfully requests {grant_amount} from {funder_name} "
        f"to support {program_name}. {org_mission} "
        f"{funding_request_summary}"
    )
    _para(doc, summary_text)


def _build_statement_of_need(
    doc: Document,
    org_name: str,
    target_population: str,
    geographic_area: str,
    program_name: str,
) -> None:
    _heading1(doc, "Statement of Need")

    need_text = (
        f"In {geographic_area}, {target_population} face significant barriers to "
        f"economic opportunity and community participation. Research consistently "
        f"demonstrates that transition-age youth with disabilities are among the most "
        f"underserved populations in the workforce development ecosystem — with unemployment "
        f"rates two to three times higher than their non-disabled peers. Many leave school "
        f"without the vocational skills, workplace experience, or professional networks "
        f"necessary to secure sustainable employment."
    )
    _para(doc, need_text)

    gap_text = (
        f"Existing resources in {geographic_area} are fragmented and inconsistently "
        f"accessible to {target_population}. Community college vocational programs often "
        f"lack the individualized supports this population requires. State vocational "
        f"rehabilitation services face long waitlists and limited employer engagement. "
        f"This gap leaves young people and their families navigating a complex system "
        f"with little coordinated support during a critical developmental window."
    )
    _para(doc, gap_text)

    org_role = (
        f"{org_name}'s {program_name} is designed to directly address this gap by "
        f"providing {target_population} with structured, evidence-informed programming "
        f"that bridges the transition from school to career. Our approach centers the "
        f"participant's strengths and builds toward self-determined employment outcomes."
    )
    _para(doc, org_role)


def _build_program_description(
    doc: Document,
    program_name: str,
    program_description: str,
    target_population: str,
    org_name: str,
) -> None:
    _heading1(doc, "Program Description")

    _heading2(doc, "Overview")
    _para(doc, program_description)

    _heading2(doc, "Activities and Methodology")
    activities = [
        f"Participant intake and individualized goal setting with a trained {org_name} staff member",
        "Skills-based instruction in vocational, communication, and workplace readiness competencies",
        "Employer site visits, informational interviews, and job shadow experiences",
        "One-on-one job coaching and career counseling sessions",
        "Post-placement support and check-ins to promote job retention",
    ]
    for activity in activities:
        _bullet(doc, activity)

    _heading2(doc, "Program Timeline")
    timeline_items = [
        "Months 1-2: Participant recruitment, intake assessments, and orientation",
        "Months 2-5: Core curriculum delivery and employer engagement",
        "Months 5-9: Job placement support, internship facilitation, and coaching",
        "Months 9-12: Post-placement follow-up, data collection, and program evaluation",
    ]
    for t in timeline_items:
        _bullet(doc, t)

    _para(
        doc,
        f"[PLACEHOLDER: Add any additional timeline milestones, cohort size details, or "
        f"partner organization roles specific to {program_name}.]",
        italic=True,
    )


def _build_goals_and_objectives(
    doc: Document,
    program_name: str,
    target_population: str,
    evaluation_metrics: list,
) -> None:
    _heading1(doc, "Goals and Objectives")

    _heading2(doc, "Program Goal")
    _para(
        doc,
        f"Increase economic independence and community integration outcomes for "
        f"{target_population} through the {program_name} program.",
    )

    _heading2(doc, "Measurable Objectives")
    if evaluation_metrics:
        for metric in evaluation_metrics:
            _bullet(doc, metric)
    else:
        default_objectives = [
            "At least 75% of enrolled participants will complete the full program",
            "At least 50% of program graduates will secure employment, internship, or post-secondary enrollment within 90 days of completion",
            "Participants will demonstrate measurable gains in workplace readiness skills as assessed by pre/post competency evaluation",
            "90-day job retention rate will meet or exceed 70% for placed participants",
        ]
        for obj in default_objectives:
            _bullet(doc, obj)


def _build_evaluation_plan(
    doc: Document,
    program_name: str,
    evaluation_metrics: list,
    org_name: str,
    funder_name: str = "[Funder]",
) -> None:
    _heading1(doc, "Evaluation Plan")

    _para(
        doc,
        f"{org_name} is committed to rigorous, ongoing program evaluation. "
        f"The following data collection methods will be used to track progress toward "
        f"the stated objectives for {program_name}:",
    )

    methods = [
        "Pre- and post-program competency assessments administered by program staff",
        "Participant attendance and engagement tracking logged in the organization's case management system",
        "90-day follow-up surveys or interviews with program graduates",
        "Employer feedback surveys collected within 60 days of participant placement",
        "Quarterly data review meetings with program staff to identify trends and adjust programming",
    ]
    for m in methods:
        _bullet(doc, m)

    _para(
        doc,
        f"Outcome data will be compiled into a funder impact report submitted within 30 days "
        f"of the grant period close. {org_name} will make this data available to {funder_name} "
        f"upon request and welcomes site visits to observe programming in action. "
        f"[PLACEHOLDER: Note any third-party evaluators, IRB protocols, or validated assessment "
        f"tools used in your specific program model.]",
        italic=False,
    )


def _build_org_capacity(
    doc: Document,
    org_name: str,
    org_mission: str,
    org_budget_total: Optional[str],
    program_name: str,
) -> None:
    _heading1(doc, "Organizational Capacity")

    capacity_text = (
        f"{org_name} has the demonstrated experience, infrastructure, and community "
        f"relationships to deliver {program_name} effectively. Our organization was founded "
        f"on the principle that {target_population_placeholder()} deserve access to the same "
        f"pathways to success available to all young people, and we have built our programming "
        f"model around evidence-informed practices in transition-age youth services."
    )
    _para(doc, capacity_text)

    track_record = (
        f"[PLACEHOLDER: Insert 2-3 specific accomplishments demonstrating your track record — "
        f"e.g., 'In our last program year, we served X participants, of whom Y% achieved Z outcome.' "
        f"Include any awards, accreditations, or notable partnerships that establish credibility.]"
    )
    _para(doc, track_record, italic=True)

    if org_budget_total:
        budget_para = doc.add_paragraph()
        _bold_label(
            budget_para,
            "Annual Organizational Budget",
            f"{org_budget_total} (audited financials available upon request)",
        )

    _para(
        doc,
        f"{org_name} is a 501(c)(3) nonprofit organization [PLACEHOLDER: EIN, year incorporated, "
        f"and state of registration]. We are governed by a [X]-member board of directors with "
        f"expertise in [PLACEHOLDER: relevant board expertise areas]. Our leadership team brings "
        f"[PLACEHOLDER: combined years and areas of relevant experience].",
    )


def target_population_placeholder() -> str:
    return "transition-age youth with disabilities"


def _build_budget_justification(
    doc: Document,
    grant_amount: str,
    program_name: str,
    org_name: str,
) -> None:
    _heading1(doc, "Budget Justification")

    budget_narrative = (
        f"The requested {grant_amount} will be deployed directly to support the core "
        f"operational components of {program_name}. {org_name} maintains a lean administrative "
        f"structure to maximize program investment. Personnel costs represent the largest portion "
        f"of the budget, reflecting the staffing intensity required to deliver individualized "
        f"services to {target_population_placeholder()}. All other line items are essential to "
        f"the program model and have been conservatively estimated based on prior program years. "
        f"[PLACEHOLDER: Add any cost-sharing, match sources, or in-kind contributions here.]"
    )
    _para(doc, budget_narrative)

    doc.add_paragraph()
    _add_budget_table(doc)

    _para(
        doc,
        "[PLACEHOLDER: Replace all $[AMOUNT] placeholders with actual figures from your program budget. "
        "Attach full budget spreadsheet as a required application attachment.]",
        italic=True,
    )


def _build_sustainability(
    doc: Document,
    program_name: str,
    org_name: str,
    funder_name: str,
    grant_amount: str,
) -> None:
    _heading1(doc, "Sustainability")

    sustainability_text = (
        f"While {org_name} is grateful for the opportunity to request support from "
        f"{funder_name}, we are committed to building a diversified funding base that "
        f"ensures {program_name} can continue to serve {target_population_placeholder()} "
        f"beyond the grant period."
    )
    _para(doc, sustainability_text)

    plan_text = (
        f"Our sustainability strategy includes: pursuit of government contracts and "
        f"public funding through state and local workforce development boards; cultivation "
        f"of individual major donors with a connection to our mission; earned revenue through "
        f"employer engagement and workforce partnership agreements; and applications to "
        f"additional foundations with aligned priorities."
    )
    _para(doc, plan_text)

    _para(
        doc,
        f"[PLACEHOLDER: Describe any multi-year funding commitments already secured, "
        f"earned revenue potential, or endowment plans specific to {program_name}. "
        f"Reference any matching funders or committed in-kind partners.]",
        italic=True,
    )


# ---------------------------------------------------------------------------
# LOI builder
# ---------------------------------------------------------------------------

def _build_loi(
    doc: Document,
    org_name: str,
    org_mission: str,
    funder_name: str,
    grant_amount: str,
    deadline: str,
    program_name: str,
    program_description: str,
    target_population: str,
    geographic_area: str,
    funding_request_summary: str,
    evaluation_metrics: list,
) -> None:
    """Build a condensed 1-page letter of inquiry."""

    _build_cover_page(doc, org_name, funder_name, grant_amount, deadline, program_name, "loi")
    doc.add_page_break()

    # Date and salutation
    today = datetime.date.today().strftime("%B %d, %Y")
    _para(doc, today)
    _para(doc, f"[Grants Manager Name]\n{funder_name}\n[Address]")
    doc.add_paragraph()

    salutation = doc.add_paragraph()
    salutation.add_run("Dear [Grants Manager Name],").bold = False

    doc.add_paragraph()

    # Opening paragraph (mission + request)
    opening = (
        f"{org_name} respectfully submits this letter of inquiry to {funder_name} "
        f"to request {grant_amount} in support of {program_name}. {org_mission}"
    )
    _para(doc, opening)

    # Statement of need
    need = (
        f"In {geographic_area}, {target_population} face disproportionately high rates of "
        f"unemployment and limited access to structured vocational development opportunities. "
        f"Existing resources are fragmented and often inaccessible to this population. "
        f"{program_name} is designed to close this gap."
    )
    _para(doc, need)

    # Program and request
    _para(doc, funding_request_summary)

    # Goals
    if evaluation_metrics:
        goals_para = doc.add_paragraph()
        goals_para.add_run("Key outcomes this grant will support:").bold = True
        for metric in evaluation_metrics:
            _bullet(doc, metric)

    # Close
    close_text = (
        f"We welcome the opportunity to discuss this proposal further and to provide any "
        f"additional information {funder_name} may require. Thank you for your consideration "
        f"of this request and for your commitment to supporting communities like ours."
    )
    _para(doc, close_text)

    doc.add_paragraph()
    _para(doc, "Sincerely,")
    doc.add_paragraph()
    _para(doc, "[Executive Director Name]\n[Title]\n{org_name}\n[Phone | Email]".format(org_name=org_name))

    _hr(doc)
    compliance = doc.add_paragraph()
    cr = compliance.add_run("GENERATED DRAFT — review with Development Director before submission.")
    cr.italic = True
    cr.font.color.rgb = RGBColor(0xAA, 0x44, 0x00)
    cr.font.size = Pt(9)


# ---------------------------------------------------------------------------
# Core render function
# ---------------------------------------------------------------------------

def render(
    funder_name: str,
    grant_amount: str,
    deadline: str,
    org_name: str,
    org_mission: str,
    program_name: str,
    program_description: str,
    target_population: str,
    geographic_area: str,
    funding_request_summary: str,
    evaluation_metrics: Optional[list] = None,
    org_budget_total: Optional[str] = None,
    proposal_type: Optional[str] = "full",
    additional_context: Optional[str] = None,
    output_dir: str = "/mnt/user-data/outputs",
) -> str:
    """
    Generate a grant proposal or LOI Word document.

    Parameters
    ----------
    funder_name : str
        Name of the funder or foundation
    grant_amount : str
        Dollar amount requested (e.g. "$50,000") or "TBD"
    deadline : str
        Application deadline (e.g. "June 15, 2026")
    org_name : str
        Full name of the applicant organization
    org_mission : str
        1-2 sentence mission statement
    program_name : str
        Name of the program this grant funds
    program_description : str
        Paragraph describing program activities and methodology
    target_population : str
        Who the program serves
    geographic_area : str
        Where the program operates
    funding_request_summary : str
        1-paragraph summary of what the grant will fund
    evaluation_metrics : list, optional
        Measurable outcome indicators
    org_budget_total : str, optional
        Total annual organizational budget
    proposal_type : str, optional
        "loi" or "full" (default "full")
    additional_context : str, optional
        Extra funder priorities or context
    output_dir : str
        Directory to save the output file

    Returns
    -------
    str
        Absolute path to the generated .docx file
    """
    # Validate required inputs
    required = {
        "funder_name": funder_name,
        "grant_amount": grant_amount,
        "deadline": deadline,
        "org_name": org_name,
        "org_mission": org_mission,
        "program_name": program_name,
        "program_description": program_description,
        "target_population": target_population,
        "geographic_area": geographic_area,
        "funding_request_summary": funding_request_summary,
    }
    for field, value in required.items():
        if not value or not str(value).strip():
            raise ValueError(f"{field} is required and cannot be empty")

    proposal_type = (proposal_type or "full").lower()
    if proposal_type not in ("loi", "full"):
        proposal_type = "full"

    evaluation_metrics = evaluation_metrics or []

    doc = Document()

    if proposal_type == "loi":
        _build_loi(
            doc,
            org_name=org_name,
            org_mission=org_mission,
            funder_name=funder_name,
            grant_amount=grant_amount,
            deadline=deadline,
            program_name=program_name,
            program_description=program_description,
            target_population=target_population,
            geographic_area=geographic_area,
            funding_request_summary=funding_request_summary,
            evaluation_metrics=evaluation_metrics,
        )
    else:
        # Full proposal
        _build_cover_page(doc, org_name, funder_name, grant_amount, deadline, program_name, "full")
        doc.add_page_break()

        _build_executive_summary(
            doc, org_name, org_mission, program_name, funding_request_summary, grant_amount, funder_name
        )
        _hr(doc)

        _build_statement_of_need(doc, org_name, target_population, geographic_area, program_name)
        _hr(doc)

        _build_program_description(doc, program_name, program_description, target_population, org_name)
        _hr(doc)

        _build_goals_and_objectives(doc, program_name, target_population, evaluation_metrics)
        _hr(doc)

        _build_evaluation_plan(doc, program_name, evaluation_metrics, org_name, funder_name)
        _hr(doc)

        _build_org_capacity(doc, org_name, org_mission, org_budget_total, program_name)
        _hr(doc)

        _build_budget_justification(doc, grant_amount, program_name, org_name)
        _hr(doc)

        _build_sustainability(doc, program_name, org_name, funder_name, grant_amount)

        if additional_context:
            doc.add_page_break()
            _heading1(doc, "Additional Information")
            _para(doc, additional_context)

    # Save
    os.makedirs(output_dir, exist_ok=True)
    timestamp = int(time.time())
    safe_funder = re.sub(r"[^\w]+", "_", funder_name)[:30].strip("_")
    filename = f"grant_proposal_{safe_funder}_{timestamp}.docx"
    out_path = os.path.join(output_dir, filename)
    doc.save(out_path)

    return out_path
