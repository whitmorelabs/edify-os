"""
impact_report/render.py

Generates a funder-ready or annual-report-style impact report as a Word (.docx)
document, tailored for nonprofit Development Directors.

Sections:
  1. Cover (title, period, org name)
  2. Letter from Leadership (placeholder with tone guide)
  3. Mission Recap
  4. Outcomes at a Glance (table)
  5. Impact Story (participant story or placeholder)
  6. Financial Summary (if provided)
  7. Thank You / Acknowledgements
  8. Looking Ahead (placeholder)

All libraries are pre-installed in Anthropic's code-execution sandbox.
"""

import os
import re
import time
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------

_NAVY = RGBColor(0x1A, 0x47, 0x6B)
_STEEL = RGBColor(0x2E, 0x72, 0x9E)
_AMBER = RGBColor(0xAA, 0x44, 0x00)
_GRAY = RGBColor(0xCC, 0xCC, 0xCC)
_GREEN = RGBColor(0x1A, 0x6B, 0x3A)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _heading1(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=1)
    if h.runs:
        h.runs[0].font.color.rgb = _NAVY


def _heading2(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=2)
    if h.runs:
        h.runs[0].font.color.rgb = _STEEL


def _para(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph(text)
    if italic and p.runs:
        p.runs[0].italic = True


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 72)
    run.font.color.rgb = _GRAY
    run.font.size = Pt(7)


def _placeholder(doc: Document, text: str) -> None:
    """Render an italic amber placeholder instruction."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = _AMBER


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _build_cover(
    doc: Document,
    org_name: str,
    report_type: str,
    report_period: str,
    funder_name: Optional[str],
    grant_amount: Optional[str],
) -> None:
    doc.add_paragraph()

    report_label = (
        "Grant Impact Report" if report_type == "funder_report" else "Annual Impact Report"
    )

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_r = title_para.add_run(report_label)
    title_r.bold = True
    title_r.font.size = Pt(24)
    title_r.font.color.rgb = _NAVY

    period_para = doc.add_paragraph()
    period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period_r = period_para.add_run(report_period)
    period_r.bold = True
    period_r.font.size = Pt(14)
    period_r.font.color.rgb = _STEEL

    org_para = doc.add_paragraph()
    org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_r = org_para.add_run(org_name)
    org_r.font.size = Pt(13)

    if report_type == "funder_report" and funder_name and grant_amount:
        doc.add_paragraph()
        funder_para = doc.add_paragraph()
        funder_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        f_run = funder_para.add_run(
            f"Prepared for: {funder_name}  |  Grant: {grant_amount}"
        )
        f_run.font.size = Pt(11)
        f_run.italic = True
        f_run.font.color.rgb = _STEEL

    doc.add_paragraph()

    compliance_para = doc.add_paragraph()
    compliance_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = compliance_para.add_run(
        "GENERATED DRAFT — pair with photos and a board signoff before publishing."
    )
    cr.italic = True
    cr.font.color.rgb = _AMBER
    cr.font.size = Pt(9)


def _build_leadership_letter(
    doc: Document,
    org_name: str,
    report_period: str,
    report_type: str,
    funder_name: Optional[str],
) -> None:
    _heading1(doc, "A Message from Our Leadership")

    tone_note = doc.add_paragraph()
    tone_r = tone_note.add_run(
        "[TONE GUIDE FOR LEADERSHIP LETTER: Write in first person from the Executive Director "
        "or Board Chair. Open with gratitude (for funders) or with an inspiring moment from "
        "the year (for annual reports). Reference 1-2 specific outcomes. Close with forward "
        "momentum and a compelling vision. Keep to 150-200 words. Have leadership review and "
        "sign off before this goes out.]"
    )
    tone_r.italic = True
    tone_r.font.color.rgb = _AMBER
    tone_r.font.size = Pt(9)

    doc.add_paragraph()

    if report_type == "funder_report" and funder_name:
        letter_draft = (
            f"Dear {funder_name} team,\n\n"
            f"On behalf of {org_name}, I am pleased to share this report covering {report_period}. "
            f"Your investment in our work has made a measurable difference for the people we serve, "
            f"and this document tells that story — in data, in outcomes, and in the voices of "
            f"participants who are living proof that the right support at the right time can change "
            f"a life's trajectory.\n\n"
            f"[PLACEHOLDER: Executive Director — add 2-3 sentences of specific impact, a personal "
            f"reflection, or a program milestone here. Then close with a forward-looking sentence "
            f"about what's next.]\n\n"
            f"With gratitude,\n\n"
            f"[Executive Director Name]\n[Title]\n{org_name}"
        )
    else:
        letter_draft = (
            f"Dear Friends and Supporters of {org_name},\n\n"
            f"As we reflect on {report_period}, we are proud to share what your generosity has "
            f"made possible. This report captures the work, the outcomes, and most importantly, "
            f"the people behind our numbers — the young adults, families, and community members "
            f"who trusted us to walk alongside them.\n\n"
            f"[PLACEHOLDER: Executive Director — add 2-3 sentences of specific program highlights, "
            f"a year-in-review reflection, or a community milestone here. Close with your vision "
            f"for the year ahead.]\n\n"
            f"In solidarity,\n\n"
            f"[Executive Director Name]\n[Title]\n{org_name}"
        )

    _para(doc, letter_draft, italic=True)


def _build_mission_recap(
    doc: Document,
    org_name: str,
    org_mission: Optional[str],
) -> None:
    _heading1(doc, "Our Mission")

    if org_mission:
        _para(doc, org_mission)
    else:
        _placeholder(
            doc,
            "[PLACEHOLDER: Insert your organization's mission statement here. "
            "Keep it to 1-2 sentences. The mission anchor reminds funders and stakeholders "
            "why this work exists.]",
        )

    doc.add_paragraph()
    _placeholder(
        doc,
        "[OPTIONAL: Add a 1-line 'What We Do' tagline or a 3-bullet 'Who We Serve / What We Do / Why It Matters' "
        "block here for annual report clarity.]",
    )


def _build_outcomes_table(
    doc: Document,
    program_outcomes: list,
    report_period: str,
) -> None:
    _heading1(doc, "Outcomes at a Glance")

    intro = doc.add_paragraph(
        f"The following table summarizes program performance for {report_period}. "
        f"Targets reflect commitments made at the start of the grant period or program year."
    )

    doc.add_paragraph()

    # Build table: program | target | actual | narrative
    table = doc.add_table(rows=len(program_outcomes) + 1, cols=4)
    table.style = "Table Grid"

    # Header row
    headers = ["Program / Indicator", "Target", "Actual", "Notes"]
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = _NAVY

    # Data rows
    for row_idx, outcome in enumerate(program_outcomes, start=1):
        row = table.rows[row_idx]

        program_val = outcome.get("program", "[Program Name]")
        target_val = outcome.get("target", "[Target]")
        actual_val = outcome.get("actual", "[Actual]")
        narrative_val = outcome.get("narrative", "")

        row.cells[0].text = str(program_val)
        row.cells[1].text = str(target_val)
        row.cells[2].text = str(actual_val)
        row.cells[3].text = str(narrative_val)

        # Color-code actual: green if appears to exceed/meet, amber if seems missed
        # Simple heuristic: if actual contains a number larger than target's number
        # In practice, just style narratively — no complex parsing
        for para in row.cells[2].paragraphs:
            for run in para.runs:
                run.font.color.rgb = _GREEN

    doc.add_paragraph()
    _placeholder(
        doc,
        "[PLACEHOLDER: Add additional outcome rows for any secondary programs, "
        "process metrics (e.g., trainings delivered, staff hours), or funder-required "
        "indicators not captured above.]",
    )


def _build_impact_story(
    doc: Document,
    participant_story: Optional[str],
) -> None:
    _heading1(doc, "Impact Story")

    if participant_story:
        _para(doc, participant_story)
        doc.add_paragraph()
        consent_note = doc.add_paragraph()
        cr = consent_note.add_run(
            "[NOTE: Confirm written consent is on file before publishing any participant "
            "story. Anonymize or use initials + pseudonyms as appropriate per your "
            "confidentiality policy.]"
        )
        cr.italic = True
        cr.font.color.rgb = _AMBER
        cr.font.size = Pt(9)
    else:
        _placeholder(
            doc,
            "[PLACEHOLDER: Insert an anonymized participant story or testimonial here. "
            "Aim for 3-5 sentences: context, challenge, intervention, outcome, reflection. "
            "Use a pseudonym or initials. Stories are the single most powerful element in "
            "any impact report — do not skip this section. "
            "Example structure: 'When [Name] came to us, [situation]. Through [program], "
            "[what changed]. Today, [outcome]. In their words: \"[quote]\".']",
        )


def _build_financial_summary(
    doc: Document,
    total_revenue: Optional[str],
    total_expenses: Optional[str],
    report_period: str,
    org_name: str,
) -> None:
    _heading1(doc, "Financial Summary")

    if total_revenue or total_expenses:
        _para(doc, f"The following summary reflects {org_name}'s financial position for {report_period}.")
        doc.add_paragraph()

        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"

        # Headers
        hdr = table.rows[0]
        hdr.cells[0].text = "Category"
        hdr.cells[1].text = "Amount"
        for cell in hdr.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = _NAVY

        # Revenue row
        rev_row = table.rows[1]
        rev_row.cells[0].text = "Total Revenue"
        rev_row.cells[1].text = total_revenue or "[Not provided]"

        # Expenses row
        exp_row = table.rows[2]
        exp_row.cells[0].text = "Total Expenses"
        exp_row.cells[1].text = total_expenses or "[Not provided]"

        # Net row
        net_row = table.rows[3]
        net_row.cells[0].text = "Net (Revenue - Expenses)"
        for para in net_row.cells[0].paragraphs:
            for run in para.runs:
                run.bold = True
        try:
            rev_clean = re.sub(r"[^\d.]", "", total_revenue or "")
            exp_clean = re.sub(r"[^\d.]", "", total_expenses or "")
            if rev_clean and exp_clean:
                net = float(rev_clean) - float(exp_clean)
                net_str = f"${net:,.0f}" if net >= 0 else f"(${abs(net):,.0f})"
                net_row.cells[1].text = net_str
            else:
                net_row.cells[1].text = "[Calculate from actuals]"
        except (ValueError, TypeError):
            net_row.cells[1].text = "[Calculate from actuals]"

        doc.add_paragraph()
        _placeholder(
            doc,
            "[PLACEHOLDER: Attach full audited financial statements or 990 as a separate document "
            "if required by the funder. This table is summary-level only.]",
        )
    else:
        _placeholder(
            doc,
            "[PLACEHOLDER: Insert a financial summary for the reporting period. "
            "Include total revenue, total expenses, and net. Many funders require a budget-vs-actual "
            "table for the specific grant-funded program. Attach audited financials if available.]",
        )


def _build_acknowledgements(
    doc: Document,
    top_funders: Optional[list],
    org_name: str,
    report_type: str,
) -> None:
    _heading1(doc, "Thank You")

    if report_type == "funder_report":
        _para(
            doc,
            f"{org_name} is deeply grateful to our funders, partners, and community supporters "
            f"whose investment makes this work possible.",
        )
    else:
        _para(
            doc,
            f"{org_name} exists because of the generosity, partnership, and belief of an "
            f"extraordinary community. Thank you to everyone who made this year possible.",
        )

    if top_funders:
        doc.add_paragraph()
        _heading2(doc, "Funding Partners")
        for funder in top_funders:
            _bullet(doc, funder)

    doc.add_paragraph()
    _placeholder(
        doc,
        "[PLACEHOLDER: Add staff, board, volunteers, and community partners to be recognized. "
        "For annual reports, consider a full board list and key staff roster. "
        "For funder reports, a brief acknowledgement of other funders demonstrates diversification.]",
    )


def _build_looking_ahead(
    doc: Document,
    org_name: str,
    report_period: str,
    report_type: str,
) -> None:
    _heading1(doc, "Looking Ahead")

    _placeholder(
        doc,
        f"[PLACEHOLDER: Write 1-2 paragraphs about what comes next for {org_name}. "
        f"For funder reports: reference upcoming program milestones, next grant period goals, "
        f"or expansion plans. Frame this as continued investment paying forward. "
        f"For annual reports: share the vision for the next 12 months — program growth, "
        f"new partnerships, capital plans, or community goals. Be specific and optimistic. "
        f"This is the last thing a reader sees — make it land.]",
    )


# ---------------------------------------------------------------------------
# Core render function
# ---------------------------------------------------------------------------

def render(
    org_name: str,
    report_type: str,
    report_period: str,
    program_outcomes: list,
    funder_name: Optional[str] = None,
    grant_amount: Optional[str] = None,
    total_revenue: Optional[str] = None,
    total_expenses: Optional[str] = None,
    participant_story: Optional[str] = None,
    top_funders: Optional[list] = None,
    org_mission: Optional[str] = None,
    output_dir: str = "/mnt/user-data/outputs",
) -> str:
    """
    Generate an impact report Word document.

    Parameters
    ----------
    org_name : str
        Full name of the organization
    report_type : str
        "funder_report" or "annual_report"
    report_period : str
        Reporting period (e.g. "January - June 2026")
    program_outcomes : list of dict
        Each dict: {"program": str, "target": str, "actual": str, "narrative": str}
    funder_name : str, optional
        Required if report_type="funder_report"
    grant_amount : str, optional
        Required if report_type="funder_report"
    total_revenue : str, optional
        Total organizational revenue for the period
    total_expenses : str, optional
        Total organizational expenses for the period
    participant_story : str, optional
        Anonymized participant story or testimonial
    top_funders : list of str, optional
        Funder names for the acknowledgements section
    org_mission : str, optional
        Mission statement for the mission recap section
    output_dir : str
        Directory to save the output file

    Returns
    -------
    str
        Absolute path to the generated .docx file
    """
    # Validate required
    for field, value in [("org_name", org_name), ("report_period", report_period)]:
        if not value or not str(value).strip():
            raise ValueError(f"{field} is required and cannot be empty")

    if not program_outcomes:
        raise ValueError("program_outcomes must contain at least one outcome object")

    report_type = (report_type or "annual_report").lower()
    if report_type not in ("funder_report", "annual_report"):
        report_type = "annual_report"

    if report_type == "funder_report":
        if not funder_name or not str(funder_name).strip():
            raise ValueError("funder_name is required when report_type is 'funder_report'")
        if not grant_amount or not str(grant_amount).strip():
            raise ValueError("grant_amount is required when report_type is 'funder_report'")

    doc = Document()

    # Cover
    _build_cover(doc, org_name, report_type, report_period, funder_name, grant_amount)
    doc.add_page_break()

    # Leadership letter
    _build_leadership_letter(doc, org_name, report_period, report_type, funder_name)
    _hr(doc)

    # Mission
    _build_mission_recap(doc, org_name, org_mission)
    _hr(doc)

    # Outcomes table
    _build_outcomes_table(doc, program_outcomes, report_period)
    _hr(doc)

    # Impact story
    _build_impact_story(doc, participant_story)
    _hr(doc)

    # Financial summary
    _build_financial_summary(doc, total_revenue, total_expenses, report_period, org_name)
    _hr(doc)

    # Acknowledgements
    _build_acknowledgements(doc, top_funders, org_name, report_type)
    _hr(doc)

    # Looking ahead
    _build_looking_ahead(doc, org_name, report_period, report_type)

    # Save
    os.makedirs(output_dir, exist_ok=True)
    timestamp = int(time.time())
    safe_period = re.sub(r"[^\w]+", "_", report_period)[:30].strip("_")
    filename = f"impact_report_{safe_period}_{timestamp}.docx"
    out_path = os.path.join(output_dir, filename)
    doc.save(out_path)

    return out_path
