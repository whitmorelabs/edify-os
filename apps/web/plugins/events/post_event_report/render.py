"""
post_event_report/render.py

Generates a post-event debrief / wrap-up report as a Word (.docx) document
for nonprofit galas, fundraisers, and programs.

Sections:
  1. Cover (event name, date, report date, org name)
  2. Executive Summary (attendance + revenue + ROI narrative)
  3. By the Numbers (table — attendance, revenue, costs, net, donor count, ROI)
  4. Cost Breakdown (table if provided, else placeholder)
  5. Survey Highlights (if provided)
  6. What Worked (bullet list)
  7. What Didn't Work (bullet list)
  8. Recommendations for Next Year (bullet list)
  9. Appendix: Vendor + Sponsor List (placeholder)

Auto-computes:
  - Net Revenue = revenue_raised - costs_total
  - ROI % = (revenue_raised - costs_total) / costs_total * 100

All libraries are pre-installed in Anthropic's code-execution sandbox.
"""

import datetime
import os
import re
import time as _time
from typing import Optional, Union

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------

_NAVY = RGBColor(0x1A, 0x47, 0x6B)
_STEEL = RGBColor(0x2E, 0x72, 0x9E)
_GREEN = RGBColor(0x1A, 0x7A, 0x3C)
_RED = RGBColor(0xAA, 0x22, 0x00)
_GRAY = RGBColor(0xCC, 0xCC, 0xCC)
_ORANGE = RGBColor(0xAA, 0x44, 0x00)
_DARK_GRAY = RGBColor(0x55, 0x55, 0x55)


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _heading1(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=1)
    if h.runs:
        h.runs[0].font.color.rgb = _NAVY


def _heading2(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=2)
    if h.runs:
        h.runs[0].font.color.rgb = _STEEL


def _para(doc: Document, text: str, italic: bool = False, bold: bool = False,
          align: str = "left", size: Optional[int] = None,
          color: Optional[RGBColor] = None) -> None:
    p = doc.add_paragraph(text)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.runs:
        run = p.runs[0]
        if italic:
            run.italic = True
        if bold:
            run.bold = True
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _bold_label(para, label: str, value: str) -> None:
    r = para.add_run(label + ": ")
    r.bold = True
    para.add_run(value)


def _hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 72)
    run.font.color.rgb = _GRAY
    run.font.size = Pt(7)


def _set_cell_shading(cell, fill_hex: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _parse_amount(value) -> float:
    """Parse a dollar amount from string or number."""
    if isinstance(value, (int, float)):
        return float(value)
    cleaned = str(value).replace(",", "").replace("$", "").strip()
    try:
        return float(cleaned)
    except ValueError:
        return 0.0


def _fmt_currency(value: float) -> str:
    """Format a float as $X,XXX."""
    return f"${value:,.0f}"


def _fmt_pct(value: float) -> str:
    """Format a float as X.X%."""
    return f"{value:.1f}%"


# ---------------------------------------------------------------------------
# Financial computations
# ---------------------------------------------------------------------------

def _compute_financials(revenue_raised, costs_total):
    revenue = _parse_amount(revenue_raised)
    costs = _parse_amount(costs_total)
    net = revenue - costs
    roi = (net / costs * 100) if costs > 0 else 0.0
    return revenue, costs, net, roi


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _build_cover(
    doc: Document,
    event_name: str,
    event_date: str,
    org_name: str,
) -> None:
    today = datetime.date.today().strftime("%B %d, %Y")

    doc.add_paragraph()
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title_para.add_run("Post-Event Report")
    t_run.bold = True
    t_run.font.size = Pt(20)
    t_run.font.color.rgb = _NAVY

    event_para = doc.add_paragraph()
    event_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    e_run = event_para.add_run(event_name)
    e_run.bold = True
    e_run.font.size = Pt(15)
    e_run.font.color.rgb = _STEEL

    doc.add_paragraph()

    for label, value in [
        ("Event Date", event_date),
        ("Report Date", today),
        ("Prepared by", org_name),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bold_label(p, label, value)

    doc.add_paragraph()

    notice_para = doc.add_paragraph()
    notice_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = notice_para.add_run(
        "INTERNAL DOCUMENT — verify all figures against actuals before sharing externally."
    )
    nr.italic = True
    nr.font.color.rgb = _ORANGE
    nr.font.size = Pt(9)


def _build_executive_summary(
    doc: Document,
    event_name: str,
    event_date: str,
    org_name: str,
    attendance: int,
    revenue: float,
    costs: float,
    net: float,
    roi: float,
    donor_acquisition_count: Optional[int],
) -> None:
    _heading1(doc, "Executive Summary")

    roi_comment = "a strong return on event investment" if roi >= 100 else "a positive event return"
    if roi < 0:
        roi_comment = "a net loss on direct costs — see recommendations for cost management opportunities"

    donor_clause = ""
    if donor_acquisition_count:
        donor_clause = f" The event also acquired {donor_acquisition_count} net new donors, expanding the organization's pipeline for future cultivation."

    summary = (
        f"{org_name} hosted {event_name} on {event_date} to {attendance:,} attendees. "
        f"The event raised {_fmt_currency(revenue)} in total revenue against {_fmt_currency(costs)} "
        f"in direct costs, generating a net of {_fmt_currency(net)} and an ROI of {_fmt_pct(roi)} — "
        f"{roi_comment}.{donor_clause} "
        f"This report summarizes financial outcomes, attendee feedback, operational highlights, "
        f"and recommendations for improving future events."
    )
    _para(doc, summary)


def _build_by_the_numbers(
    doc: Document,
    attendance: int,
    revenue: float,
    costs: float,
    net: float,
    roi: float,
    donor_acquisition_count: Optional[int],
) -> None:
    _heading1(doc, "By the Numbers")

    rows_data = [
        ("Metric", "Value"),
        ("Total Attendance", f"{attendance:,} guests"),
        ("Total Revenue Raised", _fmt_currency(revenue)),
        ("Total Event Costs", _fmt_currency(costs)),
        ("Net Revenue", _fmt_currency(net)),
        ("ROI", _fmt_pct(roi)),
    ]
    if donor_acquisition_count is not None:
        rows_data.append(("Net New Donors Acquired", str(donor_acquisition_count)))

    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = "Table Grid"

    for row_idx, (label, value) in enumerate(rows_data):
        cells = table.rows[row_idx].cells
        if row_idx == 0:
            # Header row
            cells[0].text = ""
            r1 = cells[0].paragraphs[0].add_run(label)
            r1.bold = True
            r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cells[1].text = ""
            r2 = cells[1].paragraphs[0].add_run(value)
            r2.bold = True
            r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_shading(cells[0], "1A476B")
            _set_cell_shading(cells[1], "1A476B")
        else:
            cells[0].text = ""
            r1 = cells[0].paragraphs[0].add_run(label)
            r1.bold = True

            cells[1].text = ""
            r2 = cells[1].paragraphs[0].add_run(value)

            # Colour code net and ROI
            if label == "Net Revenue":
                r2.font.color.rgb = _GREEN if net >= 0 else _RED
                r2.bold = True
            elif label == "ROI":
                r2.font.color.rgb = _GREEN if roi >= 0 else _RED
                r2.bold = True

            shade = "FFFFFF" if row_idx % 2 == 1 else "F0F6FB"
            _set_cell_shading(cells[0], shade)
            _set_cell_shading(cells[1], shade)


def _build_cost_breakdown(doc: Document, cost_breakdown: Optional[list], costs_total: float) -> None:
    _heading1(doc, "Cost Breakdown")

    if not cost_breakdown:
        _para(
            doc,
            "[PLACEHOLDER: Add itemized cost breakdown here — e.g., venue, catering, AV, printing, "
            "entertainment, staffing. Include actual vs. budgeted amounts where available.]",
            italic=True,
        )
        return

    rows_data = [("Category", "Amount")]
    running_total = 0.0
    for item in cost_breakdown:
        cat = str(item.get("category", ""))
        amt = _parse_amount(item.get("amount", 0))
        running_total += amt
        rows_data.append((cat, _fmt_currency(amt)))
    rows_data.append(("TOTAL", _fmt_currency(costs_total)))

    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = "Table Grid"

    for row_idx, (label, value) in enumerate(rows_data):
        cells = table.rows[row_idx].cells
        is_header = row_idx == 0
        is_total = label == "TOTAL"

        cells[0].text = ""
        r1 = cells[0].paragraphs[0].add_run(label)
        cells[1].text = ""
        r2 = cells[1].paragraphs[0].add_run(value)

        if is_header:
            r1.bold = True
            r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r2.bold = True
            r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_shading(cells[0], "1A476B")
            _set_cell_shading(cells[1], "1A476B")
        elif is_total:
            r1.bold = True
            r2.bold = True
            _set_cell_shading(cells[0], "D6E8F5")
            _set_cell_shading(cells[1], "D6E8F5")
        else:
            shade = "FFFFFF" if row_idx % 2 == 1 else "F0F6FB"
            _set_cell_shading(cells[0], shade)
            _set_cell_shading(cells[1], shade)


def _build_survey_highlights(doc: Document, survey_results) -> None:
    _heading1(doc, "Survey Highlights")

    if not survey_results:
        return

    if isinstance(survey_results, str):
        _para(doc, survey_results)
    elif isinstance(survey_results, list):
        for item in survey_results:
            if isinstance(item, dict):
                question = str(item.get("question", ""))
                summary = str(item.get("summary", ""))
                p = doc.add_paragraph()
                if question:
                    r = p.add_run(question + ": ")
                    r.bold = True
                p.add_run(summary)
            else:
                _para(doc, str(item))
    else:
        _para(doc, str(survey_results))


def _build_bullet_section(doc: Document, heading: str, items: Optional[list], placeholder: str) -> None:
    _heading1(doc, heading)
    if items:
        for item in items:
            _bullet(doc, str(item))
    else:
        _para(doc, f"[PLACEHOLDER: {placeholder}]", italic=True)


def _build_appendix(doc: Document) -> None:
    _heading1(doc, "Appendix: Vendor + Sponsor List")
    _para(
        doc,
        "[PLACEHOLDER: List all vendors (venue, catering, AV, entertainment, printing, transportation) "
        "with contact name, company name, and contract amount. List all event sponsors with tier level and "
        "confirmed payment received. This appendix supports accounting reconciliation and future vendor sourcing.]",
        italic=True,
    )


# ---------------------------------------------------------------------------
# Core render function
# ---------------------------------------------------------------------------

def render(
    event_name: str,
    event_date: str,
    org_name: str,
    attendance: int,
    revenue_raised: Union[float, int, str],
    costs_total: Union[float, int, str],
    cost_breakdown: Optional[list] = None,
    survey_results=None,
    donor_acquisition_count: Optional[int] = None,
    what_worked: Optional[list] = None,
    what_didnt_work: Optional[list] = None,
    next_year_recommendations: Optional[list] = None,
    output_dir: str = "/mnt/user-data/outputs",
) -> str:
    """
    Generate a post-event debrief report as a Word document.

    Parameters
    ----------
    event_name : str
        Name of the event
    event_date : str
        Date of the event
    org_name : str
        Full name of the organization
    attendance : int
        Number of attendees
    revenue_raised : float | int | str
        Total revenue raised (tickets + donations + sponsorships)
    costs_total : float | int | str
        Total event costs
    cost_breakdown : list of dict, optional
        Each dict: {category, amount}
    survey_results : str or list of dict, optional
        Survey feedback — freeform string or list of {question, summary}
    donor_acquisition_count : int, optional
        Net new donors from the event
    what_worked : list of str, optional
        Things that went well
    what_didnt_work : list of str, optional
        Things to improve
    next_year_recommendations : list of str, optional
        Suggestions for future events
    output_dir : str
        Directory to save the output file

    Returns
    -------
    str
        Absolute path to the generated .docx file
    """
    required = {
        "event_name": event_name,
        "event_date": event_date,
        "org_name": org_name,
    }
    for field, value in required.items():
        if not value or not str(value).strip():
            raise ValueError(f"{field} is required and cannot be empty")

    if attendance is None:
        raise ValueError("attendance is required")
    if revenue_raised is None:
        raise ValueError("revenue_raised is required")
    if costs_total is None:
        raise ValueError("costs_total is required")

    revenue, costs, net, roi = _compute_financials(revenue_raised, costs_total)

    os.makedirs(output_dir, exist_ok=True)
    safe_name = re.sub(r"[^\w]+", "_", event_name)[:30].strip("_")
    safe_date = re.sub(r"[^\w]+", "_", event_date)[:15].strip("_")
    filename = f"post_event_report_{safe_name}_{safe_date}.docx"
    out_path = os.path.join(output_dir, filename)

    doc = Document()

    _build_cover(doc, event_name, event_date, org_name)
    doc.add_page_break()

    _build_executive_summary(
        doc, event_name, event_date, org_name, int(attendance),
        revenue, costs, net, roi, donor_acquisition_count
    )
    _hr(doc)

    _build_by_the_numbers(doc, int(attendance), revenue, costs, net, roi, donor_acquisition_count)
    _hr(doc)

    _build_cost_breakdown(doc, cost_breakdown, costs)
    _hr(doc)

    if survey_results:
        _build_survey_highlights(doc, survey_results)
        _hr(doc)

    _build_bullet_section(
        doc, "What Worked", what_worked,
        "List what went well — setup, programming, vendor performance, attendee experience, etc."
    )
    _hr(doc)

    _build_bullet_section(
        doc, "What Didn't Work", what_didnt_work,
        "List what to improve — logistics issues, timeline overruns, vendor problems, budget overages, etc."
    )
    _hr(doc)

    _build_bullet_section(
        doc, "Recommendations for Next Year", next_year_recommendations,
        "List specific, actionable changes for the next event — process improvements, timeline shifts, budget adjustments, new ideas."
    )
    _hr(doc)

    _build_appendix(doc)

    doc.save(out_path)
    return out_path
