"""
sponsor_package/render.py

Generates a sponsorship prospectus + 3 outreach email templates as a single
Word (.docx) document for nonprofit galas, fundraisers, and programs.

Sections:
  1. Cover page
  2. About the Event
  3. About the Organization
  4. Why Sponsor
  5. Sponsor Tiers (table)
  6. Logo + Recognition
  7. Contact / Next Steps
  --- page break ---
  8. Outreach Email Templates (cold, warm, last-chance)

All libraries are pre-installed in Anthropic's code-execution sandbox.
"""

import os
import re
import time as _time
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------

_NAVY = RGBColor(0x1A, 0x47, 0x6B)
_STEEL = RGBColor(0x2E, 0x72, 0x9E)
_GOLD = RGBColor(0xB8, 0x86, 0x00)
_GRAY = RGBColor(0xCC, 0xCC, 0xCC)
_ORANGE = RGBColor(0xAA, 0x44, 0x00)
_DARK_GRAY = RGBColor(0x55, 0x55, 0x55)


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _heading1(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=1)
    if h.runs:
        h.runs[0].font.color.rgb = _NAVY


def _heading2(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=2)
    if h.runs:
        h.runs[0].font.color.rgb = _STEEL


def _para(doc: Document, text: str, italic: bool = False, bold: bool = False,
          align: str = "left", size: Optional[int] = None,
          color: Optional[RGBColor] = None) -> None:
    p = doc.add_paragraph(text)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.runs:
        run = p.runs[0]
        if italic:
            run.italic = True
        if bold:
            run.bold = True
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _bold_label(para, label: str, value: str) -> None:
    r = para.add_run(label + ": ")
    r.bold = True
    para.add_run(value)


def _hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 72)
    run.font.color.rgb = _GRAY
    run.font.size = Pt(7)


def _set_cell_shading(cell, fill_hex: str) -> None:
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _format_currency(amount) -> str:
    """Format a numeric amount as $X,XXX."""
    try:
        val = int(float(str(amount).replace(",", "").replace("$", "")))
        return f"${val:,}"
    except (ValueError, TypeError):
        return str(amount)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _build_cover(
    doc: Document,
    event_name: str,
    event_date: str,
    org_name: str,
    fundraising_goal: str,
) -> None:
    doc.add_paragraph()
    doc.add_paragraph()

    # Event name
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(event_name)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = _NAVY

    # Subtitle
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("Sponsorship Opportunities")
    sub_run.bold = True
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = _STEEL

    doc.add_paragraph()

    # Details
    for label, value in [("Date", event_date), ("Organization", org_name), ("Fundraising Goal", fundraising_goal)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bold_label(p, label, value)

    doc.add_paragraph()

    # Logo placeholder
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_para.add_run("[INSERT ORGANIZATION LOGO HERE]")
    logo_run.italic = True
    logo_run.font.color.rgb = _DARK_GRAY
    logo_run.font.size = Pt(9)

    doc.add_paragraph()

    # Draft notice
    notice_para = doc.add_paragraph()
    notice_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = notice_para.add_run(
        "GENERATED DRAFT — add logo, confirm tier amounts, and have leadership review before sending."
    )
    nr.italic = True
    nr.font.color.rgb = _ORANGE
    nr.font.size = Pt(9)


def _build_about_event(
    doc: Document,
    event_name: str,
    event_date: str,
    venue: Optional[str],
    org_name: str,
    event_purpose: str,
    fundraising_goal: str,
    attendee_estimate,
) -> None:
    _heading1(doc, "About the Event")
    _para(doc, event_purpose)

    doc.add_paragraph()
    key_facts_para = doc.add_paragraph()
    key_facts_run = key_facts_para.add_run("Key Event Details")
    key_facts_run.bold = True
    key_facts_run.font.color.rgb = _NAVY

    facts = [("Event", event_name), ("Date", event_date)]
    if venue:
        facts.append(("Venue", venue))
    if attendee_estimate:
        facts.append(("Expected Attendance", str(attendee_estimate) + " guests"))
    facts.append(("Fundraising Goal", fundraising_goal))
    facts.append(("Presented by", org_name))

    for label, value in facts:
        p = doc.add_paragraph()
        _bold_label(p, label, value)


def _build_about_org(doc: Document, org_name: str, org_mission: str) -> None:
    _heading1(doc, f"About {org_name}")
    _para(doc, org_mission)
    doc.add_paragraph()
    _para(
        doc,
        "[PLACEHOLDER: Insert 2-3 impact statistics here — e.g., '200+ youth served annually', "
        "'85% employment rate among graduates', 'Operating since [year] in [city/region]'. "
        "Replace with actual organization data before sending.]",
        italic=True,
    )


def _build_why_sponsor(doc: Document, event_name: str, org_name: str) -> None:
    _heading1(doc, "Why Sponsor?")
    _para(
        doc,
        f"Sponsoring {event_name} is an opportunity to align your brand with a mission that matters. "
        f"Here is what your investment delivers:"
    )
    doc.add_paragraph()
    why_items = [
        f"Visibility — Reach {org_name}'s network of supporters, community leaders, and elected officials. "
        "Your brand will appear across event materials, social media, and post-event communications reaching hundreds of engaged stakeholders.",
        f"Mission Alignment — Demonstrate your organization's commitment to the community by investing in programs "
        f"that create real, measurable change. Sponsoring {org_name} is not just a donation — it is a statement about your values.",
        "Community — Join a growing coalition of corporate and community partners who believe in building a stronger, "
        "more equitable community. Your sponsorship helps sustain programs that would not exist without this support.",
    ]
    for item in why_items:
        _bullet(doc, item)


def _build_tiers_table(doc: Document, sponsor_tiers: list) -> None:
    _heading1(doc, "Sponsorship Tiers")

    if not sponsor_tiers:
        _para(doc, "[PLACEHOLDER: Add sponsor tier details here.]", italic=True)
        return

    # Table: Tier | Amount | Benefits
    col_count = 3
    table = doc.add_table(rows=len(sponsor_tiers) + 1, cols=col_count)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(["Sponsorship Tier", "Investment", "Benefits"]):
        cell = hdr_cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(cell, "1A476B")

    # Tier rows — alternate shading
    alt_colors = ["FFFFFF", "D6E8F5"]
    for row_idx, tier in enumerate(sponsor_tiers):
        row_cells = table.rows[row_idx + 1].cells
        shade = alt_colors[row_idx % 2]

        # Tier name
        row_cells[0].text = ""
        name_run = row_cells[0].paragraphs[0].add_run(str(tier.get("name", "")))
        name_run.bold = True

        # Amount
        row_cells[1].text = ""
        amt_val = _format_currency(tier.get("amount", ""))
        amt_run = row_cells[1].paragraphs[0].add_run(amt_val)
        amt_run.bold = True

        # Benefits
        row_cells[2].text = ""
        benefits = tier.get("benefits", [])
        if isinstance(benefits, list):
            for bi, benefit in enumerate(benefits):
                if bi == 0:
                    p = row_cells[2].paragraphs[0]
                else:
                    p = row_cells[2].add_paragraph()
                p.add_run(f"• {benefit}")
        else:
            row_cells[2].paragraphs[0].add_run(str(benefits))

        for cell in row_cells:
            _set_cell_shading(cell, shade)


def _build_logo_recognition(doc: Document) -> None:
    _heading1(doc, "Logo & Recognition")
    _para(
        doc,
        "Sponsor logos will appear on all event materials in a size and placement proportional to the "
        "sponsorship level. See below for recognition schedule:"
    )
    recognition_items = [
        "Event Banner — Title and Gold sponsors",
        "Event Program (printed) — All sponsors, sized by tier",
        "Event Signage (venue) — Title, Gold, and Silver sponsors",
        "Organization Website — Title and Gold sponsors (12-month listing)",
        "Social Media — Pre-event recognition posts per tier schedule",
        "Email Communications — Title sponsors featured in event emails",
    ]
    for item in recognition_items:
        _bullet(doc, item)
    doc.add_paragraph()
    _para(
        doc,
        "[PLACEHOLDER: Insert a visual mockup or logo placement diagram here. "
        "Designer to add sample banner/program layout showing tier placement.]",
        italic=True,
    )


def _build_contact_next_steps(
    doc: Document,
    signer_name: str,
    signer_title: str,
    org_name: str,
    event_name: str,
) -> None:
    _heading1(doc, "Contact & Next Steps")
    _para(
        doc,
        f"We would love to have you as a sponsor of {event_name}. To confirm your sponsorship or ask questions, "
        f"please contact us by the deadline below."
    )
    doc.add_paragraph()

    deadline_para = doc.add_paragraph()
    dl_run = deadline_para.add_run("Sponsorship Deadline: ")
    dl_run.bold = True
    deadline_para.add_run("[INSERT DEADLINE DATE — we recommend 6-8 weeks before the event]")

    doc.add_paragraph()

    for label, value in [
        ("Contact", signer_name),
        ("Title", signer_title),
        ("Organization", org_name),
        ("Email", "[INSERT EMAIL]"),
        ("Phone", "[INSERT PHONE]"),
    ]:
        p = doc.add_paragraph()
        _bold_label(p, label, value)

    doc.add_paragraph()
    _para(
        doc,
        "To confirm your sponsorship, please complete the attached Sponsorship Commitment Form and return it "
        "to the contact above. A member of our team will follow up within 2 business days to confirm receipt, "
        "collect your logo files, and provide payment instructions.",
    )


# ---------------------------------------------------------------------------
# Email template builders
# ---------------------------------------------------------------------------

def _build_email_section_header(doc: Document, label: str, email_type_desc: str) -> None:
    doc.add_paragraph()
    h = doc.add_heading(label, level=2)
    if h.runs:
        h.runs[0].font.color.rgb = _STEEL
    _para(doc, email_type_desc, italic=True)
    _hr(doc)


def _build_cold_outreach(
    doc: Document,
    event_name: str,
    event_date: str,
    org_name: str,
    org_mission: str,
    fundraising_goal: str,
    signer_name: str,
    signer_title: str,
) -> None:
    _build_email_section_header(
        doc,
        "Email Template 1 — Cold Outreach",
        "Use for sponsors you have not contacted before. Formal introduction to the organization and event.",
    )
    _para(doc, "Subject: Sponsorship Opportunity — " + event_name + " | " + event_date)
    doc.add_paragraph()
    _para(doc, "Dear [Contact Name],")
    doc.add_paragraph()
    body = (
        f"I am reaching out on behalf of {org_name} to share an exciting sponsorship opportunity. "
        f"On {event_date}, we will host {event_name} — an annual event that brings together our community "
        f"of supporters, partners, and advocates to celebrate the impact of our work and raise critical funds "
        f"for the year ahead. Our goal this year is {fundraising_goal}.\n\n"
        f"{org_mission}\n\n"
        f"We are inviting a select group of community-minded organizations to join us as event sponsors. "
        f"Sponsorship is an opportunity to align your brand with our mission, gain meaningful visibility among "
        f"our network, and demonstrate your commitment to our community.\n\n"
        f"I have attached our Sponsorship Prospectus for your review. Tiers start at [LOWEST TIER AMOUNT] and "
        f"include a range of recognition benefits. I would welcome a brief conversation to share more about the "
        f"event and explore whether a partnership makes sense."
    )
    _para(doc, body)
    doc.add_paragraph()
    _para(doc, "Would you have 15 minutes in the next week or two for a quick call?")
    doc.add_paragraph()
    _para(doc, "Thank you for your consideration. I look forward to hearing from you.")
    doc.add_paragraph()
    _para(doc, f"Warm regards,\n{signer_name}\n{signer_title}\n{org_name}\n[Phone] | [Email]")


def _build_warm_outreach(
    doc: Document,
    event_name: str,
    event_date: str,
    org_name: str,
    fundraising_goal: str,
    signer_name: str,
    signer_title: str,
) -> None:
    _build_email_section_header(
        doc,
        "Email Template 2 — Warm Outreach",
        "Use for contacts with an existing relationship — past sponsors, board connections, or known community partners.",
    )
    _para(doc, "Subject: " + event_name + " — " + event_date + " | Sponsorship Invitation")
    doc.add_paragraph()
    _para(doc, "Dear [Contact Name],")
    doc.add_paragraph()
    body = (
        f"I hope you are doing well! I am writing with an exciting invitation. {event_name} is coming up on "
        f"{event_date}, and I immediately thought of you and [Company/Organization Name] as a perfect fit "
        f"for our sponsor family this year.\n\n"
        f"[PERSONALIZATION NOTE: Reference your prior connection here — e.g., 'It was so great to connect at "
        f"[event/meeting] last [season/year]' or 'We so appreciated [Company]'s past support of [program/event].']\n\n"
        f"This year we are aiming to raise {fundraising_goal} to support [briefly describe program impact]. "
        f"As a sponsor, you would receive [reference 1-2 tier benefits most relevant to this contact], "
        f"along with the satisfaction of knowing your investment directly supports [mission outcome].\n\n"
        f"I have attached our Sponsorship Prospectus with full tier details. Given our relationship, "
        f"I wanted to reach out personally before we opened this more broadly. The deadline to confirm is "
        f"[INSERT DEADLINE], and the most visible tiers fill quickly."
    )
    _para(doc, body)
    doc.add_paragraph()
    _para(doc, "Would you have time for a quick call this week or next? I would love to reconnect and share more.")
    doc.add_paragraph()
    _para(doc, f"With gratitude,\n{signer_name}\n{signer_title}\n{org_name}\n[Phone] | [Email]")


def _build_last_chance(
    doc: Document,
    event_name: str,
    event_date: str,
    org_name: str,
    signer_name: str,
    signer_title: str,
) -> None:
    _build_email_section_header(
        doc,
        "Email Template 3 — Last-Chance Follow-Up",
        "Use approximately 1 week before the sponsorship deadline. Friendly urgency — not pushy.",
    )
    _para(doc, "Subject: Last Chance — " + event_name + " Sponsorship Deadline [DATE]")
    doc.add_paragraph()
    _para(doc, "Dear [Contact Name],")
    doc.add_paragraph()
    body = (
        f"I wanted to send a quick note as we approach our sponsorship deadline for {event_name} on {event_date}.\n\n"
        f"We still have a few sponsorship opportunities available, and I did not want you to miss the chance "
        f"to be part of this event. [OPTIONAL: 'Several of our higher tiers have already been claimed — "
        f"we have [X] spots remaining at the [Tier Name] level.']\n\n"
        f"If you have been considering a sponsorship and are ready to move forward, please reply to this email "
        f"or reach me at [phone] by [DEADLINE DATE]. I am happy to answer any remaining questions.\n\n"
        f"If the timing is not right this year, no worries at all — I would love to keep {org_name} on your "
        f"radar for future opportunities. Either way, thank you for your time and your support of our mission."
    )
    _para(doc, body)
    doc.add_paragraph()
    _para(doc, f"Thank you,\n{signer_name}\n{signer_title}\n{org_name}\n[Phone] | [Email]")


# ---------------------------------------------------------------------------
# Core render function
# ---------------------------------------------------------------------------

def render(
    event_name: str,
    event_date: str,
    org_name: str,
    org_mission: str,
    event_purpose: str,
    fundraising_goal: str,
    sponsor_tiers: list,
    signer_name: str,
    signer_title: str,
    venue: Optional[str] = None,
    attendee_estimate=None,
    output_dir: str = "/mnt/user-data/outputs",
) -> str:
    """
    Generate a sponsorship prospectus + 3 outreach email templates as a Word doc.

    Parameters
    ----------
    event_name : str
        Name of the event
    event_date : str
        Date of the event
    org_name : str
        Full name of the organization
    org_mission : str
        1-2 sentence mission statement
    event_purpose : str
        What the event is for, who it serves, fundraising goal description
    fundraising_goal : str
        Dollar goal for the event (e.g. "$150,000")
    sponsor_tiers : list of dict
        Each dict: {name, amount, benefits (list of strings)}
    signer_name : str
        Name of the contact signing the package
    signer_title : str
        Title of the signer
    venue : str, optional
        Venue name and address
    attendee_estimate : int or str, optional
        Expected attendance
    output_dir : str
        Directory to save the output file

    Returns
    -------
    str
        Absolute path to the generated .docx file
    """
    required = {
        "event_name": event_name,
        "event_date": event_date,
        "org_name": org_name,
        "org_mission": org_mission,
        "event_purpose": event_purpose,
        "fundraising_goal": fundraising_goal,
        "signer_name": signer_name,
        "signer_title": signer_title,
    }
    for field, value in required.items():
        if not value or not str(value).strip():
            raise ValueError(f"{field} is required and cannot be empty")

    if not sponsor_tiers or not isinstance(sponsor_tiers, list) or len(sponsor_tiers) == 0:
        raise ValueError("sponsor_tiers is required and must be a non-empty list")

    os.makedirs(output_dir, exist_ok=True)
    timestamp = int(_time.time())
    safe_name = re.sub(r"[^\w]+", "_", event_name)[:30].strip("_")
    filename = f"sponsor_package_{safe_name}_{timestamp}.docx"
    out_path = os.path.join(output_dir, filename)

    doc = Document()

    # --- Prospectus sections ---
    _build_cover(doc, event_name, event_date, org_name, fundraising_goal)
    doc.add_page_break()

    _build_about_event(doc, event_name, event_date, venue, org_name, event_purpose, fundraising_goal, attendee_estimate)
    _hr(doc)

    _build_about_org(doc, org_name, org_mission)
    _hr(doc)

    _build_why_sponsor(doc, event_name, org_name)
    _hr(doc)

    _build_tiers_table(doc, sponsor_tiers)
    _hr(doc)

    _build_logo_recognition(doc)
    _hr(doc)

    _build_contact_next_steps(doc, signer_name, signer_title, org_name, event_name)

    # --- Email templates (separate page) ---
    doc.add_page_break()

    email_header_para = doc.add_paragraph()
    email_header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eh_run = email_header_para.add_run("OUTREACH EMAIL TEMPLATES")
    eh_run.bold = True
    eh_run.font.size = Pt(14)
    eh_run.font.color.rgb = _NAVY

    note_para = doc.add_paragraph()
    note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note_para.add_run(
        "Customize each template with the recipient's name, company, and any personalization notes "
        "before sending. Do not send without reviewing for accuracy."
    )
    nr.italic = True
    nr.font.color.rgb = _ORANGE
    nr.font.size = Pt(9)

    _build_cold_outreach(doc, event_name, event_date, org_name, org_mission, fundraising_goal, signer_name, signer_title)
    doc.add_page_break()

    _build_warm_outreach(doc, event_name, event_date, org_name, fundraising_goal, signer_name, signer_title)
    doc.add_page_break()

    _build_last_chance(doc, event_name, event_date, org_name, signer_name, signer_title)

    doc.save(out_path)
    return out_path
