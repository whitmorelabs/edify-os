"""
action_item_extractor/render.py

Extracts structured action items from pasted meeting notes or transcripts and
generates a Word (.docx) document with:
  1. Cover (meeting title + date)
  2. Action Items Table (#, Action, Owner, Deadline, Priority, Notes)
  3. Decisions Captured (numbered list)
  4. Open Questions (numbered list)

Parsing approach:
  - Split by lines, apply heuristic keyword matching
  - Action indicators: "action:", "TODO:", "[action]", "will X", "needs to X",
    "responsible for", "follow up on", "assigned to", "to complete", "to schedule",
    "to send", "to draft", "to prepare", "to coordinate"
  - Decision indicators: "decided", "agreed", "approved", "confirmed", "resolved",
    "it was decided", "board voted", "motion passed"
  - Question indicators: "?", "open question:", "TBD:", "unclear:", "pending:",
    "need to determine", "need to find out"
  - Owner extraction: patterns like "Owner: X", "[X]", "(X)", "assigned to X",
    "X will", "X to", "X is responsible"
  - Date extraction: "by YYYY-MM-DD", "by Month Day", "by Friday/Monday/etc.",
    "due [date]", "deadline [date]", relative ("next week", "end of month")
  - Priority: "urgent", "critical", "asap", "high priority" → "High"; else "Normal"
  - Falls back gracefully when dateutil is unavailable

The output is explicitly marked as a DRAFT — the EA should review and refine.

All libraries are pre-installed in Anthropic's code-execution sandbox.
"""

import datetime
import os
import re
import time
from typing import Optional, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ---------------------------------------------------------------------------
# Optional dateutil import (available in Anthropic sandbox; graceful fallback)
# ---------------------------------------------------------------------------

try:
    from dateutil import parser as dateutil_parser
    from dateutil.relativedelta import relativedelta
    _DATEUTIL_AVAILABLE = True
except ImportError:
    _DATEUTIL_AVAILABLE = False


# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------

_NAVY = RGBColor(0x1A, 0x2E, 0x5C)
_SLATE = RGBColor(0x3A, 0x55, 0x77)
_GRAY = RGBColor(0xCC, 0xCC, 0xCC)
_LIGHT_GRAY = RGBColor(0x88, 0x88, 0x88)
_DRAFT_ORANGE = RGBColor(0xAA, 0x44, 0x00)
_WHITE_TEXT = RGBColor(0xFF, 0xFF, 0xFF)

_HEADER_FILL = "1A2E5C"
_ALT_ROW_FILL = "EEF1F7"
_HIGH_PRI_FILL = "FFC7CE"
_HIGH_PRI_TEXT = RGBColor(0x9C, 0x00, 0x06)
_NORMAL_PRI_TEXT = RGBColor(0x37, 0x6C, 0x23)


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------

def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top: int = 50, bottom: int = 50, left: int = 80, right: int = 80) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


# ---------------------------------------------------------------------------
# Parsing helpers
# ---------------------------------------------------------------------------

# Keywords that signal an action item
_ACTION_KEYWORDS = [
    r"\baction\b", r"\btodo\b", r"\bto[- ]do\b",
    r"\bwill\b", r"\bneeds? to\b", r"\bresponsible for\b",
    r"\bfollow[- ]?up\b", r"\bassigned to\b", r"\bto complete\b",
    r"\bto schedule\b", r"\bto send\b", r"\bto draft\b",
    r"\bto prepare\b", r"\bto coordinate\b", r"\bto review\b",
    r"\bto submit\b", r"\bto confirm\b", r"\bto create\b",
    r"\bto reach out\b", r"\bto contact\b", r"\bto update\b",
    r"\bto share\b", r"\bto present\b", r"\bto finalize\b",
    r"\bto book\b", r"\bto arrange\b", r"\bto plan\b",
    r"\bto report\b", r"\bto investigate\b", r"\bto explore\b",
    r"\bwill handle\b", r"\bwill take\b", r"\bwill reach\b",
    r"\bwill prepare\b", r"\bwill draft\b", r"\bwill send\b",
    r"\bwill follow\b", r"\bwill coordinate\b", r"\bwill schedule\b",
    r"\bwill update\b", r"\bwill submit\b", r"\baction item\b",
    r"\bai\b",  # "AI:" prefix convention in some orgs
]
_ACTION_RE = re.compile("|".join(_ACTION_KEYWORDS), re.IGNORECASE)

# Keywords that signal a decision
_DECISION_KEYWORDS = [
    r"\bdecided\b", r"\bagreed\b", r"\bapproved\b", r"\bconfirmed\b",
    r"\bresolved\b", r"\bit was decided\b", r"\bvoted\b", r"\bmotion passed\b",
    r"\bmotion carried\b", r"\bwe will\b", r"\bdecision\b", r"\bconcluded\b",
    r"\bcommitted to\b", r"\bwill proceed\b", r"\bgoing with\b",
    r"\bchose\b", r"\bselected\b", r"\badopted\b",
]
_DECISION_RE = re.compile("|".join(_DECISION_KEYWORDS), re.IGNORECASE)

# Keywords that signal an open question
_QUESTION_KEYWORDS = [
    r"\?", r"\bopen question\b", r"\btbd\b", r"\bunclear\b",
    r"\bpending\b", r"\bneed to determine\b", r"\bneed to find out\b",
    r"\bneed to confirm\b", r"\bunresolved\b", r"\bto be determined\b",
    r"\bto be confirmed\b", r"\bto be decided\b", r"\bstill unknown\b",
    r"\bopen issue\b",
]
_QUESTION_RE = re.compile("|".join(_QUESTION_KEYWORDS), re.IGNORECASE)

# Priority signals
_HIGH_PRIORITY_RE = re.compile(
    r"\b(urgent|critical|asap|as soon as possible|high priority|high-priority|immediately)\b",
    re.IGNORECASE
)

# Owner patterns — try to extract a name after ownership signals
_OWNER_PATTERNS = [
    re.compile(r"owner[:\s]+([A-Z][a-z]+(?: [A-Z][a-z]+)?)", re.IGNORECASE),
    re.compile(r"assigned to[:\s]+([A-Z][a-z]+(?: [A-Z][a-z]+)?)", re.IGNORECASE),
    re.compile(r"responsible[:\s]+([A-Z][a-z]+(?: [A-Z][a-z]+)?)", re.IGNORECASE),
    re.compile(r"@([A-Z][a-z]+(?: [A-Z][a-z]+)?)", re.IGNORECASE),
    # Pattern: "Name will X" or "Name to X" at start of sentence
    re.compile(r"^([A-Z][a-z]+(?: [A-Z][a-z]+)?) (?:will|to) ", re.IGNORECASE),
    # Bracketed names: [Name] or (Name)
    re.compile(r"[\[\(]([A-Z][a-z]+(?: [A-Z][a-z]+)?)[\]\)]"),
]

# Absolute date patterns
_DATE_PATTERNS = [
    # ISO: 2026-04-26
    re.compile(r"\b(\d{4}-\d{2}-\d{2})\b"),
    # Month Day Year: April 26, 2026 / Apr 26 2026
    re.compile(
        r"\b(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
        r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|"
        r"Dec(?:ember)?)\s+\d{1,2}(?:,?\s*\d{4})?",
        re.IGNORECASE
    ),
    # "by Friday", "by Monday", etc.
    re.compile(
        r"\bby\s+(Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday)\b",
        re.IGNORECASE
    ),
    # "by end of week/month/quarter"
    re.compile(r"\bby\s+end\s+of\s+(week|month|quarter|year)\b", re.IGNORECASE),
    # "next week", "next month"
    re.compile(r"\bnext\s+(week|month|quarter)\b", re.IGNORECASE),
    # "due [date phrase]"
    re.compile(r"\bdue\s+(\S+(?:\s+\S+)?)", re.IGNORECASE),
    # "deadline: [date]"
    re.compile(r"\bdeadline[:\s]+(\S+(?:\s+\S+)?)", re.IGNORECASE),
]

# Relative weekday offsets from current date
_WEEKDAY_OFFSETS = {
    "monday": 0, "tuesday": 1, "wednesday": 2, "thursday": 3,
    "friday": 4, "saturday": 5, "sunday": 6,
}


def _extract_date(line: str) -> str:
    """
    Try to extract a deadline date from a line of text.
    Returns a date string (absolute preferred, relative label as fallback).
    Returns "" if no date found.
    """
    today = datetime.date.today()

    # ISO date — highest confidence
    iso_match = re.search(r"\b(\d{4}-\d{2}-\d{2})\b", line)
    if iso_match:
        return iso_match.group(1)

    # "by Month Day Year"
    month_match = re.search(
        r"\b(January|February|March|April|May|June|July|August|September|"
        r"October|November|December|Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Oct|Nov|Dec)"
        r"\s+(\d{1,2})(?:,?\s*(\d{4}))?",
        line, re.IGNORECASE
    )
    if month_match:
        raw = month_match.group(0)
        if _DATEUTIL_AVAILABLE:
            try:
                parsed = dateutil_parser.parse(raw, default=today.replace(day=1))
                return parsed.strftime("%Y-%m-%d")
            except Exception:
                return raw
        return raw

    # "by Friday" → next occurrence
    weekday_match = re.search(
        r"\bby\s+(Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday)\b",
        line, re.IGNORECASE
    )
    if weekday_match:
        day_name = weekday_match.group(1).lower()
        target = _WEEKDAY_OFFSETS.get(day_name, -1)
        if target >= 0:
            days_ahead = (target - today.weekday()) % 7
            if days_ahead == 0:
                days_ahead = 7
            result = today + datetime.timedelta(days=days_ahead)
            return result.strftime("%Y-%m-%d")

    # "by end of week/month/quarter"
    eow_match = re.search(r"\bby\s+end\s+of\s+(week|month|quarter|year)\b", line, re.IGNORECASE)
    if eow_match:
        period = eow_match.group(1).lower()
        if period == "week":
            days_to_friday = (4 - today.weekday()) % 7
            if days_to_friday == 0:
                days_to_friday = 7
            return (today + datetime.timedelta(days=days_to_friday)).strftime("%Y-%m-%d")
        elif period == "month":
            # Last day of current month
            if today.month == 12:
                last_day = today.replace(year=today.year + 1, month=1, day=1) - datetime.timedelta(days=1)
            else:
                last_day = today.replace(month=today.month + 1, day=1) - datetime.timedelta(days=1)
            return last_day.strftime("%Y-%m-%d")
        elif period == "quarter":
            quarter_end_months = {1: 3, 2: 3, 3: 3, 4: 6, 5: 6, 6: 6,
                                  7: 9, 8: 9, 9: 9, 10: 12, 11: 12, 12: 12}
            end_month = quarter_end_months[today.month]
            if end_month == 12:
                last_day = today.replace(month=12, day=31)
            else:
                last_day = today.replace(month=end_month + 1, day=1) - datetime.timedelta(days=1)
            return last_day.strftime("%Y-%m-%d")
        elif period == "year":
            return today.replace(month=12, day=31).strftime("%Y-%m-%d")

    # "next week/month/quarter"
    next_match = re.search(r"\bnext\s+(week|month|quarter)\b", line, re.IGNORECASE)
    if next_match:
        period = next_match.group(1).lower()
        if period == "week":
            days_to_next_monday = (7 - today.weekday()) % 7
            if days_to_next_monday == 0:
                days_to_next_monday = 7
            return (today + datetime.timedelta(days=days_to_next_monday)).strftime("%Y-%m-%d")
        elif period == "month":
            if today.month == 12:
                return today.replace(year=today.year + 1, month=1, day=1).strftime("%Y-%m-%d")
            return today.replace(month=today.month + 1, day=1).strftime("%Y-%m-%d")
        elif period == "quarter":
            return "Next quarter"

    return ""


def _extract_owner(line: str) -> str:
    """Try to extract an owner name from a line. Returns "" if not found."""
    for pattern in _OWNER_PATTERNS:
        m = pattern.search(line)
        if m:
            candidate = m.group(1).strip()
            # Filter out common false positives (single common words)
            stopwords = {"The", "A", "An", "This", "That", "We", "They", "All",
                         "Our", "Your", "His", "Her", "Its", "Each", "Both"}
            if candidate not in stopwords and len(candidate) > 1:
                return candidate
    return ""


def _clean_line(line: str) -> str:
    """Remove common prefixes and markdown formatting."""
    line = line.strip()
    # Remove bullet markers
    line = re.sub(r"^[-*•·]\s+", "", line)
    # Remove numbered list markers
    line = re.sub(r"^\d+[.)]\s+", "", line)
    # Remove common prefix labels
    line = re.sub(r"^(Action|TODO|AI|Action Item|Decision|Q|Open Question)[:\s]+", "", line,
                  flags=re.IGNORECASE)
    return line.strip()


def _classify_line(line: str) -> str:
    """
    Classify a line as: 'action', 'decision', 'question', or 'context'.
    Priority: action > decision > question > context.
    """
    if _ACTION_RE.search(line):
        return "action"
    if _DECISION_RE.search(line):
        return "decision"
    if _QUESTION_RE.search(line):
        return "question"
    return "context"


def _parse_notes(
    meeting_notes: str,
    default_owner: str,
) -> Tuple[list, list, list]:
    """
    Parse freeform meeting notes into (actions, decisions, questions).

    Returns
    -------
    actions : list of dict
        {action, owner, deadline, priority, notes}
    decisions : list of str
    questions : list of str
    """
    actions = []
    decisions = []
    questions = []

    lines = meeting_notes.splitlines()

    for raw_line in lines:
        stripped = raw_line.strip()
        if not stripped or len(stripped) < 5:
            continue

        category = _classify_line(stripped)
        cleaned = _clean_line(stripped)

        if not cleaned:
            continue

        if category == "action":
            owner = _extract_owner(stripped) or default_owner
            deadline = _extract_date(stripped)
            priority = "High" if _HIGH_PRIORITY_RE.search(stripped) else "Normal"

            # Build notes: capture any extra context (e.g., if it's a long sentence)
            notes = ""
            # If the cleaned line contains a colon after an action label, the "notes"
            # are the part after the colon
            colon_split = re.split(r":\s+", cleaned, maxsplit=1)
            if len(colon_split) == 2 and len(colon_split[0]) < 30:
                cleaned = colon_split[1]
            # Truncate if very long — keep first ~120 chars as action, rest as notes
            if len(cleaned) > 140:
                notes = cleaned[140:].strip()
                cleaned = cleaned[:140].strip() + "..."

            actions.append({
                "action": cleaned,
                "owner": owner,
                "deadline": deadline,
                "priority": priority,
                "notes": notes,
            })

        elif category == "decision":
            decisions.append(cleaned)

        elif category == "question":
            questions.append(cleaned)

    return actions, decisions, questions


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _heading1(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=1)
    if h.runs:
        h.runs[0].font.color.rgb = _NAVY


def _para(doc: Document, text: str, italic: bool = False, size: int = 10,
          color: Optional[RGBColor] = None) -> None:
    p = doc.add_paragraph(text)
    if p.runs:
        if italic:
            p.runs[0].italic = True
        p.runs[0].font.size = Pt(size)
        if color:
            p.runs[0].font.color.rgb = color


def _hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 80)
    run.font.color.rgb = _GRAY
    run.font.size = Pt(7)


def _numbered(doc: Document, text: str, size: int = 10) -> None:
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(size)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _build_cover(
    doc: Document,
    meeting_title: str,
    meeting_date: str,
    org_name: Optional[str],
    today_str: str,
    action_count: int,
    decision_count: int,
    question_count: int,
) -> None:
    for _ in range(2):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title_para.add_run("Meeting Action Items")
    t_run.bold = True
    t_run.font.size = Pt(24)
    t_run.font.color.rgb = _NAVY

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub_para.add_run(meeting_title)
    s_run.bold = True
    s_run.font.size = Pt(14)
    s_run.font.color.rgb = _SLATE

    doc.add_paragraph()

    fields = [
        ("Meeting Date", meeting_date),
    ]
    if org_name:
        fields.append(("Organization", org_name))
    fields.append(("Extracted", today_str))

    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_label = p.add_run(f"{label}: ")
        r_label.bold = True
        r_value = p.add_run(value)
        _ = r_value

    doc.add_paragraph()

    # Summary counts
    summary_para = doc.add_paragraph()
    summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = summary_para.add_run(
        f"{action_count} action item(s)  |  {decision_count} decision(s)  |  {question_count} open question(s)"
    )
    s_run.font.size = Pt(10)
    s_run.font.color.rgb = _SLATE

    doc.add_paragraph()

    draft_para = doc.add_paragraph()
    draft_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d_run = draft_para.add_run(
        "DRAFT — Heuristically extracted. Review and refine before distributing."
    )
    d_run.italic = True
    d_run.font.size = Pt(9)
    d_run.font.color.rgb = _DRAFT_ORANGE


def _build_action_table(doc: Document, actions: list) -> None:
    _heading1(doc, "Action Items")

    _para(
        doc,
        "Extracted from meeting notes. Review owners, deadlines, and priorities before sending.",
        italic=True,
        size=9,
        color=_DRAFT_ORANGE,
    )
    doc.add_paragraph()

    if not actions:
        _para(doc, "[No action items were extracted from the provided notes.]", italic=True)
        return

    headers = ["#", "Action", "Owner", "Deadline", "Priority", "Notes"]
    col_widths = [Inches(0.3), Inches(2.4), Inches(1.0), Inches(1.0), Inches(0.7), Inches(2.1)]

    table = doc.add_table(rows=len(actions) + 1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr_row.cells[i]
        cell.width = w
        _shade_cell(cell, _HEADER_FILL)
        _set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = _WHITE_TEXT
        run.font.size = Pt(9)

    for row_idx, action in enumerate(actions):
        action_text = str(action.get("action", ""))
        owner = str(action.get("owner", "Unassigned"))
        deadline = str(action.get("deadline", ""))
        priority = str(action.get("priority", "Normal"))
        notes = str(action.get("notes", ""))

        table_row = table.rows[row_idx + 1]

        is_high = priority == "High"

        if is_high:
            for ci in range(6):
                _shade_cell(table_row.cells[ci], _HIGH_PRI_FILL)
        elif row_idx % 2 == 1:
            for ci in range(6):
                _shade_cell(table_row.cells[ci], _ALT_ROW_FILL)

        # Number
        table_row.cells[0].width = col_widths[0]
        _set_cell_margins(table_row.cells[0])
        p_num = table_row.cells[0].paragraphs[0]
        p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_num = p_num.add_run(str(row_idx + 1))
        r_num.font.size = Pt(9)

        # Action
        table_row.cells[1].width = col_widths[1]
        _set_cell_margins(table_row.cells[1])
        r_action = table_row.cells[1].paragraphs[0].add_run(action_text)
        r_action.font.size = Pt(9)

        # Owner
        table_row.cells[2].width = col_widths[2]
        _set_cell_margins(table_row.cells[2])
        r_owner = table_row.cells[2].paragraphs[0].add_run(owner)
        r_owner.font.size = Pt(9)
        r_owner.bold = True

        # Deadline
        table_row.cells[3].width = col_widths[3]
        _set_cell_margins(table_row.cells[3])
        p_dl = table_row.cells[3].paragraphs[0]
        p_dl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_dl = p_dl.add_run(deadline or "—")
        r_dl.font.size = Pt(9)

        # Priority
        table_row.cells[4].width = col_widths[4]
        _set_cell_margins(table_row.cells[4])
        p_pri = table_row.cells[4].paragraphs[0]
        p_pri.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_pri = p_pri.add_run(priority)
        r_pri.font.size = Pt(9)
        r_pri.bold = True
        r_pri.font.color.rgb = _HIGH_PRI_TEXT if is_high else _NORMAL_PRI_TEXT

        # Notes
        table_row.cells[5].width = col_widths[5]
        _set_cell_margins(table_row.cells[5])
        r_notes = table_row.cells[5].paragraphs[0].add_run(notes)
        r_notes.font.size = Pt(8)
        r_notes.italic = True

    doc.add_paragraph()

    key = doc.add_paragraph()
    key.add_run("Priority key: ").bold = True
    r_h = key.add_run("Red = High (urgent/critical)  ")
    r_h.font.color.rgb = _HIGH_PRI_TEXT
    r_h.font.size = Pt(9)
    r_n = key.add_run("Green = Normal")
    r_n.font.color.rgb = _NORMAL_PRI_TEXT
    r_n.font.size = Pt(9)


def _build_decisions(doc: Document, decisions: list) -> None:
    _heading1(doc, "Decisions Captured")
    _para(
        doc,
        "These lines were identified as decisions made during the meeting.",
        italic=True,
        size=9,
        color=_LIGHT_GRAY,
    )
    doc.add_paragraph()

    if not decisions:
        _para(doc, "[No decisions were extracted from the provided notes.]", italic=True)
        return

    for decision in decisions:
        _numbered(doc, str(decision))


def _build_questions(doc: Document, questions: list) -> None:
    _heading1(doc, "Open Questions")
    _para(
        doc,
        "These items were raised during the meeting but not resolved. Assign owners before the next meeting.",
        italic=True,
        size=9,
        color=_LIGHT_GRAY,
    )
    doc.add_paragraph()

    if not questions:
        _para(doc, "[No open questions were identified from the provided notes.]", italic=True)
        return

    for question in questions:
        _numbered(doc, str(question))


# ---------------------------------------------------------------------------
# Core render function
# ---------------------------------------------------------------------------

def render(
    meeting_title: str,
    meeting_date: str,
    meeting_notes: str,
    default_owner: str = "Unassigned",
    org_name: Optional[str] = None,
    output_dir: str = "/mnt/user-data/outputs",
) -> str:
    """
    Extract structured action items from freeform meeting notes and produce a Word doc.

    Parameters
    ----------
    meeting_title : str
        Title of the meeting (e.g. "April Board Meeting")
    meeting_date : str
        Date of the meeting
    meeting_notes : str
        Freeform meeting notes, transcript, or summary text (pasted by the user)
    default_owner : str
        Default owner when no owner is identified (default: "Unassigned")
    org_name : str, optional
        Organization name for the cover page
    output_dir : str
        Directory to save the output file

    Returns
    -------
    str
        Absolute path to the generated .docx file

    Notes
    -----
    Parsing is heuristic and intended to produce a DRAFT for EA review.
    The output is explicitly marked as a draft. Common limitations:
    - May miss action items if phrased unusually
    - Owner extraction works best with explicit "Name will..." or "Owner: Name" patterns
    - Relative date resolution (e.g., "by Friday") is computed from today's date
    """
    # Validate required fields
    if not meeting_title or not str(meeting_title).strip():
        raise ValueError("meeting_title is required")
    if not meeting_date or not str(meeting_date).strip():
        raise ValueError("meeting_date is required")
    if not meeting_notes or not str(meeting_notes).strip():
        raise ValueError("meeting_notes is required and cannot be empty")

    default_owner = default_owner or "Unassigned"
    today_str = datetime.date.today().strftime("%B %d, %Y")

    # Parse the notes
    actions, decisions, questions = _parse_notes(meeting_notes, default_owner)

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # 1. Cover
    _build_cover(
        doc, meeting_title, meeting_date, org_name, today_str,
        len(actions), len(decisions), len(questions)
    )
    doc.add_page_break()

    # 2. Action Items Table
    _build_action_table(doc, actions)
    _hr(doc)

    # 3. Decisions
    _build_decisions(doc, decisions)
    _hr(doc)

    # 4. Open Questions
    _build_questions(doc, questions)

    # Footer note
    doc.add_paragraph()
    _hr(doc)
    footer = doc.add_paragraph()
    f_run = footer.add_run(
        f"Generated {today_str} | Parsing is heuristic — review all items before distributing. "
        f"{'dateutil available — relative dates resolved.' if _DATEUTIL_AVAILABLE else 'dateutil not available — relative dates kept as labels.'}"
    )
    f_run.italic = True
    f_run.font.size = Pt(8)
    f_run.font.color.rgb = _LIGHT_GRAY

    # Save
    os.makedirs(output_dir, exist_ok=True)
    safe_title = re.sub(r"[^\w]+", "_", meeting_title)[:25].strip("_")
    safe_date = re.sub(r"[^\w]+", "_", meeting_date)[:12].strip("_")
    filename = f"action_items_{safe_title}_{safe_date}.docx"
    out_path = os.path.join(output_dir, filename)
    doc.save(out_path)

    return out_path
