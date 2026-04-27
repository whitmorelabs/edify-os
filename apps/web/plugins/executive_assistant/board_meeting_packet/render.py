"""
board_meeting_packet/render.py

Generates a complete nonprofit board meeting packet as a Word (.docx) document.
Covers the conventions nonprofit boards use: consent calendars, committee reports,
action item trackers, and executive summaries — none of which generic docx tools know.

Sections:
  1. Cover Page
  2. Agenda (table: Time | Item | Presenter | Type | Supporting Material)
  3. Consent Calendar (numbered list + recommended motion)
  4. Action Item Tracker (table from prior meeting)
  5. Committee Reports
  6. Executive Summary
  7. Appendix placeholder

All libraries are pre-installed in Anthropic's code-execution sandbox.
"""

import datetime
import os
import re
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------

_NAVY = RGBColor(0x1A, 0x2E, 0x5C)
_SLATE = RGBColor(0x3A, 0x55, 0x77)
_GRAY = RGBColor(0xCC, 0xCC, 0xCC)
_LIGHT_GRAY_TEXT = RGBColor(0x88, 0x88, 0x88)
_DRAFT_ORANGE = RGBColor(0xAA, 0x44, 0x00)
_WHITE_TEXT = RGBColor(0xFF, 0xFF, 0xFF)

_HEADER_FILL = "1A2E5C"
_ALT_ROW_FILL = "EEF1F7"

# Agenda item type colors
_TYPE_FILLS = {
    "decision": "FFE0B2",
    "discussion": "E3F2FD",
    "informational": "E8F5E9",
    "consent": "F3E5F5",
}
_TYPE_TEXT = {
    "decision": RGBColor(0x7F, 0x3B, 0x00),
    "discussion": RGBColor(0x0D, 0x47, 0xA1),
    "informational": RGBColor(0x1B, 0x5E, 0x20),
    "consent": RGBColor(0x4A, 0x14, 0x8C),
}

# Status badge colors
_STATUS_FILLS = {
    "complete": "C6EFCE",
    "in progress": "FFEB9C",
    "overdue": "FFC7CE",
    "deferred": "E0E0E0",
    "pending": "FFEB9C",
}
_STATUS_TEXT = {
    "complete": RGBColor(0x37, 0x6C, 0x23),
    "in progress": RGBColor(0x9C, 0x65, 0x00),
    "overdue": RGBColor(0x9C, 0x00, 0x06),
    "deferred": RGBColor(0x55, 0x55, 0x55),
    "pending": RGBColor(0x9C, 0x65, 0x00),
}


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------

def _shade_cell(cell, hex_color: str) -> None:
    """Fill a table cell background with a solid hex color (no #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top: int = 60, bottom: int = 60, left: int = 80, right: int = 80) -> None:
    """Set cell internal margins (twips)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _add_header_footer(doc: Document, meeting_date_str: str, org_name: str) -> None:
    """Add header and footer with meeting info."""
    section = doc.sections[0]

    # Header
    section.header.is_linked_to_previous = False
    hdr_para = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    hdr_para.clear()
    hdr_run = hdr_para.add_run(f"{org_name}  |  Board Meeting Packet  |  {meeting_date_str}")
    hdr_run.font.size = Pt(8)
    hdr_run.font.color.rgb = _LIGHT_GRAY_TEXT
    hdr_para.paragraph_format.space_after = Pt(0)

    # Footer
    section.footer.is_linked_to_previous = False
    ftr_para = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    ftr_para.clear()
    ftr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ftr_run = ftr_para.add_run("CONFIDENTIAL — Board Use Only")
    ftr_run.font.size = Pt(8)
    ftr_run.italic = True
    ftr_run.font.color.rgb = _LIGHT_GRAY_TEXT


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _heading1(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=1)
    if h.runs:
        h.runs[0].font.color.rgb = _NAVY


def _heading2(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=2)
    if h.runs:
        h.runs[0].font.color.rgb = _SLATE


def _para(doc: Document, text: str, italic: bool = False, size: Optional[int] = None,
          color: Optional[RGBColor] = None) -> None:
    p = doc.add_paragraph(text)
    if p.runs:
        if italic:
            p.runs[0].italic = True
        if size:
            p.runs[0].font.size = Pt(size)
        if color:
            p.runs[0].font.color.rgb = color


def _hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 80)
    run.font.color.rgb = _GRAY
    run.font.size = Pt(7)


def _table_header_row(table, headers: list, col_widths: list) -> None:
    """Style the first row of a table as a dark header."""
    hdr_row = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr_row.cells[i]
        cell.width = w
        _shade_cell(cell, _HEADER_FILL)
        _set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = _WHITE_TEXT
        run.font.size = Pt(9)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _build_cover(
    doc: Document,
    org_name: str,
    meeting_date_str: str,
    meeting_type: str,
    prepared_by: str,
    signer_role: str,
    today_str: str,
) -> None:
    for _ in range(3):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title_para.add_run("Board Meeting Packet")
    t_run.bold = True
    t_run.font.size = Pt(26)
    t_run.font.color.rgb = _NAVY

    org_para = doc.add_paragraph()
    org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    o_run = org_para.add_run(org_name)
    o_run.bold = True
    o_run.font.size = Pt(16)
    o_run.font.color.rgb = _SLATE

    doc.add_paragraph()

    type_label = meeting_type.replace("_", " ").title()
    fields = [
        ("Meeting Type", f"{type_label} Board Meeting"),
        ("Meeting Date", meeting_date_str),
        ("Prepared By", prepared_by),
        ("Role", signer_role),
        ("Date Prepared", today_str),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_label = p.add_run(f"{label}: ")
        r_label.bold = True
        r_label.font.size = Pt(11)
        r_value = p.add_run(value)
        r_value.font.size = Pt(11)

    doc.add_paragraph()
    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_run = conf_para.add_run("CONFIDENTIAL — Intended for Board Members Only")
    c_run.italic = True
    c_run.font.size = Pt(9)
    c_run.font.color.rgb = _DRAFT_ORANGE


def _build_agenda(doc: Document, agenda_items: list) -> None:
    _heading1(doc, "Meeting Agenda")

    headers = ["Time", "Agenda Item", "Presenter", "Type", "Supporting Material"]
    col_widths = [Inches(0.8), Inches(2.8), Inches(1.4), Inches(1.0), Inches(1.5)]

    rows = len(agenda_items) + 1
    table = doc.add_table(rows=rows, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _table_header_row(table, headers, col_widths)

    for row_idx, item in enumerate(agenda_items):
        title = str(item.get("title", "[Item]"))
        presenter = str(item.get("presenter", ""))
        time_alloc = item.get("time_allocation_minutes", "")
        item_type = str(item.get("type", "informational")).lower()
        supporting = str(item.get("supporting_doc_note", ""))

        time_str = f"{time_alloc} min" if time_alloc else ""
        type_label = item_type.title()
        fill = _TYPE_FILLS.get(item_type, "FFFFFF")
        text_color = _TYPE_TEXT.get(item_type, _NAVY)

        table_row = table.rows[row_idx + 1]

        if row_idx % 2 == 1:
            _shade_cell(table_row.cells[0], _ALT_ROW_FILL)
            _shade_cell(table_row.cells[1], _ALT_ROW_FILL)
            _shade_cell(table_row.cells[2], _ALT_ROW_FILL)
            _shade_cell(table_row.cells[4], _ALT_ROW_FILL)

        table_row.cells[0].width = col_widths[0]
        _set_cell_margins(table_row.cells[0])
        p_time = table_row.cells[0].paragraphs[0]
        p_time.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_time = p_time.add_run(time_str)
        r_time.font.size = Pt(9)

        table_row.cells[1].width = col_widths[1]
        _set_cell_margins(table_row.cells[1])
        p_title = table_row.cells[1].paragraphs[0]
        r_title = p_title.add_run(title)
        r_title.font.size = Pt(9)
        r_title.bold = True

        table_row.cells[2].width = col_widths[2]
        _set_cell_margins(table_row.cells[2])
        p_pres = table_row.cells[2].paragraphs[0]
        r_pres = p_pres.add_run(presenter)
        r_pres.font.size = Pt(9)

        table_row.cells[3].width = col_widths[3]
        _shade_cell(table_row.cells[3], fill)
        _set_cell_margins(table_row.cells[3])
        p_type = table_row.cells[3].paragraphs[0]
        p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_type = p_type.add_run(type_label)
        r_type.font.size = Pt(8)
        r_type.bold = True
        r_type.font.color.rgb = text_color

        table_row.cells[4].width = col_widths[4]
        _set_cell_margins(table_row.cells[4])
        p_sup = table_row.cells[4].paragraphs[0]
        r_sup = p_sup.add_run(supporting)
        r_sup.font.size = Pt(8)
        r_sup.italic = True

    doc.add_paragraph()
    legend = doc.add_paragraph()
    legend.add_run("Type key: ").bold = True
    for t in ["Decision", "Discussion", "Informational", "Consent"]:
        key = t.lower()
        run = legend.add_run(f"{t}  ")
        run.font.size = Pt(8)
        run.bold = True
        run.font.color.rgb = _TYPE_TEXT.get(key, _NAVY)


def _build_consent_calendar(doc: Document, consent_items: list) -> None:
    _heading1(doc, "Consent Calendar")

    _para(
        doc,
        "Consent calendar items have been reviewed in advance and are recommended for approval "
        "en bloc without individual discussion. Any board member may request that an item be "
        "removed from the consent calendar and moved to the regular agenda before the vote.",
        italic=True,
        size=10,
    )
    doc.add_paragraph()

    if consent_items:
        for i, item in enumerate(consent_items, 1):
            p = doc.add_paragraph(style="List Number")
            p.add_run(str(item)).font.size = Pt(10)
    else:
        _para(doc, "[No consent calendar items for this meeting.]", italic=True)

    doc.add_paragraph()
    motion_para = doc.add_paragraph()
    motion_para.paragraph_format.left_indent = Inches(0.4)
    m_run = motion_para.add_run(
        "Recommended motion: \"I move to approve all items on the consent calendar as presented.\""
    )
    m_run.bold = True
    m_run.font.size = Pt(10)
    m_run.font.color.rgb = _NAVY


def _build_action_tracker(doc: Document, prior_action_items: list) -> None:
    _heading1(doc, "Action Item Tracker")
    _para(
        doc,
        "Carried forward from the previous board meeting. Items marked 'Complete' may be "
        "formally closed at this meeting by motion of the board.",
        italic=True,
        size=10,
    )
    doc.add_paragraph()

    if not prior_action_items:
        _para(doc, "[No open action items from the previous meeting.]", italic=True)
        return

    headers = ["#", "Action Item", "Owner", "Status", "Notes"]
    col_widths = [Inches(0.4), Inches(2.8), Inches(1.2), Inches(1.0), Inches(2.1)]

    rows = len(prior_action_items) + 1
    table = doc.add_table(rows=rows, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _table_header_row(table, headers, col_widths)

    for row_idx, action in enumerate(prior_action_items):
        item_text = str(action.get("item", "[Action]"))
        owner = str(action.get("owner", "Unassigned"))
        status = str(action.get("status", "Pending")).lower().strip()
        notes = str(action.get("notes", ""))

        status_fill = _STATUS_FILLS.get(status, "FFFFFF")
        status_text_color = _STATUS_TEXT.get(status, _NAVY)

        table_row = table.rows[row_idx + 1]

        if row_idx % 2 == 1:
            _shade_cell(table_row.cells[0], _ALT_ROW_FILL)
            _shade_cell(table_row.cells[1], _ALT_ROW_FILL)
            _shade_cell(table_row.cells[2], _ALT_ROW_FILL)
            _shade_cell(table_row.cells[4], _ALT_ROW_FILL)

        table_row.cells[0].width = col_widths[0]
        _set_cell_margins(table_row.cells[0])
        p_num = table_row.cells[0].paragraphs[0]
        p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_num = p_num.add_run(str(row_idx + 1))
        r_num.font.size = Pt(9)

        table_row.cells[1].width = col_widths[1]
        _set_cell_margins(table_row.cells[1])
        r_item = table_row.cells[1].paragraphs[0].add_run(item_text)
        r_item.font.size = Pt(9)

        table_row.cells[2].width = col_widths[2]
        _set_cell_margins(table_row.cells[2])
        r_owner = table_row.cells[2].paragraphs[0].add_run(owner)
        r_owner.font.size = Pt(9)

        table_row.cells[3].width = col_widths[3]
        _shade_cell(table_row.cells[3], status_fill)
        _set_cell_margins(table_row.cells[3])
        p_status = table_row.cells[3].paragraphs[0]
        p_status.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_status = p_status.add_run(status.title())
        r_status.font.size = Pt(9)
        r_status.bold = True
        r_status.font.color.rgb = status_text_color

        table_row.cells[4].width = col_widths[4]
        _set_cell_margins(table_row.cells[4])
        r_notes = table_row.cells[4].paragraphs[0].add_run(notes)
        r_notes.font.size = Pt(9)


def _build_committee_reports(doc: Document, committee_reports: list) -> None:
    _heading1(doc, "Committee Reports")

    if not committee_reports:
        _para(doc, "[No committee reports for this meeting.]", italic=True)
        return

    for report in committee_reports:
        committee = str(report.get("committee", "Committee"))
        chair = str(report.get("chair", ""))
        summary = str(report.get("summary", "[No summary provided.]"))

        _heading2(doc, committee)
        if chair:
            chair_para = doc.add_paragraph()
            r_label = chair_para.add_run("Chair: ")
            r_label.bold = True
            r_label.font.size = Pt(10)
            r_chair = chair_para.add_run(chair)
            r_chair.font.size = Pt(10)

        _para(doc, summary, size=10)
        doc.add_paragraph()


def _build_executive_summary(doc: Document, executive_summary: Optional[str]) -> None:
    _heading1(doc, "Executive Director Summary")

    if executive_summary and executive_summary.strip():
        _para(doc, executive_summary)
    else:
        _para(
            doc,
            "[PLACEHOLDER: Executive Director provides a 1-paragraph summary of organizational "
            "highlights, strategic updates, and key themes for board attention since the last meeting.]",
            italic=True,
        )


def _build_appendix(doc: Document, agenda_items: list) -> None:
    _heading1(doc, "Appendix — Supporting Materials")

    _para(
        doc,
        "The following supporting documents are referenced in the agenda. "
        "Attach printed copies or share digital links as appropriate.",
        italic=True,
        size=10,
    )
    doc.add_paragraph()

    has_docs = False
    for item in agenda_items:
        doc_note = item.get("supporting_doc_note", "")
        if doc_note and str(doc_note).strip():
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(f"{item.get('title', 'Item')} — {doc_note}")
            r.font.size = Pt(10)
            has_docs = True

    if not has_docs:
        _para(doc, "[No supporting documents referenced in this packet.]", italic=True)


# ---------------------------------------------------------------------------
# Core render function
# ---------------------------------------------------------------------------

def render(
    org_name: str,
    meeting_date: str,
    agenda_items: list,
    prepared_by: str,
    meeting_type: str = "regular",
    prior_action_items: Optional[list] = None,
    consent_calendar_items: Optional[list] = None,
    committee_reports: Optional[list] = None,
    executive_summary: Optional[str] = None,
    signer_role: str = "Executive Assistant",
    output_dir: str = "/mnt/user-data/outputs",
) -> str:
    """
    Generate a complete nonprofit board meeting packet as a Word document.

    Parameters
    ----------
    org_name : str
        Full name of the organization
    meeting_date : str
        Meeting date string (e.g. "April 26, 2026" or "2026-04-26")
    agenda_items : list of dict
        Each item: {title, presenter, time_allocation_minutes, type, supporting_doc_note}
        type is one of: discussion / decision / informational / consent
    prepared_by : str
        Name of the EA who prepared the packet
    meeting_type : str
        "regular", "special", or "annual" (default: "regular")
    prior_action_items : list of dict, optional
        {item, owner, status, notes} — carryover from last meeting
    consent_calendar_items : list of str, optional
        Items for en-bloc approval
    committee_reports : list of dict, optional
        {committee, chair, summary}
    executive_summary : str, optional
        1-paragraph ED summary
    signer_role : str
        Signer's title (default: "Executive Assistant")
    output_dir : str
        Directory to save the output file

    Returns
    -------
    str
        Absolute path to the generated .docx file
    """
    # Validate required fields
    if not org_name or not str(org_name).strip():
        raise ValueError("org_name is required")
    if not meeting_date or not str(meeting_date).strip():
        raise ValueError("meeting_date is required")
    if not agenda_items or not isinstance(agenda_items, list):
        raise ValueError("agenda_items is required and must be a non-empty list")
    if not prepared_by or not str(prepared_by).strip():
        raise ValueError("prepared_by is required")

    valid_meeting_types = ("regular", "special", "annual")
    if meeting_type not in valid_meeting_types:
        meeting_type = "regular"

    prior_action_items = prior_action_items or []
    consent_calendar_items = consent_calendar_items or []
    committee_reports = committee_reports or []
    today_str = datetime.date.today().strftime("%B %d, %Y")

    doc = Document()

    # Set page margins (standard letter)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    _add_header_footer(doc, meeting_date, org_name)

    # 1. Cover page
    _build_cover(doc, org_name, meeting_date, meeting_type, prepared_by, signer_role, today_str)
    doc.add_page_break()

    # 2. Agenda
    _build_agenda(doc, agenda_items)
    _hr(doc)

    # 3. Consent Calendar
    _build_consent_calendar(doc, consent_calendar_items)
    doc.add_page_break()

    # 4. Action Item Tracker
    _build_action_tracker(doc, prior_action_items)
    _hr(doc)

    # 5. Committee Reports
    _build_committee_reports(doc, committee_reports)

    if committee_reports:
        doc.add_page_break()

    # 6. Executive Summary
    _build_executive_summary(doc, executive_summary)
    _hr(doc)

    # 7. Appendix
    _build_appendix(doc, agenda_items)

    # Save
    os.makedirs(output_dir, exist_ok=True)
    safe_date = re.sub(r"[^\w]+", "_", meeting_date)[:20].strip("_")
    filename = f"board_packet_{safe_date}.docx"
    out_path = os.path.join(output_dir, filename)
    doc.save(out_path)

    return out_path
