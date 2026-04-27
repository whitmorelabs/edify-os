"""
executive_brief/render.py

Generates a 1-page briefing note for the Executive Director as a Word (.docx) document.
Designed for prep before external meetings: foundation pitches, government meetings,
partner calls, funder check-ins.

Sections:
  1. Header block (topic, date, attendees)
  2. Background
  3. Key Decisions Needed (numbered list)
  4. Stakeholder Positions (table)
  5. Recommended Stance
  6. Risks (bullets)
  7. Prepared by / date

Formatting: compact 0.7" margins, tight paragraph spacing — prints on one page.

All libraries are pre-installed in Anthropic's code-execution sandbox.
"""

import datetime
import os
import re
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------

_NAVY = RGBColor(0x1A, 0x2E, 0x5C)
_SLATE = RGBColor(0x3A, 0x55, 0x77)
_GRAY = RGBColor(0xCC, 0xCC, 0xCC)
_LIGHT_GRAY_TEXT = RGBColor(0x77, 0x77, 0x77)
_AMBER_WARN = RGBColor(0x9C, 0x65, 0x00)
_RED_WARN = RGBColor(0x9C, 0x00, 0x06)
_WHITE_TEXT = RGBColor(0xFF, 0xFF, 0xFF)

_HEADER_FILL = "1A2E5C"
_ALT_ROW_FILL = "EEF1F7"
_SECTION_BAR_FILL = "3A5577"


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------

def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top: int = 40, bottom: int = 40, left: int = 70, right: int = 70) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


# ---------------------------------------------------------------------------
# Compact formatting helpers
# ---------------------------------------------------------------------------

def _compact_para(doc: Document, text: str, bold: bool = False, italic: bool = False,
                  size: int = 9, color: Optional[RGBColor] = None,
                  space_before: int = 2, space_after: int = 2) -> None:
    """Add a paragraph with tight spacing — optimized for 1-page output."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color


def _section_heading(doc: Document, text: str) -> None:
    """Compact colored section heading — replaces Heading 1 to keep one-page format."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = _WHITE_TEXT
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _SECTION_BAR_FILL)
    pPr.append(shd)
    pFmt = p.paragraph_format
    pFmt.left_indent = Inches(0.1)
    pFmt.right_indent = Inches(0.1)


def _hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("─" * 95)
    run.font.color.rgb = _GRAY
    run.font.size = Pt(6)


def _bullet_compact(doc: Document, text: str, size: int = 9,
                    color: Optional[RGBColor] = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _numbered_compact(doc: Document, text: str, size: int = 9) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(size)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _build_header_block(
    doc: Document,
    org_name: str,
    meeting_topic: str,
    meeting_date: str,
    attendees: list,
) -> None:
    """Top metadata block: topic, date, attendees table."""
    # Title bar
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(3)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title_para.add_run("EXECUTIVE DIRECTOR BRIEFING")
    t_run.bold = True
    t_run.font.size = Pt(13)
    t_run.font.color.rgb = _NAVY

    # Sub-heading: org and topic
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_before = Pt(0)
    sub_para.paragraph_format.space_after = Pt(6)
    s_run = sub_para.add_run(f"{org_name}  |  {meeting_topic}")
    s_run.font.size = Pt(10)
    s_run.bold = True
    s_run.font.color.rgb = _SLATE

    # Date line
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_before = Pt(0)
    date_para.paragraph_format.space_after = Pt(6)
    d_run = date_para.add_run(f"Meeting Date: {meeting_date}")
    d_run.font.size = Pt(9)
    d_run.font.color.rgb = _LIGHT_GRAY_TEXT

    _hr(doc)

    # Attendees table
    if attendees:
        _section_heading(doc, "Attendees")

        headers = ["Name", "Role", "Organization"]
        col_widths = [Inches(1.8), Inches(2.0), Inches(2.2)]

        table = doc.add_table(rows=len(attendees) + 1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_row = table.rows[0]
        for i, (h, w) in enumerate(zip(headers, col_widths)):
            cell = hdr_row.cells[i]
            cell.width = w
            _shade_cell(cell, _HEADER_FILL)
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.bold = True
            run.font.color.rgb = _WHITE_TEXT
            run.font.size = Pt(8)

        for row_idx, person in enumerate(attendees):
            name = str(person.get("name", ""))
            role = str(person.get("role", ""))
            organization = str(person.get("organization", org_name))

            table_row = table.rows[row_idx + 1]
            if row_idx % 2 == 1:
                for ci in range(3):
                    _shade_cell(table_row.cells[ci], _ALT_ROW_FILL)

            for ci, val in enumerate([name, role, organization]):
                cell = table_row.cells[ci]
                cell.width = col_widths[ci]
                _set_cell_margins(cell)
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(8)


def _build_background(doc: Document, background: str) -> None:
    _section_heading(doc, "Background")
    _compact_para(doc, background, size=9, space_before=3, space_after=3)


def _build_key_decisions(doc: Document, decisions: list) -> None:
    _section_heading(doc, "Key Decisions Needed")
    for decision in decisions:
        _numbered_compact(doc, str(decision))


def _build_stakeholder_positions(doc: Document, positions: list) -> None:
    if not positions:
        return
    _section_heading(doc, "Stakeholder Positions")

    headers = ["Stakeholder", "Position / Interest"]
    col_widths = [Inches(2.0), Inches(4.0)]

    table = doc.add_table(rows=len(positions) + 1, cols=2)
    table.style = "Table Grid"

    hdr_row = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr_row.cells[i]
        cell.width = w
        _shade_cell(cell, _HEADER_FILL)
        _set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = _WHITE_TEXT
        run.font.size = Pt(8)

    for row_idx, pos in enumerate(positions):
        stakeholder = str(pos.get("stakeholder", ""))
        position = str(pos.get("position", ""))

        table_row = table.rows[row_idx + 1]
        if row_idx % 2 == 1:
            for ci in range(2):
                _shade_cell(table_row.cells[ci], _ALT_ROW_FILL)

        cell0 = table_row.cells[0]
        cell0.width = col_widths[0]
        _set_cell_margins(cell0)
        run0 = cell0.paragraphs[0].add_run(stakeholder)
        run0.font.size = Pt(9)
        run0.bold = True

        cell1 = table_row.cells[1]
        cell1.width = col_widths[1]
        _set_cell_margins(cell1)
        run1 = cell1.paragraphs[0].add_run(position)
        run1.font.size = Pt(9)


def _build_recommended_stance(doc: Document, recommended_stance: str) -> None:
    _section_heading(doc, "Recommended Stance")
    _compact_para(doc, recommended_stance, size=9, space_before=3, space_after=2)


def _build_risks(doc: Document, risks: list) -> None:
    if not risks:
        return
    _section_heading(doc, "Risks & Watch Points")
    for risk in risks:
        _bullet_compact(doc, str(risk), color=_AMBER_WARN)


def _build_footer_block(doc: Document, prepared_by: str, today_str: str) -> None:
    _hr(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("Prepared by: ")
    r1.bold = True
    r1.font.size = Pt(8)
    r2 = p.add_run(f"{prepared_by}  |  ")
    r2.font.size = Pt(8)
    r3 = p.add_run("Date: ")
    r3.bold = True
    r3.font.size = Pt(8)
    r4 = p.add_run(today_str)
    r4.font.size = Pt(8)
    r5 = p.add_run("  |  CONFIDENTIAL")
    r5.font.size = Pt(7)
    r5.italic = True
    r5.font.color.rgb = _LIGHT_GRAY_TEXT


# ---------------------------------------------------------------------------
# Core render function
# ---------------------------------------------------------------------------

def render(
    org_name: str,
    meeting_topic: str,
    meeting_date: str,
    attendees: list,
    background: str,
    key_decisions_needed: list,
    recommended_stance: str,
    prepared_by: str,
    stakeholder_positions: Optional[list] = None,
    risks: Optional[list] = None,
    output_dir: str = "/mnt/user-data/outputs",
) -> str:
    """
    Generate a 1-page Executive Director briefing note as a Word document.

    Parameters
    ----------
    org_name : str
        Name of the organization
    meeting_topic : str
        Short topic label, e.g. "Smith Foundation Funding Meeting"
    meeting_date : str
        Date string for the meeting
    attendees : list of dict
        {name, role, organization} for each attendee on both sides
    background : str
        What brought us here; prior interactions; context paragraph
    key_decisions_needed : list of str
        What needs to come out of this meeting
    recommended_stance : str
        What the ED should say or push for
    prepared_by : str
        EA's name
    stakeholder_positions : list of dict, optional
        {stakeholder, position} for each party
    risks : list of str, optional
        Things to watch for
    output_dir : str
        Directory to save the output file

    Returns
    -------
    str
        Absolute path to the generated .docx file
    """
    # Validate required fields
    required = {
        "org_name": org_name,
        "meeting_topic": meeting_topic,
        "meeting_date": meeting_date,
        "background": background,
        "recommended_stance": recommended_stance,
        "prepared_by": prepared_by,
    }
    for field, value in required.items():
        if not value or not str(value).strip():
            raise ValueError(f"{field} is required and cannot be empty")
    if not attendees or not isinstance(attendees, list):
        raise ValueError("attendees is required and must be a non-empty list")
    if not key_decisions_needed or not isinstance(key_decisions_needed, list):
        raise ValueError("key_decisions_needed is required and must be a non-empty list")

    stakeholder_positions = stakeholder_positions or []
    risks = risks or []
    today_str = datetime.date.today().strftime("%B %d, %Y")

    doc = Document()

    # Narrow margins — designed to fit on one page
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    _build_header_block(doc, org_name, meeting_topic, meeting_date, attendees)
    _build_background(doc, background)
    _build_key_decisions(doc, key_decisions_needed)
    _build_stakeholder_positions(doc, stakeholder_positions)
    _build_recommended_stance(doc, recommended_stance)
    _build_risks(doc, risks)
    _build_footer_block(doc, prepared_by, today_str)

    # Save
    os.makedirs(output_dir, exist_ok=True)
    safe_topic = re.sub(r"[^\w]+", "_", meeting_topic)[:30].strip("_")
    safe_date = re.sub(r"[^\w]+", "_", meeting_date)[:12].strip("_")
    filename = f"executive_brief_{safe_topic}_{safe_date}.docx"
    out_path = os.path.join(output_dir, filename)
    doc.save(out_path)

    return out_path
