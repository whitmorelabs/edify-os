"""
volunteer_handbook_section/render.py

Generates a plain-language volunteer handbook section as a Word (.docx) document.

Supported topics (enum):
  - code_of_conduct
  - safety_protocols
  - mandatory_reporting
  - grievance_procedure
  - confidentiality
  - boundaries_with_youth
  - social_media_policy

Each topic has built-in standard nonprofit prose, customized with org_name,
population_served, and org_specific_notes.

All libraries are pre-installed in Anthropic's code-execution sandbox.
"""

import os
import time
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Supported topics
# ---------------------------------------------------------------------------

VALID_TOPICS = {
    "code_of_conduct",
    "safety_protocols",
    "mandatory_reporting",
    "grievance_procedure",
    "confidentiality",
    "boundaries_with_youth",
    "social_media_policy",
}


# ---------------------------------------------------------------------------
# Topic template data
# ---------------------------------------------------------------------------

def _get_topic_template(
    topic: str,
    org_name: str,
    population_served: str,
    org_specific_notes: Optional[str],
) -> dict:
    """
    Return a dict with keys:
      title, purpose, scope, procedures (list of str),
      faqs (list of {q, a}), compliance_note
    """
    pop = population_served or "the communities we serve"
    note_suffix = f"\n\nOrganization-Specific Note: {org_specific_notes}" if org_specific_notes else ""

    if topic == "code_of_conduct":
        return {
            "title": "Code of Conduct",
            "purpose": (
                f"This Code of Conduct establishes the behavioral standards that {org_name} "
                f"expects of all volunteers. Our mission depends on a culture of respect, "
                f"integrity, and professionalism. These standards protect the people we serve, "
                f"our staff, and our volunteers alike."
            ),
            "scope": (
                f"This policy applies to all {org_name} volunteers in any setting where "
                f"volunteer work occurs — on-site, off-site, at events, during virtual activities, "
                f"or when representing the organization in the community."
            ),
            "procedures": [
                "Treat all individuals — participants, staff, fellow volunteers, and community members — with dignity and respect at all times.",
                "Maintain professional boundaries and avoid relationships that could create conflicts of interest or appear inappropriate.",
                "Arrive on time and notify your supervisor as early as possible if you cannot fulfill a scheduled shift.",
                "Follow all {org_name} policies, procedures, and staff instructions while volunteering.".replace("{org_name}", org_name),
                "Do not use alcohol, illegal substances, or any impairing substances before or during volunteer activities.",
                "Avoid engaging in discriminatory language or behavior based on race, ethnicity, gender identity, sexual orientation, religion, disability, age, or any other protected characteristic.",
                "Protect the confidentiality of participants, staff, and organizational information (see Confidentiality policy).",
                "Use organizational resources — including supplies, equipment, and technology — only for approved volunteer purposes.",
                "Report any concerns about misconduct, safety risks, or policy violations to your supervisor or the Volunteer Coordinator promptly.",
                "Represent the organization responsibly on social media (see Social Media policy).",
            ],
            "faqs": [
                {
                    "q": "What happens if I witness a fellow volunteer violating this code?",
                    "a": f"Report what you observed to the Volunteer Coordinator or a staff supervisor as soon as possible. You are not expected to confront the individual yourself. {org_name} will investigate and respond appropriately.",
                },
                {
                    "q": "Can I take photos of participants during my volunteer shift?",
                    "a": f"No photos of participants may be taken without explicit written consent authorized by {org_name}. If you want to document your experience, speak with the Volunteer Coordinator first.",
                },
                {
                    "q": "What if I disagree with a staff member's decision?",
                    "a": "Raise your concern respectfully and privately with the staff member after the immediate situation is resolved, or bring it to the Volunteer Coordinator. Do not argue with staff in front of participants.",
                },
                {
                    "q": "Is this policy different from the staff code of conduct?",
                    "a": f"The same core values apply to everyone at {org_name}. Some specific procedures differ for staff (e.g., disciplinary processes), but the behavioral expectations are substantively the same.",
                },
            ],
            "compliance_note": (
                "This policy is informed by nonprofit best practices for volunteer management. "
                "Organizations working with vulnerable populations should review this section "
                "with employment counsel before distribution."
            ),
        }

    elif topic == "safety_protocols":
        pop_note = f"working with {pop}" if pop != "the communities we serve" else "in our programs"
        return {
            "title": "Safety Protocols",
            "purpose": (
                f"The safety of participants, volunteers, and staff is {org_name}'s highest operational "
                f"priority. This section outlines the safety practices all volunteers must follow. "
                f"These protocols are not optional — they are the minimum standard for participation."
            ),
            "scope": (
                f"All volunteers {pop_note}. Requirements may vary by program site; your supervisor "
                f"will provide site-specific safety information during orientation."
            ),
            "procedures": [
                "Complete all required safety training before beginning volunteer work. Training completion is tracked and enforced.",
                f"Know the location of emergency exits, first aid kits, and AED devices at your volunteer site.",
                "In the event of an emergency: ensure your own safety first, then assist others if trained to do so. Call 911 for life-threatening emergencies.",
                "Report any injury — to yourself or anyone in your care — to staff immediately, regardless of severity. Incident reports are required.",
                f"Never leave {pop.replace('the communities we serve', 'participants')} unsupervised. Two-adult supervision is required whenever minors or vulnerable individuals are present.",
                "Do not use personal vehicles to transport participants unless explicitly authorized in writing by {org_name} and covered by appropriate insurance.".replace("{org_name}", org_name),
                "Report any physical hazards (broken equipment, unsafe conditions, etc.) to staff before beginning activities in that area.",
                "Maintain current personal health requirements if required by your program (e.g., TB test, flu vaccination).",
                "In the event of a fire alarm or evacuation: exit immediately using the designated route and proceed to the assembly area. Take attendance of participants in your group if possible.",
                "Do not administer medications to participants under any circumstances — this responsibility belongs to designated staff.",
            ],
            "faqs": [
                {
                    "q": "What do I do if a participant is injured during an activity I'm leading?",
                    "a": "Stay calm. Call for staff immediately. Call 911 if there is any risk to life. Do not move the person if a spinal injury is possible. Complete an incident report as soon as the situation is stable.",
                },
                {
                    "q": "Is two-adult supervision always required?",
                    "a": f"Yes, when working with minors or vulnerable adults. {org_name} follows a two-adult rule to protect both participants and volunteers. Never be alone with a participant — if your co-volunteer must leave, pause the activity and get another staff member.",
                },
                {
                    "q": "What if a participant discloses a personal crisis or safety concern to me?",
                    "a": "Listen without judgment and do not promise confidentiality. Immediately notify a staff supervisor. You are not expected to resolve the crisis — escalation is the correct response.",
                },
                {
                    "q": "Can I use my cell phone during my volunteer shift?",
                    "a": f"Personal cell phone use is restricted during direct-service activities with participants. Exceptions: contacting staff in an emergency. Follow any site-specific guidance provided at orientation.",
                },
            ],
            "compliance_note": (
                "Organizations serving youth or vulnerable adults are subject to additional safety "
                "requirements under state law (e.g., Megan's Law, mandated reporting). Consult "
                "your HR or legal counsel to ensure these protocols meet applicable standards."
            ),
        }

    elif topic == "mandatory_reporting":
        pop_detail = pop if pop != "the communities we serve" else "participants in our programs"
        return {
            "title": "Mandatory Reporting",
            "purpose": (
                f"Volunteers at {org_name} who work with {pop_detail} are mandatory reporters "
                f"under applicable state law. This means you are legally required to report "
                f"reasonable suspicions of abuse, neglect, or exploitation — regardless of "
                f"whether you have direct proof. This policy exists to protect the people we serve "
                f"and to protect you."
            ),
            "scope": (
                f"All {org_name} volunteers who have direct contact with {pop_detail}. "
                f"This obligation applies at all times — not only during volunteer hours — "
                f"if information comes to your attention through your volunteer role."
            ),
            "procedures": [
                "You are not required to investigate or confirm abuse — reporting is required when you have a reasonable suspicion.",
                "If a participant discloses abuse, neglect, or exploitation to you: listen without pressing for details, do not promise confidentiality, and do not confront the alleged abuser.",
                "Immediately report your concern to the designated supervisor or Volunteer Coordinator. Do not delay.",
                f"If you cannot reach a supervisor and you believe a participant is in immediate danger, call the appropriate authorities (Child Protective Services, Adult Protective Services, or 911) directly.",
                "Provide your supervisor with the following information: what you observed or were told, any identifying details about the participant and alleged abuser, and the approximate date/time of the disclosure.",
                f"{org_name} staff will coordinate the formal report to the appropriate agency. Document what you reported and when in writing.",
                "Your identity as a reporter is protected by law. Retaliation against mandatory reporters is prohibited.",
                "You may be asked to cooperate with an investigation. Follow staff guidance and consult with the Volunteer Coordinator.",
            ],
            "faqs": [
                {
                    "q": "What if I'm not sure whether what I observed qualifies as abuse?",
                    "a": "Report it anyway. The standard is reasonable suspicion — not certainty. Investigators are trained to determine what happened. Your job is to report what you know.",
                },
                {
                    "q": "What if the alleged abuser is a fellow volunteer or staff member?",
                    "a": f"Report it to the Volunteer Coordinator or a supervisor not involved in the situation. If you believe organizational leadership is involved, contact the board chair or your state's reporting hotline directly.",
                },
                {
                    "q": "Can I keep what a participant tells me confidential?",
                    "a": "No. As a mandatory reporter, you cannot promise confidentiality on matters of safety. If a participant asks you to keep something secret, explain gently that you may need to share it with someone who can help.",
                },
                {
                    "q": "What happens to me if I fail to report?",
                    "a": "Failure to report is a criminal offense in most states and is also a violation of this policy. It may result in removal from the volunteer program and civil or criminal liability.",
                },
                {
                    "q": "Where do I report if I need to contact authorities directly?",
                    "a": f"Your supervisor will provide the current contact information for the appropriate reporting agency during orientation. Keep this reference accessible during your volunteer shifts.",
                },
            ],
            "compliance_note": (
                "Mandatory reporting laws vary by state and by the populations served. "
                "This section describes general principles. Consult your state's child protective "
                "services statute and adult protective services regulations for jurisdiction-specific "
                "requirements. Legal review is strongly recommended before distribution."
            ),
        }

    elif topic == "grievance_procedure":
        return {
            "title": "Grievance Procedure",
            "purpose": (
                f"{org_name} is committed to a volunteer experience that is fair, respectful, "
                f"and safe. If you experience or witness a problem — a conflict with staff, "
                f"a policy concern, harassment, or discrimination — you have the right to raise "
                f"it through this process without fear of retaliation."
            ),
            "scope": (
                f"This procedure applies to all {org_name} volunteers who wish to formally raise "
                f"a concern about their volunteer experience, treatment by staff or other volunteers, "
                f"or the application of organizational policies."
            ),
            "procedures": [
                "Informal resolution (first): if you are comfortable doing so, address the concern directly with the person involved. Many issues can be resolved through a respectful conversation.",
                "If informal resolution is not appropriate or does not resolve the issue: document the incident in writing. Include: what happened, when, who was involved, and any witnesses.",
                f"Submit your written concern to the Volunteer Coordinator. If the Volunteer Coordinator is the subject of your grievance, submit directly to the Executive Director.",
                f"{org_name} will acknowledge receipt of your grievance within 5 business days and will conduct a review within 15 business days.",
                "You may be asked to participate in a fact-finding conversation. You may bring a support person (another volunteer or personal advisor — not legal counsel) to any meeting about your grievance.",
                f"You will receive a written response with the outcome and any corrective actions taken. {org_name} will keep this response confidential to the extent possible.",
                "If you are not satisfied with the outcome, you may appeal in writing to the Executive Director within 10 business days of receiving the response.",
                "Retaliation against anyone who files a good-faith grievance is strictly prohibited and is itself grounds for disciplinary action.",
            ],
            "faqs": [
                {
                    "q": "What kinds of concerns can I raise through this process?",
                    "a": f"Any concern related to your experience as a volunteer at {org_name}: harassment, discrimination, unsafe conditions, unfair treatment, policy violations, or conflicts with staff or other volunteers.",
                },
                {
                    "q": "Will my grievance be kept confidential?",
                    "a": f"{org_name} will handle all grievances with discretion. We cannot guarantee complete confidentiality because investigating a concern may require sharing information with relevant parties. We will share only what is necessary.",
                },
                {
                    "q": "What if my concern involves something illegal?",
                    "a": "You always retain the right to contact relevant government agencies (e.g., EEOC, state labor board, law enforcement) regardless of whether you use this internal process.",
                },
                {
                    "q": "Can I be removed from the volunteer program for filing a grievance?",
                    "a": f"Filing a good-faith grievance is protected. You may not be removed or retaliated against for raising a concern. However, {org_name} retains the right to act on any findings of misconduct on your part that may emerge during an investigation.",
                },
            ],
            "compliance_note": (
                "This grievance procedure reflects best practices for nonprofit volunteer management. "
                "If your organization is a covered employer under the NLRA or other labor law, "
                "consult employment counsel regarding any additional procedural requirements."
            ),
        }

    elif topic == "confidentiality":
        return {
            "title": "Confidentiality Policy",
            "purpose": (
                f"Volunteers at {org_name} frequently encounter sensitive information about "
                f"participants, staff, donors, and organizational operations. This policy exists "
                f"to protect the dignity, privacy, and legal rights of the people we serve — "
                f"and to maintain the trust that makes our work possible."
            ),
            "scope": (
                f"All {org_name} volunteers, in all program areas and settings. This obligation "
                f"continues after your volunteer relationship with {org_name} ends."
            ),
            "procedures": [
                "Treat all participant information — names, contact details, personal histories, service records, and any disclosures — as strictly confidential.",
                "Do not share participant information with family members, friends, other volunteers, or anyone outside the organization unless directed by staff.",
                "Do not discuss individual participants in public spaces, on social media, or in any setting where you could be overheard.",
                f"Access organizational information (files, databases, donor records) only as required for your approved volunteer role. Do not browse records beyond your scope of work.",
                "If you accidentally encounter confidential information you were not meant to see, inform your supervisor immediately.",
                f"Do not share {org_name}'s internal policies, financials, or strategy with outside parties without explicit permission from leadership.",
                "Use a secure, password-protected device for any organizational communications. Do not store participant information on personal devices.",
                "When in doubt about whether something is confidential, treat it as confidential and ask your supervisor.",
            ],
            "faqs": [
                {
                    "q": "Can I share a participant's success story — without using their name?",
                    "a": f"Even anonymized stories require staff approval before sharing. What feels anonymous may still be identifiable in a small community. Always get clearance from the Volunteer Coordinator before sharing any participant story.",
                },
                {
                    "q": "What if a participant asks me not to tell anyone what they disclosed?",
                    "a": "You cannot promise full confidentiality — especially if safety is involved (see Mandatory Reporting policy). Be honest: explain that you care about them and that some disclosures require you to involve a supervisor.",
                },
                {
                    "q": "Does this policy apply to what I see in a participant's file?",
                    "a": "Yes. If your role gives you access to participant files or databases, everything in those records is confidential. Do not share file contents with anyone who does not have a need to know.",
                },
                {
                    "q": "I posted something about my volunteer experience on Instagram. Did I violate this policy?",
                    "a": f"It depends. Posting general enthusiasm about volunteering is usually fine. Posting anything that identifies a participant (by name, photo, or detail that makes them recognizable) without explicit consent from {org_name} is a violation. When in doubt, check first.",
                },
            ],
            "compliance_note": (
                "Depending on your program area, additional confidentiality requirements may apply "
                "(e.g., HIPAA for health-related programs, FERPA for educational records). "
                "Legal counsel should review this section for compliance with applicable law."
            ),
        }

    elif topic == "boundaries_with_youth":
        pop_age = pop if pop != "the communities we serve" else "youth participants"
        return {
            "title": "Professional Boundaries with Youth",
            "purpose": (
                f"{org_name} serves {pop_age}. Maintaining clear, consistent professional boundaries "
                f"is one of the most important responsibilities of every volunteer. Boundaries protect "
                f"young people from exploitation, protect volunteers from false allegations, and "
                f"preserve the integrity of our programs. This is not optional guidance — it is policy."
            ),
            "scope": (
                f"All {org_name} volunteers who have any contact with {pop_age}, including in-person, "
                f"virtual, or digital communication."
            ),
            "procedures": [
                "Never be alone with a youth participant. Two-adult supervision is required at all times. If you find yourself alone with a youth, immediately move to a visible area or call for another adult.",
                "Do not exchange personal contact information (phone numbers, social media handles, personal email) with participants.",
                "Do not contact participants outside of official program channels unless explicitly authorized by staff.",
                "Do not give gifts to individual participants without prior staff approval. Group activities and small program supplies are different from individual gifts.",
                "Physical touch must be appropriate, consensual, and program-relevant (e.g., a high-five or handshake). Avoid prolonged physical contact, private contact, or touch that participants appear to be uncomfortable with.",
                "Do not share personal details about your own life (relationships, struggles, finances, health) beyond what is professionally appropriate.",
                "Do not communicate with participants via personal social media accounts. Official organizational channels only.",
                "If a youth participant attempts to establish a personal relationship outside appropriate boundaries, redirect them kindly and report the interaction to staff.",
                "If you are unsure whether an interaction or communication crosses a boundary, assume it does and check with your supervisor.",
                "Violations of this policy are subject to immediate removal from the volunteer program and may be reported to appropriate authorities.",
            ],
            "faqs": [
                {
                    "q": "A participant wants to follow me on Instagram. What do I do?",
                    "a": "Decline politely and explain that you keep personal and professional accounts separate. Do not follow them back. If they persist, inform your supervisor.",
                },
                {
                    "q": "Can I give a participant my phone number in case they need help?",
                    "a": f"No. Direct them to official {org_name} contact numbers and crisis resources instead. Providing your personal number — even with the best intentions — creates boundary confusion and potential liability.",
                },
                {
                    "q": "A participant told me something personal that concerns me. What do I do?",
                    "a": "Listen without judgment, do not promise to keep it secret, and report it to staff immediately. See the Mandatory Reporting policy for disclosures involving safety.",
                },
                {
                    "q": "I genuinely care about a participant. Isn't staying connected a good thing?",
                    "a": f"Caring about participants is exactly why we need boundaries. Unchecked one-on-one relationships — however well-meaning — put youth at risk and can expose {org_name} to serious liability. Boundaries are an act of care, not indifference.",
                },
                {
                    "q": "What counts as a boundary violation?",
                    "a": "Any interaction that would make a parent, supervisor, or colleague uncomfortable if they observed it. When in doubt, choose the more conservative path.",
                },
            ],
            "compliance_note": (
                "Organizations serving minors are subject to child welfare laws, mandatory reporting "
                "statutes, and background check requirements that vary by state. This policy reflects "
                "national best practices (e.g., Darkness to Light / Stewards of Children framework). "
                "Legal review is required before distribution. Consider requiring annual training "
                "for all volunteers on boundary-setting and child abuse prevention."
            ),
        }

    elif topic == "social_media_policy":
        return {
            "title": "Social Media Policy",
            "purpose": (
                f"Social media is a powerful communication tool — and one that carries real risks "
                f"when used carelessly in a nonprofit context. {org_name}'s Social Media Policy "
                f"sets clear expectations for how volunteers represent the organization and interact "
                f"with participants online. These guidelines protect participants, protect the "
                f"organization's reputation, and protect you."
            ),
            "scope": (
                f"All {org_name} volunteers when posting on any public or semi-public platform "
                f"(Instagram, Facebook, LinkedIn, X/Twitter, TikTok, YouTube, Reddit, etc.) "
                f"in connection with your volunteer role or when identifying yourself as a "
                f"{org_name} volunteer."
            ),
            "procedures": [
                f"Do not post photos, videos, or identifying information about participants without explicit written consent authorized by {org_name}.",
                f"Do not tag {org_name} in personal posts that do not reflect our values, mission, or professional standards.",
                f"If you post about your volunteer experience, keep it general (e.g., 'Had an amazing morning volunteering with kids') — not specific (e.g., names, locations, personal stories).",
                f"Do not use personal social media to communicate with participants. All participant communication must go through official {org_name} channels.",
                f"Do not post {org_name}'s confidential or non-public information: financial details, personnel matters, internal conflicts, or pending announcements.",
                f"Be mindful that your personal social media reflects on you — and, by association, on {org_name}. Avoid posts that could be considered discriminatory, harassing, or inflammatory.",
                f"If you manage an official {org_name} social media account as part of your volunteer role, follow all additional platform-specific guidelines provided by the Communications team.",
                f"If you see a post about {org_name} that contains misinformation or that is damaging, do not engage publicly. Screenshot it and notify the Volunteer Coordinator.",
            ],
            "faqs": [
                {
                    "q": "Can I post a selfie from my volunteer shift?",
                    "a": f"Yes, if no participants are visible and identifiable. Ensure the setting and caption reflect well on {org_name}. If uncertain, ask a supervisor before posting.",
                },
                {
                    "q": "What if a participant tags me in a social media post?",
                    "a": f"Untag yourself as soon as possible and let your supervisor know. You may choose to report the post if it contains identifying information. Do not engage with the participant about it directly — this is a staff matter.",
                },
                {
                    "q": "I saw something on social media that involves a participant's safety. What do I do?",
                    "a": f"Report it to your supervisor immediately. Do not respond to the post yourself. If you believe there is imminent danger, call 911.",
                },
                {
                    "q": "Can I share the organization's public posts on my own social media?",
                    "a": f"Yes. Sharing official content from {org_name}'s public accounts is encouraged — it amplifies our mission and uses information we have already approved.",
                },
            ],
            "compliance_note": (
                "Social media policies intersect with confidentiality law, COPPA (for youth-serving "
                "organizations), and the organization's broader communications strategy. "
                "This section should be reviewed by your Communications Director or legal counsel "
                "before distribution. Update it as platforms and legal standards evolve."
            ),
        }

    else:
        raise ValueError(f"Unknown topic: {topic}. Valid topics: {', '.join(sorted(VALID_TOPICS))}")


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _add_section_heading(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x47, 0x6B)


def _hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)


# ---------------------------------------------------------------------------
# Core render function
# ---------------------------------------------------------------------------

def render(
    topic: str,
    org_name: str,
    population_served: Optional[str] = None,
    org_specific_notes: Optional[str] = None,
    output_dir: str = "/mnt/user-data/outputs",
) -> str:
    """
    Generate a volunteer handbook section Word document.

    Parameters
    ----------
    topic : str
        Policy topic. Must be one of:
        code_of_conduct, safety_protocols, mandatory_reporting,
        grievance_procedure, confidentiality, boundaries_with_youth,
        social_media_policy
    org_name : str
        Name of the organization
    population_served : str, optional
        Description of population, e.g. "youth ages 17-21". Affects safety
        and mandatory reporting nuance.
    org_specific_notes : str, optional
        Freeform additional context or org-specific rules to append.
    output_dir : str
        Directory to save the output file

    Returns
    -------
    str
        Absolute path to the generated .docx file
    """
    # --- Validate ---
    if not topic:
        raise ValueError("topic is required")
    if topic not in VALID_TOPICS:
        raise ValueError(
            f"Invalid topic '{topic}'. Must be one of: {', '.join(sorted(VALID_TOPICS))}"
        )
    if not org_name:
        raise ValueError("org_name is required")

    # --- Get template content ---
    tmpl = _get_topic_template(
        topic=topic,
        org_name=org_name,
        population_served=population_served or "the communities we serve",
        org_specific_notes=org_specific_notes,
    )

    # --- Build document ---
    doc = Document()

    # Document title
    title_para = doc.add_heading(f"Volunteer Handbook", level=0)
    title_para.runs[0].font.color.rgb = RGBColor(0x1A, 0x47, 0x6B)

    org_sub = doc.add_paragraph()
    run = org_sub.add_run(org_name)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x2E, 0x72, 0x9E)

    doc.add_paragraph(
        f"Section: {tmpl['title']}  |  "
        f"Effective: {time.strftime('%B %Y')}  |  "
        f"Review annually or as regulations change"
    ).runs[0].font.size = Pt(9)

    _hr(doc)

    # =========================================================
    # SECTION TITLE
    # =========================================================
    _add_section_heading(doc, tmpl["title"])

    # =========================================================
    # 1. PURPOSE
    # =========================================================
    doc.add_heading("1. Purpose", level=2).runs[0].font.color.rgb = RGBColor(0x2E, 0x72, 0x9E)
    doc.add_paragraph(tmpl["purpose"])

    # =========================================================
    # 2. SCOPE
    # =========================================================
    doc.add_heading("2. Scope", level=2).runs[0].font.color.rgb = RGBColor(0x2E, 0x72, 0x9E)
    doc.add_paragraph(tmpl["scope"])

    # Org-specific note (if provided), inserted early for context
    if org_specific_notes:
        note_para = doc.add_paragraph()
        note_para.add_run("Organization Note: ").bold = True
        note_para.add_run(org_specific_notes).italic = True

    # =========================================================
    # 3. PROCEDURES / EXPECTATIONS
    # =========================================================
    doc.add_heading("3. Procedures and Expectations", level=2).runs[0].font.color.rgb = RGBColor(0x2E, 0x72, 0x9E)
    doc.add_paragraph(
        "All volunteers are expected to uphold the following standards:"
    )

    for procedure in tmpl["procedures"]:
        doc.add_paragraph(procedure, style="List Number")

    # =========================================================
    # 4. WHAT TO DO IF...
    # =========================================================
    doc.add_heading("4. What To Do If...", level=2).runs[0].font.color.rgb = RGBColor(0x2E, 0x72, 0x9E)
    doc.add_paragraph(
        "Common scenarios and how to handle them:"
    )

    for faq in tmpl["faqs"]:
        q_para = doc.add_paragraph()
        q_para.add_run(faq["q"]).bold = True
        a_para = doc.add_paragraph(faq["a"])
        a_para.paragraph_format.left_indent = Pt(18)
        doc.add_paragraph()  # spacer

    # =========================================================
    # 5. ACKNOWLEDGEMENT
    # =========================================================
    _hr(doc)
    doc.add_heading("5. Acknowledgement", level=2).runs[0].font.color.rgb = RGBColor(0x2E, 0x72, 0x9E)

    doc.add_paragraph(
        f"I have read and understood the {tmpl['title']} section of the {org_name} "
        f"Volunteer Handbook. I agree to uphold these standards throughout my volunteer "
        f"relationship with {org_name}."
    )

    doc.add_paragraph()

    # Signature lines
    sig_lines = [
        ("Volunteer Printed Name", "_" * 45),
        ("Volunteer Signature", "_" * 45),
        ("Date", "_" * 20),
        ("Witness / Volunteer Coordinator", "_" * 45),
    ]
    for label, line in sig_lines:
        sig_para = doc.add_paragraph()
        sig_para.add_run(f"{label}: ").bold = True
        sig_para.add_run(line)

    # =========================================================
    # COMPLIANCE NOTE (footer)
    # =========================================================
    _hr(doc)
    note_para = doc.add_paragraph()
    note_para.add_run("Compliance Note: ").bold = True
    note_run = note_para.add_run(tmpl["compliance_note"])
    note_run.italic = True
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    legal_para = doc.add_paragraph()
    legal_para.add_run(
        f"This document was generated by {org_name}'s HR & Volunteer Coordinator "
        f"using Edify OS. Always have current handbook sections reviewed by qualified "
        f"HR or legal counsel before distributing to volunteers."
    ).font.size = Pt(9)

    # --- Save ---
    os.makedirs(output_dir, exist_ok=True)
    timestamp = int(time.time())
    filename = f"volunteer_handbook_{topic}_{timestamp}.docx"
    out_path = os.path.join(output_dir, filename)
    doc.save(out_path)

    return out_path
