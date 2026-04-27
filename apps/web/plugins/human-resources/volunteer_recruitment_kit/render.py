"""
volunteer_recruitment_kit/render.py

Generates a complete volunteer-role recruitment package as a Word (.docx) document.

Sections:
  1. Role Description — title, mission line, responsibilities, skills, time commitment
  2. Outreach Drafts — social media post, email to community partners, flyer text
  3. Screening Questions — 5-7 question list with scoring rubric guidance

All libraries are pre-installed in Anthropic's code-execution sandbox.
"""

import os
import time
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _add_section_heading(doc: Document, text: str) -> None:
    """Add a Level-1 section heading."""
    h = doc.add_heading(text, level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x47, 0x6B)  # dark navy


def _add_subsection(doc: Document, text: str) -> None:
    """Add a Level-2 subsection heading."""
    h = doc.add_heading(text, level=2)
    h.runs[0].font.color.rgb = RGBColor(0x2E, 0x72, 0x9E)


def _add_bold_label(para, label: str, value: str) -> None:
    """Add 'Label: value' with the label bolded in a paragraph."""
    run = para.add_run(label + ": ")
    run.bold = True
    para.add_run(value)


def _bulleted(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _hr(doc: Document) -> None:
    """Insert a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)


# ---------------------------------------------------------------------------
# Screening question generator
# ---------------------------------------------------------------------------

def _build_screening_questions(role_name: str, required_skills: list, program_area: str) -> list:
    """Return 5-7 screening questions tailored to the role."""
    questions = [
        {
            "q": f"What draws you to volunteering as a {role_name} with our {program_area} program?",
            "type": "Open-ended",
            "rubric": "Look for genuine alignment with mission; awareness of population served."
        },
        {
            "q": "Describe a time you worked with a diverse group of people. How did you navigate differences in communication styles or backgrounds?",
            "type": "Behavioral",
            "rubric": "Score 1–5: 1 = vague or avoidant, 5 = specific example demonstrating cultural humility."
        },
        {
            "q": f"Can you commit to the required time of {program_area}-related volunteer work as described? Are there any foreseeable scheduling conflicts?",
            "type": "Logistical / Availability",
            "rubric": "Confirm honest acknowledgment of schedule. A 'yes with caveats' is preferable to an overcommitment."
        },
        {
            "q": "Have you worked with youth, vulnerable adults, or underserved populations before? What did you learn from that experience?",
            "type": "Experience check",
            "rubric": "Relevant experience is a plus; no experience is not disqualifying if answered with curiosity."
        },
        {
            "q": "What would you do if a participant shared something with you that made you concerned for their safety?",
            "type": "Judgment / Safety scenario",
            "rubric": "Look for: escalation to staff/supervisor, not attempting to handle alone, no promise of confidentiality. Critical flag if candidate says they'd keep it private."
        },
    ]

    # Add skill-specific question if required_skills provided
    if required_skills:
        skill_str = ", ".join(required_skills[:2])
        questions.append({
            "q": f"This role requires {skill_str}. Can you tell us how you've applied those skills in a professional, academic, or volunteer context?",
            "type": "Skill demonstration",
            "rubric": f"Verify competency in: {skill_str}. Concrete examples score higher than general claims."
        })

    # Closing motivation question
    questions.append({
        "q": "What does a successful volunteer experience look like to you six months from now?",
        "type": "Goal alignment",
        "rubric": "Look for realistic expectations, interest in growth, and alignment with what the org can offer."
    })

    return questions[:7]  # cap at 7


# ---------------------------------------------------------------------------
# Core render function
# ---------------------------------------------------------------------------

def render(
    role_name: str,
    time_commitment: str,
    program_area: str,
    org_name: str,
    required_skills: Optional[list] = None,
    nice_to_have_skills: Optional[list] = None,
    commitment_length: Optional[str] = None,
    output_dir: str = "/mnt/user-data/outputs",
) -> str:
    """
    Generate a volunteer recruitment kit Word document.

    Parameters
    ----------
    role_name : str
        Title of the volunteer role, e.g. "Youth Mentor"
    time_commitment : str
        Time expectation, e.g. "Saturdays 9am-1pm"
    program_area : str
        Description of the program the volunteer supports
    org_name : str
        Name of the organization
    required_skills : list, optional
        Skills that are required for the role
    nice_to_have_skills : list, optional
        Skills that are preferred but not required
    commitment_length : str, optional
        Duration expectation, e.g. "6 months minimum"
    output_dir : str
        Directory to save the output file

    Returns
    -------
    str
        Absolute path to the generated .docx file
    """
    # --- Validate required inputs ---
    if not role_name:
        raise ValueError("role_name is required")
    if not time_commitment:
        raise ValueError("time_commitment is required")
    if not program_area:
        raise ValueError("program_area is required")
    if not org_name:
        raise ValueError("org_name is required")

    required_skills = required_skills or []
    nice_to_have_skills = nice_to_have_skills or []

    # --- Build document ---
    doc = Document()

    # Title page header
    title = doc.add_heading("Volunteer Recruitment Kit", level=0)
    title.runs[0].font.color.rgb = RGBColor(0x1A, 0x47, 0x6B)

    sub_title = doc.add_paragraph()
    sub_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = sub_title.add_run(f"{role_name}  |  {org_name}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x72, 0x9E)

    doc.add_paragraph(
        f"Prepared for internal use by the HR & Volunteer Coordinator. "
        f"This kit covers the role description, three outreach channel drafts, "
        f"and a structured screening question guide."
    )

    _hr(doc)

    # =========================================================
    # SECTION 1: ROLE DESCRIPTION
    # =========================================================
    _add_section_heading(doc, "Section 1: Role Description")

    doc.add_heading(role_name, level=2)

    # Mission line
    mission_para = doc.add_paragraph()
    mission_para.add_run(
        f"At {org_name}, volunteers are the heartbeat of our work in {program_area}. "
        f"The {role_name} plays a vital role in supporting participants and extending "
        f"our capacity to serve."
    ).italic = True

    doc.add_paragraph()  # spacer

    # Time commitment box
    p = doc.add_paragraph()
    _add_bold_label(p, "Time Commitment", time_commitment)

    if commitment_length:
        p2 = doc.add_paragraph()
        _add_bold_label(p2, "Commitment Length", commitment_length)

    p3 = doc.add_paragraph()
    _add_bold_label(p3, "Program Area", program_area)

    doc.add_paragraph()

    # Responsibilities
    _add_subsection(doc, "Key Responsibilities")
    responsibilities = [
        f"Support {org_name}'s {program_area} program activities as directed by staff",
        "Build positive, professional relationships with program participants",
        "Maintain confidentiality and uphold organizational policies at all times",
        "Attend required orientation, training, and volunteer check-in meetings",
        "Track and report volunteer hours using the organization's designated system",
        "Communicate proactively with the volunteer coordinator about any concerns or scheduling conflicts",
    ]
    for r in responsibilities:
        _bulleted(doc, r)

    doc.add_paragraph()

    # Required Skills
    if required_skills:
        _add_subsection(doc, "Required Skills & Qualifications")
        for skill in required_skills:
            _bulleted(doc, skill)
    else:
        _add_subsection(doc, "Required Skills & Qualifications")
        _bulleted(doc, "No specialized credentials required — a commitment to the mission is the most important qualification")
        _bulleted(doc, "Ability to pass a background check (required for all volunteers working with vulnerable populations)")
        _bulleted(doc, "Reliable transportation or ability to commute to program site")

    doc.add_paragraph()

    # Nice to have
    if nice_to_have_skills:
        _add_subsection(doc, "Nice-to-Have Skills")
        for skill in nice_to_have_skills:
            _bulleted(doc, skill)
        doc.add_paragraph()

    # Equity statement
    equity_para = doc.add_paragraph()
    equity_para.add_run(
        f"{org_name} is committed to building a volunteer team that reflects the communities "
        f"we serve. We strongly encourage individuals from all backgrounds, including those "
        f"with lived experience related to our mission, to apply."
    ).italic = True

    _hr(doc)

    # =========================================================
    # SECTION 2: OUTREACH DRAFTS
    # =========================================================
    _add_section_heading(doc, "Section 2: Outreach Drafts")

    doc.add_paragraph(
        "Use or adapt these three outreach formats across your recruitment channels. "
        "Replace [APPLICATION LINK] with your actual application URL before publishing."
    )

    # --- 2A: Social Media Post ---
    _add_subsection(doc, "2A. Social Media Post (Instagram / Facebook / LinkedIn)")
    social_text = (
        f"We're looking for passionate people to join us as {role_name}s! "
        f"{org_name}'s {program_area} program creates real change — and we can't do it without volunteers like you. "
        f"If you can commit to {time_commitment}, we'd love to meet you. "
        f"No special experience required — just heart, reliability, and a genuine desire to show up for others. "
        f"Apply today and be the reason someone's week gets better. "
        f"[APPLICATION LINK]  #Volunteer #Nonprofit #{org_name.replace(' ', '')}"
    )
    doc.add_paragraph(social_text)

    doc.add_paragraph()

    # --- 2B: Partner Email ---
    _add_subsection(doc, "2B. Email to Community Partners")

    subject_para = doc.add_paragraph()
    _add_bold_label(subject_para, "Subject", f"Volunteer Opportunity: {role_name} at {org_name}")

    email_body = (
        f"Dear [Partner Contact Name],\n\n"
        f"I hope this message finds you well. We are actively recruiting volunteers for our {role_name} role "
        f"in support of {org_name}'s {program_area} program, and I thought your network might include individuals "
        f"who would be a great fit.\n\n"
        f"Volunteers commit to {time_commitment}"
        + (f" for {commitment_length}" if commitment_length else "")
        + f". We provide full orientation and training — no prior experience is required.\n\n"
        f"If you know someone who might be interested, please share our application link: [APPLICATION LINK]. "
        f"I am happy to answer any questions or arrange a brief call to tell you more about the program.\n\n"
        f"Thank you for your continued partnership in supporting our community.\n\n"
        f"Warm regards,\n[Your Name]\n[Title]\n{org_name}"
    )
    doc.add_paragraph(email_body)

    doc.add_paragraph()

    # --- 2C: Flyer Text ---
    _add_subsection(doc, "2C. Flyer Text (Print / Digital Handout)")

    headline_para = doc.add_paragraph()
    run = headline_para.add_run(f"BECOME A {role_name.upper()}")
    run.bold = True
    run.font.size = Pt(16)

    bullet_para = doc.add_paragraph()
    bullet_para.add_run(
        f"Join {org_name} and make a direct impact in our {program_area} program. "
        f"We need volunteers who can give {time_commitment}. "
        f"Training provided. All backgrounds welcome."
    )

    cta_para = doc.add_paragraph()
    run2 = cta_para.add_run(f"Apply now: [APPLICATION LINK]  |  Questions? Contact: [EMAIL/PHONE]")
    run2.bold = True

    _hr(doc)

    # =========================================================
    # SECTION 3: SCREENING QUESTIONS
    # =========================================================
    _add_section_heading(doc, "Section 3: Volunteer Screening Questions")

    doc.add_paragraph(
        f"Use these questions during phone screens or in-person interviews for the {role_name} role. "
        f"A structured interview process promotes equitable, bias-resistant selection. "
        f"Score each behavioral question on the 1–5 rubric provided."
    )

    questions = _build_screening_questions(role_name, required_skills, program_area)

    for i, item in enumerate(questions, start=1):
        q_heading = doc.add_heading(f"Q{i}. {item['q']}", level=3)
        q_heading.runs[0].font.color.rgb = RGBColor(0x22, 0x22, 0x22)

        meta_para = doc.add_paragraph()
        _add_bold_label(meta_para, "Type", item["type"])

        rubric_para = doc.add_paragraph()
        _add_bold_label(rubric_para, "Scoring Guidance", item["rubric"])

        # Notes line
        notes_para = doc.add_paragraph()
        notes_para.add_run("Notes: ").bold = True
        notes_para.add_run("_______________________________________________")
        doc.add_paragraph()

    # Reference check note
    _hr(doc)
    ref_note = doc.add_paragraph()
    ref_note.add_run("Reference Check Note: ").bold = True
    ref_note.add_run(
        "All finalist candidates should complete a reference check with at least one professional or "
        "volunteer reference. For roles working with youth or vulnerable populations, a background "
        "check is required before the volunteer begins service."
    )

    # --- Save ---
    os.makedirs(output_dir, exist_ok=True)
    timestamp = int(time.time())
    filename = f"volunteer_recruitment_kit_{timestamp}.docx"
    out_path = os.path.join(output_dir, filename)
    doc.save(out_path)

    return out_path
