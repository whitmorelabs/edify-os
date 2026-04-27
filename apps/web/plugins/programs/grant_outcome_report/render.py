"""
grant_outcome_report/render.py

Generates a funder-ready grant outcome report as a Word (.docx) document
for nonprofit Programs Directors.

Sections:
  1. Cover Page
  2. Executive Summary
  3. Performance Against Deliverables (table with auto-computed % + R/Y/G coding)
  4. Program Narrative
  5. Participant Story (anonymized, if provided)
  6. Challenges and Adaptations
  7. Next Reporting Period Plan
  8. Compliance note

All libraries are pre-installed in Anthropic's code-execution sandbox.
"""

import datetime
import os
import re
import time
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------

_FOREST = RGBColor(0x1A, 0x5C, 0x38)
_TEAL = RGBColor(0x2E, 0x86, 0x6B)
_GRAY = RGBColor(0xCC, 0xCC, 0xCC)
_ORANGE_DRAFT = RGBColor(0xAA, 0x44, 0x00)

# R/Y/G fills (no #)
_GREEN_FILL = "C6EFCE"    # Excel-style light green
_AMBER_FILL = "FFEB9C"    # Excel-style light amber
_RED_FILL = "FFC7CE"      # Excel-style light red
_HEADER_FILL = "1A5C38"
_WHITE_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
_GREEN_TEXT = RGBColor(0x37, 0x6C, 0x23)
_AMBER_TEXT = RGBColor(0x9C, 0x65, 0x00)
_RED_TEXT = RGBColor(0x9C, 0x00, 0x06)


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------

def _shade_cell(cell, hex_color: str) -> None:
    """Fill a table cell background with a hex color (no #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _heading1(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=1)
    if h.runs:
        h.runs[0].font.color.rgb = _FOREST


def _heading2(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=2)
    if h.runs:
        h.runs[0].font.color.rgb = _TEAL


def _para(doc: Document, text: str, italic: bool = False, size: Optional[int] = None) -> None:
    p = doc.add_paragraph(text)
    if p.runs:
        if italic:
            p.runs[0].italic = True
        if size:
            p.runs[0].font.size = Pt(size)


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 72)
    run.font.color.rgb = _GRAY
    run.font.size = Pt(7)


def _bold_label(para, label: str, value: str) -> None:
    r = para.add_run(label + ": ")
    r.bold = True
    para.add_run(value)


# ---------------------------------------------------------------------------
# Deliverable computation
# ---------------------------------------------------------------------------

def _compute_status(actual: float, target: float) -> tuple:
    """
    Return (pct_str, fill_hex, text_color, status_label).
    Green >= 100%, Amber 75-99%, Red < 75%.
    """
    if target <= 0:
        return ("N/A", _AMBER_FILL, _AMBER_TEXT, "N/A")
    pct = (actual / target) * 100
    pct_str = f"{pct:.1f}%"
    if pct >= 100:
        return (pct_str, _GREEN_FILL, _GREEN_TEXT, "On Target")
    elif pct >= 75:
        return (pct_str, _AMBER_FILL, _AMBER_TEXT, "Approaching")
    else:
        return (pct_str, _RED_FILL, _RED_TEXT, "Below Target")


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _build_cover(
    doc: Document,
    grant_name: str,
    funder_name: str,
    grant_amount: str,
    org_name: str,
    program_name: str,
    report_period: str,
    today_str: str,
) -> None:
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Grant Outcome Report")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = _FOREST

    prog_para = doc.add_paragraph()
    prog_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prog_run = prog_para.add_run(program_name)
    prog_run.bold = True
    prog_run.font.size = Pt(14)
    prog_run.font.color.rgb = _TEAL

    doc.add_paragraph()

    fields = [
        ("Grant", grant_name),
        ("Submitted to", funder_name),
        ("Grant Amount", grant_amount),
        ("Reporting Organization", org_name),
        ("Reporting Period", report_period),
        ("Report Date", today_str),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bold_label(p, label, value)

    doc.add_paragraph()

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "Generated draft — pair with photos and program-specific details before submission to funder."
    )
    nr.italic = True
    nr.font.color.rgb = _ORANGE_DRAFT
    nr.font.size = Pt(9)


def _build_executive_summary(
    doc: Document,
    org_name: str,
    program_name: str,
    funder_name: str,
    grant_amount: str,
    report_period: str,
    deliverable_targets: list,
) -> None:
    _heading1(doc, "Executive Summary")

    # Auto-generate a summary from deliverables
    total = len(deliverable_targets)
    on_target = sum(
        1 for d in deliverable_targets
        if d.get("target_value", 0) > 0 and
        d.get("actual_value", 0) / d["target_value"] >= 1.0
    )
    approaching = sum(
        1 for d in deliverable_targets
        if d.get("target_value", 0) > 0 and
        0.75 <= d.get("actual_value", 0) / d["target_value"] < 1.0
    )
    below = total - on_target - approaching

    status_phrase = (
        f"Of {total} deliverable(s), {on_target} are on or above target, "
        f"{approaching} are approaching target, and {below} are below target."
        if total > 0 else "Deliverable data is detailed in the performance table below."
    )

    summary = (
        f"{org_name} is pleased to submit this outcome report to {funder_name} "
        f"for the {report_period} reporting period. This report summarizes progress on "
        f"{program_name}, supported by a {grant_amount} grant. "
        f"{status_phrase} "
        f"The narrative sections below provide context for performance, highlight participant impact, "
        f"and outline plans for the next reporting period."
    )
    _para(doc, summary)


def _build_deliverables_table(
    doc: Document,
    deliverable_targets: list,
) -> None:
    _heading1(doc, "Performance Against Deliverables")

    headers = ["Deliverable", "Target", "Actual", "% of Target", "Status", "Notes"]
    col_widths = [Inches(2.0), Inches(0.7), Inches(0.7), Inches(0.85), Inches(0.85), Inches(1.4)]

    rows = len(deliverable_targets) + 1
    table = doc.add_table(rows=rows, cols=6)
    table.style = "Table Grid"

    for i, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = w

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        _shade_cell(cell, _HEADER_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = _WHITE_TEXT
        run.font.size = Pt(9)

    # Data rows
    for row_idx, d in enumerate(deliverable_targets):
        deliverable = str(d.get("deliverable", "[Deliverable]"))
        target_val = d.get("target_value", 0)
        actual_val = d.get("actual_value", 0)
        narrative = str(d.get("narrative", ""))

        pct_str, fill_hex, text_color, status_label = _compute_status(
            float(actual_val), float(target_val)
        )

        table_row = table.rows[row_idx + 1]

        # Deliverable
        table_row.cells[0].text = deliverable
        if table_row.cells[0].paragraphs[0].runs:
            table_row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)

        # Target
        cell_t = table_row.cells[1]
        cell_t.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_t = cell_t.paragraphs[0].add_run(str(target_val))
        run_t.font.size = Pt(9)

        # Actual
        cell_a = table_row.cells[2]
        cell_a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_a = cell_a.paragraphs[0].add_run(str(actual_val))
        run_a.font.size = Pt(9)

        # % of Target
        cell_pct = table_row.cells[3]
        _shade_cell(cell_pct, fill_hex)
        cell_pct.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_pct = cell_pct.paragraphs[0].add_run(pct_str)
        run_pct.bold = True
        run_pct.font.size = Pt(9)
        run_pct.font.color.rgb = text_color

        # Status
        cell_s = table_row.cells[4]
        _shade_cell(cell_s, fill_hex)
        cell_s.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_s = cell_s.paragraphs[0].add_run(status_label)
        run_s.font.size = Pt(9)
        run_s.font.color.rgb = text_color

        # Notes
        table_row.cells[5].text = narrative
        if table_row.cells[5].paragraphs[0].runs:
            table_row.cells[5].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # Color key
    key_para = doc.add_paragraph()
    key_para.add_run("Color key: ").bold = True
    green_run = key_para.add_run("Green = On/Above Target (≥100%)  ")
    green_run.font.color.rgb = _GREEN_TEXT
    green_run.font.size = Pt(9)
    amber_run = key_para.add_run("Amber = Approaching Target (75-99%)  ")
    amber_run.font.color.rgb = _AMBER_TEXT
    amber_run.font.size = Pt(9)
    red_run = key_para.add_run("Red = Below Target (<75%)")
    red_run.font.color.rgb = _RED_TEXT
    red_run.font.size = Pt(9)


def _build_program_narrative(doc: Document, program_narrative: Optional[str]) -> None:
    _heading1(doc, "Program Narrative")

    if program_narrative and program_narrative.strip():
        _para(doc, program_narrative)
    else:
        _para(
            doc,
            "[PLACEHOLDER: Describe what happened during the reporting period — key activities, "
            "events, partnerships, staff changes, and programmatic highlights. 2-3 paragraphs.]",
            italic=True,
        )


def _build_participant_story(doc: Document, participant_story: Optional[str]) -> None:
    if not participant_story or not participant_story.strip():
        return

    _heading1(doc, "Participant Story")
    _para(
        doc,
        "The following story is shared with participant permission. Identifying details have been "
        "changed or omitted to protect privacy.",
        italic=True,
    )
    doc.add_paragraph()

    # Format as a pull-quote block with indentation
    quote_para = doc.add_paragraph()
    quote_para.paragraph_format.left_indent = Inches(0.5)
    quote_para.paragraph_format.right_indent = Inches(0.5)
    qr = quote_para.add_run(participant_story)
    qr.italic = True
    qr.font.size = Pt(11)
    qr.font.color.rgb = _TEAL


def _build_challenges(doc: Document, challenges: list) -> None:
    _heading1(doc, "Challenges and Adaptations")

    _para(
        doc,
        "The following challenges were encountered during the reporting period. "
        "Each is framed with what was learned and how the program adapted:",
    )

    if challenges:
        for challenge in challenges:
            _bullet(doc, challenge)
    else:
        _para(doc, "[PLACEHOLDER: List any challenges encountered and how they were addressed.]", italic=True)


def _build_next_period_plan(doc: Document, next_period_plan: Optional[str]) -> None:
    _heading1(doc, "Next Reporting Period Plan")

    if next_period_plan and next_period_plan.strip():
        _para(doc, next_period_plan)
    else:
        _para(
            doc,
            "[PLACEHOLDER: Describe what the organization plans to accomplish in the next "
            "reporting period — key milestones, enrollment targets, activities, and any "
            "course corrections based on this period's performance.]",
            italic=True,
        )


# ---------------------------------------------------------------------------
# Core render function
# ---------------------------------------------------------------------------

def render(
    grant_name: str,
    funder_name: str,
    grant_amount: str,
    org_name: str,
    program_name: str,
    report_period: str,
    deliverable_targets: list,
    participant_story: Optional[str] = None,
    program_narrative: Optional[str] = None,
    challenges_encountered: Optional[list] = None,
    next_period_plan: Optional[str] = None,
    output_dir: str = "/mnt/user-data/outputs",
) -> str:
    """
    Generate a funder-ready grant outcome report Word document.

    Parameters
    ----------
    grant_name : str
        Name of the grant
    funder_name : str
        Name of the funder
    grant_amount : str
        Total grant award (e.g. "$75,000")
    org_name : str
        Full name of the reporting organization
    program_name : str
        What the grant funded
    report_period : str
        Reporting period (e.g. "January-June 2026")
    deliverable_targets : list of dict
        Each dict: {deliverable: str, target_value: int/float,
                    actual_value: int/float, narrative: str}
        Percentage and color coding are auto-computed.
    participant_story : str, optional
        Anonymized participant story
    program_narrative : str, optional
        What happened during the reporting period
    challenges_encountered : list of str, optional
        Challenges and how they were addressed
    next_period_plan : str, optional
        What the organization plans next
    output_dir : str
        Directory to save the output file

    Returns
    -------
    str
        Absolute path to the generated .docx file
    """
    # Validate required inputs
    required = {
        "grant_name": grant_name,
        "funder_name": funder_name,
        "grant_amount": grant_amount,
        "org_name": org_name,
        "program_name": program_name,
        "report_period": report_period,
    }
    for field, value in required.items():
        if not value or not str(value).strip():
            raise ValueError(f"{field} is required and cannot be empty")

    if not deliverable_targets or not isinstance(deliverable_targets, list):
        raise ValueError("deliverable_targets is required and must be a non-empty list")

    challenges_encountered = challenges_encountered or []
    today_str = datetime.date.today().strftime("%B %d, %Y")

    doc = Document()

    _build_cover(
        doc, grant_name, funder_name, grant_amount, org_name,
        program_name, report_period, today_str,
    )
    doc.add_page_break()

    _build_executive_summary(
        doc, org_name, program_name, funder_name, grant_amount,
        report_period, deliverable_targets,
    )
    _hr(doc)

    _build_deliverables_table(doc, deliverable_targets)
    _hr(doc)

    _build_program_narrative(doc, program_narrative)
    _hr(doc)

    _build_participant_story(doc, participant_story)
    if participant_story and participant_story.strip():
        _hr(doc)

    _build_challenges(doc, challenges_encountered)
    _hr(doc)

    _build_next_period_plan(doc, next_period_plan)

    # Compliance footer
    doc.add_paragraph()
    _hr(doc)
    footer = doc.add_paragraph()
    fr = footer.add_run(
        "Generated draft — pair with photos and program-specific details before submission to funder."
    )
    fr.italic = True
    fr.font.color.rgb = _ORANGE_DRAFT
    fr.font.size = Pt(9)

    # Save
    os.makedirs(output_dir, exist_ok=True)
    timestamp = int(time.time())
    safe_name = re.sub(r"[^\w]+", "_", grant_name)[:30].strip("_")
    filename = f"grant_outcome_report_{safe_name}_{timestamp}.docx"
    out_path = os.path.join(output_dir, filename)
    doc.save(out_path)

    return out_path
