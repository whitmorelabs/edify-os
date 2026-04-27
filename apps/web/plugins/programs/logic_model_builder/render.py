"""
logic_model_builder/render.py

Generates a logic model and theory of change Word document for nonprofit
Programs Directors.

Sections:
  1. Cover Page
  2. Theory of Change Narrative
  3. Logic Model Table (5 columns: Inputs | Activities | Outputs |
     Short-Term Outcomes | Long-Term Outcomes)
  4. Measurement Indicators Table (Outcome | Indicator | Data Source)
  5. Assumptions (numbered list)
  6. External Factors (numbered list)

All libraries are pre-installed in Anthropic's code-execution sandbox.
"""

import datetime
import os
import re
import time
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------

_FOREST = RGBColor(0x1A, 0x5C, 0x38)   # deep green — programs domain
_TEAL = RGBColor(0x2E, 0x86, 0x6B)     # medium teal for subheadings
_GRAY = RGBColor(0xCC, 0xCC, 0xCC)
_HEADER_FILL = "1A5C38"                 # hex for table header shading (no #)
_WHITE_TEXT = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------------------
# XML helpers for table cell shading
# ---------------------------------------------------------------------------

def _shade_cell(cell, hex_color: str) -> None:
    """Fill a table cell background with a hex color (no #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _heading1(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=1)
    if h.runs:
        h.runs[0].font.color.rgb = _FOREST


def _para(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph(text)
    if italic and p.runs:
        p.runs[0].italic = True


def _numbered(doc: Document, items: list) -> None:
    for i, item in enumerate(items, start=1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def _hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 72)
    run.font.color.rgb = _GRAY
    run.font.size = Pt(7)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _build_cover(
    doc: Document,
    program_name: str,
    org_name: str,
    today_str: str,
) -> None:
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Logic Model & Theory of Change")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = _FOREST

    prog_para = doc.add_paragraph()
    prog_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prog_run = prog_para.add_run(program_name)
    prog_run.bold = True
    prog_run.font.size = Pt(14)
    prog_run.font.color.rgb = _TEAL

    doc.add_paragraph()

    for label, value in [("Organization", org_name), ("Date", today_str)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label + ": ")
        r.bold = True
        p.add_run(value)

    doc.add_paragraph()

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "Generated draft — review with Programs Director before submitting to funder."
    )
    nr.italic = True
    nr.font.color.rgb = RGBColor(0xAA, 0x44, 0x00)
    nr.font.size = Pt(9)


def _build_theory_of_change(
    doc: Document,
    program_name: str,
    org_name: str,
    target_population: str,
    program_inputs: list,
    activities: list,
    short_term_outcomes: list,
    long_term_outcomes: list,
) -> None:
    _heading1(doc, "Theory of Change")

    inputs_summary = ", ".join(program_inputs[:3])
    if len(program_inputs) > 3:
        inputs_summary += f", and {len(program_inputs) - 3} additional resource(s)"

    activities_summary = "; ".join(activities[:3])
    if len(activities) > 3:
        activities_summary += f"; and {len(activities) - 3} additional activity(ies)"

    short_summary = " and ".join(short_term_outcomes[:2])
    long_summary = " and ".join(long_term_outcomes[:2])

    narrative = (
        f"If {org_name} invests {inputs_summary} into {program_name}, "
        f"and implements {activities_summary}, "
        f"then {target_population} will experience {short_summary} in the short term. "
        f"Over time, these gains will translate into {long_summary}, "
        f"advancing {org_name}'s mission through measurable, participant-centered impact. "
        f"This theory rests on the belief that structured, evidence-informed programming — "
        f"delivered with adequate resources and responsive to participant needs — "
        f"creates the conditions for lasting change."
    )

    _para(doc, narrative)


def _build_logic_model_table(
    doc: Document,
    program_inputs: list,
    activities: list,
    outputs: list,
    short_term_outcomes: list,
    long_term_outcomes: list,
) -> None:
    _heading1(doc, "Logic Model")

    columns = ["Inputs", "Activities", "Outputs", "Short-Term Outcomes", "Long-Term Outcomes"]
    col_data = [program_inputs, activities, outputs, short_term_outcomes, long_term_outcomes]

    # Determine row count: max items in any column + 1 header row
    max_rows = max(len(col) for col in col_data)
    table = doc.add_table(rows=max_rows + 1, cols=5)
    table.style = "Table Grid"

    # Set equal column widths across 6.5" usable width
    col_width = Inches(6.5 / 5)
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width

    hdr_row = table.rows[0]
    for i, col_name in enumerate(columns):
        cell = hdr_row.cells[i]
        _shade_cell(cell, _HEADER_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(col_name)
        run.bold = True
        run.font.color.rgb = _WHITE_TEXT
        run.font.size = Pt(9)

    for row_idx in range(max_rows):
        table_row = table.rows[row_idx + 1]
        for col_idx, col_items in enumerate(col_data):
            cell = table_row.cells[col_idx]
            if row_idx < len(col_items):
                p = cell.paragraphs[0]
                run = p.add_run(col_items[row_idx])
                run.font.size = Pt(9)

    doc.add_paragraph()


def _build_indicators_table(
    doc: Document,
    short_term_outcomes: list,
    long_term_outcomes: list,
    measurement_indicators: list,
) -> None:
    _heading1(doc, "Measurement Indicators")

    _para(
        doc,
        "The following indicators will be used to track progress toward each outcome. "
        "Update the 'Data Source' column with your organization's actual data collection tools.",
    )

    all_outcomes = [("Short-Term", o) for o in short_term_outcomes] + \
                   [("Long-Term", o) for o in long_term_outcomes]

    rows = len(all_outcomes) + 1  # +1 for header row
    table = doc.add_table(rows=rows, cols=3)
    table.style = "Table Grid"

    widths = [Inches(2.5), Inches(2.5), Inches(1.5)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    headers = ["Outcome", "Indicator", "Data Source"]
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        _shade_cell(cell, _HEADER_FILL)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = _WHITE_TEXT
        run.font.size = Pt(9)

    for i, (term_type, outcome) in enumerate(all_outcomes):
        row = table.rows[i + 1]
        row.cells[0].text = f"[{term_type}] {outcome}"
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)

        if i < len(measurement_indicators):
            row.cells[1].text = measurement_indicators[i]
        else:
            row.cells[1].text = "[Add indicator]"
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

        row.cells[2].text = "[Add data source]"
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()


def _build_assumptions(doc: Document, assumptions: list) -> None:
    _heading1(doc, "Assumptions")
    _para(
        doc,
        "This logic model rests on the following assumptions. If these conditions are not met, "
        "program outcomes may be affected:",
    )
    if assumptions:
        _numbered(doc, assumptions)
    else:
        _para(doc, "[No assumptions specified — add conditions required for program success.]", italic=True)


def _build_external_factors(doc: Document, external_factors: list) -> None:
    _heading1(doc, "External Factors")
    _para(
        doc,
        "The following contextual factors are outside the program's direct control but may "
        "influence outcomes:",
    )
    if external_factors:
        _numbered(doc, external_factors)
    else:
        _para(doc, "[No external factors specified — add relevant contextual conditions.]", italic=True)


# ---------------------------------------------------------------------------
# Core render function
# ---------------------------------------------------------------------------

def render(
    program_name: str,
    target_population: str,
    org_name: str,
    program_inputs: list,
    activities: list,
    outputs: list,
    short_term_outcomes: list,
    long_term_outcomes: list,
    measurement_indicators: Optional[list] = None,
    assumptions: Optional[list] = None,
    external_factors: Optional[list] = None,
    output_dir: str = "/mnt/user-data/outputs",
) -> str:
    """
    Generate a logic model and theory of change Word document.

    Parameters
    ----------
    program_name : str
        Name of the program
    target_population : str
        Who the program serves
    org_name : str
        Full name of the organization
    program_inputs : list of str
        Resources that go into the program
    activities : list of str
        What the program does
    outputs : list of str
        Directly countable products of activities
    short_term_outcomes : list of str
        Knowledge/attitude/skill changes (within 1 year)
    long_term_outcomes : list of str
        Behavior changes and mission impact (1-3+ years)
    measurement_indicators : list of str, optional
        How each outcome is tracked (pairs with outcomes in order)
    assumptions : list of str, optional
        What must be true for the model to work
    external_factors : list of str, optional
        Context outside the program's control that affects success
    output_dir : str
        Directory to save the output file

    Returns
    -------
    str
        Absolute path to the generated .docx file
    """
    # Validate required inputs
    required = {
        "program_name": program_name,
        "target_population": target_population,
        "org_name": org_name,
    }
    for field, value in required.items():
        if not value or not str(value).strip():
            raise ValueError(f"{field} is required and cannot be empty")

    for list_field, list_value in [
        ("program_inputs", program_inputs),
        ("activities", activities),
        ("outputs", outputs),
        ("short_term_outcomes", short_term_outcomes),
        ("long_term_outcomes", long_term_outcomes),
    ]:
        if not list_value or not isinstance(list_value, list) or len(list_value) == 0:
            raise ValueError(f"{list_field} is required and must be a non-empty list")

    measurement_indicators = measurement_indicators or []
    assumptions = assumptions or []
    external_factors = external_factors or []

    today_str = datetime.date.today().strftime("%B %d, %Y")
    doc = Document()

    _build_cover(doc, program_name, org_name, today_str)
    doc.add_page_break()

    _build_theory_of_change(
        doc, program_name, org_name, target_population,
        program_inputs, activities, short_term_outcomes, long_term_outcomes,
    )
    _hr(doc)

    _build_logic_model_table(
        doc, program_inputs, activities, outputs,
        short_term_outcomes, long_term_outcomes,
    )
    _hr(doc)

    _build_indicators_table(
        doc, short_term_outcomes, long_term_outcomes, measurement_indicators,
    )
    _hr(doc)

    _build_assumptions(doc, assumptions)
    _hr(doc)

    _build_external_factors(doc, external_factors)

    # Save
    os.makedirs(output_dir, exist_ok=True)
    timestamp = int(time.time())
    safe_name = re.sub(r"[^\w]+", "_", program_name)[:30].strip("_")
    filename = f"logic_model_{safe_name}_{timestamp}.docx"
    out_path = os.path.join(output_dir, filename)
    doc.save(out_path)

    return out_path
