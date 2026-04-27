"""
participant_survey/render.py

Generates a participant survey instrument as a ready-to-print Word (.docx)
document for nonprofit Programs Directors.

Survey types:
  - intake     : demographics, baseline self-assessment, expectations, accessibility
  - satisfaction: facility, staff, content, NPS, open feedback
  - outcome    : baseline recall, learning gains, behavior changes, impact, next steps
  - exit       : reasons for leaving, what worked, what didn't, would-return, suggestions

Each survey has 10-15 questions grouped into 3-5 thematic sections.

All libraries are pre-installed in Anthropic's code-execution sandbox.
"""

import datetime
import os
import re
import time
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------

_FOREST = RGBColor(0x1A, 0x5C, 0x38)
_TEAL = RGBColor(0x2E, 0x86, 0x6B)
_GRAY = RGBColor(0xCC, 0xCC, 0xCC)
_ORANGE_DRAFT = RGBColor(0xAA, 0x44, 0x00)
_HEADER_FILL = "1A5C38"
_WHITE_TEXT = RGBColor(0xFF, 0xFF, 0xFF)

_SURVEY_TYPE_LABELS = {
    "intake": "Participant Intake Form",
    "satisfaction": "Program Satisfaction Survey",
    "outcome": "Program Outcome Survey",
    "exit": "Program Exit Survey",
}

# ---------------------------------------------------------------------------
# Question templates (hardcoded nonprofit standards)
# ---------------------------------------------------------------------------

# Each question is a dict: {section, text, q_type}
# q_type: "likert5" | "mc" (multiple choice) | "open" | "rank" | "yesno"
# mc items: list of strings
# rank items: list of strings

_INTAKE_QUESTIONS = [
    # Section 1: About You
    {"section": "About You", "text": "What is your age?", "q_type": "mc",
     "items": ["Under 16", "16-17", "18-19", "20-21", "22 or older", "Prefer not to say"]},
    {"section": "About You", "text": "What is your gender?", "q_type": "mc",
     "items": ["Man", "Woman", "Non-binary / gender non-conforming", "Prefer to self-describe: ___________", "Prefer not to say"]},
    {"section": "About You", "text": "What is your primary language at home?", "q_type": "mc",
     "items": ["English", "Spanish", "Other: ___________"]},
    {"section": "About You", "text": "Do you identify as having a disability?", "q_type": "mc",
     "items": ["Yes", "No", "Prefer not to say"]},

    # Section 2: Your Starting Point
    {"section": "Your Starting Point", "text": "Before joining this program, how would you rate your readiness for the program's goals?",
     "q_type": "likert5", "anchors": ("Not at all ready", "Very ready")},
    {"section": "Your Starting Point", "text": "What do you most hope to gain from this program? (Check all that apply)", "q_type": "mc",
     "items": ["New skills", "A job or career path", "Confidence", "Connections / networking", "Certificate or credential", "Other: ___________"]},
    {"section": "Your Starting Point", "text": "What is your current employment or school status?", "q_type": "mc",
     "items": ["Currently employed (full-time)", "Currently employed (part-time)", "Not currently working", "In school", "Other: ___________"]},

    # Section 3: What to Expect
    {"section": "What to Expect", "text": "How did you hear about this program?", "q_type": "mc",
     "items": ["School referral", "Family or friend", "Social media", "Flyer or brochure", "Other organization", "Other: ___________"]},
    {"section": "What to Expect", "text": "How many hours per week can you commit to program activities?", "q_type": "mc",
     "items": ["Less than 5 hours", "5-10 hours", "10-20 hours", "More than 20 hours"]},

    # Section 4: Your Needs
    {"section": "Your Needs", "text": "Do you need any accommodations to fully participate in program activities?", "q_type": "open"},
    {"section": "Your Needs", "text": "Do you have access to a smartphone, tablet, or computer?", "q_type": "mc",
     "items": ["Yes — smartphone", "Yes — computer or tablet", "Both", "No"]},
    {"section": "Your Needs", "text": "How would you prefer to receive program updates and reminders?", "q_type": "mc",
     "items": ["Text message", "Email", "Phone call", "Program app", "No preference"]},
    {"section": "Your Needs", "text": "Is there anything else you'd like the program team to know about you or your needs?", "q_type": "open"},
]

_SATISFACTION_QUESTIONS = [
    # Section 1: Overall Experience
    {"section": "Overall Experience", "text": "Overall, how satisfied are you with this program?",
     "q_type": "likert5", "anchors": ("Very dissatisfied", "Very satisfied")},
    {"section": "Overall Experience", "text": "The program met my expectations.", "q_type": "likert5",
     "anchors": ("Strongly disagree", "Strongly agree")},

    # Section 2: Staff and Services
    {"section": "Staff and Services", "text": "Staff treated me with respect and dignity.", "q_type": "likert5",
     "anchors": ("Strongly disagree", "Strongly agree")},
    {"section": "Staff and Services", "text": "Staff were knowledgeable and helpful when I had questions.", "q_type": "likert5",
     "anchors": ("Strongly disagree", "Strongly agree")},
    {"section": "Staff and Services", "text": "I felt comfortable asking staff for help.", "q_type": "likert5",
     "anchors": ("Strongly disagree", "Strongly agree")},

    # Section 3: Program Content
    {"section": "Program Content", "text": "The program content was relevant to my goals.", "q_type": "likert5",
     "anchors": ("Strongly disagree", "Strongly agree")},
    {"section": "Program Content", "text": "The pace of the program was:", "q_type": "mc",
     "items": ["Too slow", "Just right", "Too fast"]},
    {"section": "Program Content", "text": "Which part of the program was most valuable to you?", "q_type": "open"},

    # Section 4: Would You Recommend Us?
    {"section": "Would You Recommend Us?",
     "text": "On a scale of 0-10, how likely are you to recommend this program to a friend or family member?\n(0 = Not at all likely | 10 = Extremely likely)",
     "q_type": "open"},
    {"section": "Would You Recommend Us?", "text": "Why did you give that score?", "q_type": "open"},

    # Section 5: Your Voice
    {"section": "Your Voice", "text": "What is one thing the program does really well?", "q_type": "open"},
    {"section": "Your Voice", "text": "What is one thing the program should improve?", "q_type": "open"},
    {"section": "Your Voice", "text": "Is there anything else you'd like us to know?", "q_type": "open"},
]

_OUTCOME_QUESTIONS = [
    # Section 1: Thinking Back
    {"section": "Thinking Back", "text": "Before you started the program, how would you rate your skills and readiness in the program's focus area?",
     "q_type": "likert5", "anchors": ("Very low", "Very high")},
    {"section": "Thinking Back", "text": "What was your main goal when you joined the program?", "q_type": "open"},

    # Section 2: What You Learned
    {"section": "What You Learned", "text": "Because of this program, I have new skills I didn't have before.", "q_type": "likert5",
     "anchors": ("Strongly disagree", "Strongly agree")},
    {"section": "What You Learned", "text": "Because of this program, I feel more confident.", "q_type": "likert5",
     "anchors": ("Strongly disagree", "Strongly agree")},
    {"section": "What You Learned", "text": "What is the most important thing you learned in this program?", "q_type": "open"},

    # Section 3: Changes You've Made
    {"section": "Changes You've Made", "text": "Since participating in this program, I have taken steps toward my goals.", "q_type": "likert5",
     "anchors": ("Strongly disagree", "Strongly agree")},
    {"section": "Changes You've Made", "text": "What actions have you taken since starting the program? (Check all that apply)", "q_type": "mc",
     "items": ["Applied for a job", "Started a new job", "Enrolled in school or training", "Improved a specific skill", "Made a new professional connection", "Other: ___________"]},
    {"section": "Changes You've Made", "text": "How would you rate your readiness NOW compared to when you started?",
     "q_type": "likert5", "anchors": ("Much lower", "Much higher")},

    # Section 4: Your Impact Story
    {"section": "Your Impact Story",
     "text": "In your own words, how has this program made a difference in your life?", "q_type": "open"},
    {"section": "Your Impact Story",
     "text": "Would you be willing to share your story (anonymously) to help others learn about this program?", "q_type": "mc",
     "items": ["Yes — you can use my first name", "Yes — anonymously only", "No, I'd prefer to keep it private"]},

    # Section 5: Looking Ahead
    {"section": "Looking Ahead", "text": "What is your next goal after this program?", "q_type": "open"},
    {"section": "Looking Ahead", "text": "What additional support would help you reach your next goal?", "q_type": "open"},
]

_EXIT_QUESTIONS = [
    # Section 1: Your Time with Us
    {"section": "Your Time with Us", "text": "How long were you in the program?", "q_type": "mc",
     "items": ["Less than 1 month", "1-3 months", "4-6 months", "More than 6 months"]},
    {"section": "Your Time with Us", "text": "Why are you leaving the program at this time? (Check all that apply)", "q_type": "mc",
     "items": ["I completed the program", "I got a job or reached my goal", "Life circumstances changed (family, housing, health)", "The program wasn't the right fit", "Transportation or schedule barriers", "I'm transferring to another program", "Other: ___________"]},

    # Section 2: What Worked
    {"section": "What Worked", "text": "What was most helpful about this program?", "q_type": "open"},
    {"section": "What Worked", "text": "Which program activity or service made the biggest difference for you?", "q_type": "open"},
    {"section": "What Worked", "text": "Staff were supportive and responsive to my needs.", "q_type": "likert5",
     "anchors": ("Strongly disagree", "Strongly agree")},

    # Section 3: What Could Be Better
    {"section": "What Could Be Better", "text": "What was most difficult or frustrating about this program?", "q_type": "open"},
    {"section": "What Could Be Better", "text": "What would you change to make this program better for future participants?", "q_type": "open"},
    {"section": "What Could Be Better", "text": "Were there any barriers that made it hard for you to participate?", "q_type": "mc",
     "items": ["Transportation", "Schedule / timing", "Childcare", "Language or communication", "Health or disability-related", "No major barriers", "Other: ___________"]},

    # Section 4: Next Steps
    {"section": "Next Steps", "text": "What are your plans after leaving the program?", "q_type": "open"},
    {"section": "Next Steps", "text": "If we offered future programs, would you return or refer someone you know?", "q_type": "likert5",
     "anchors": ("Definitely not", "Definitely yes")},

    # Section 5: Final Thoughts
    {"section": "Final Thoughts", "text": "Is there anything else you'd like to share before you go?", "q_type": "open"},
    {"section": "Final Thoughts", "text": "May we contact you in the future to follow up on your progress?", "q_type": "mc",
     "items": ["Yes — you have my contact information", "Yes — please update my contact info below: ___________", "No, I'd prefer not to be contacted"]},
]

_QUESTION_BANKS = {
    "intake": _INTAKE_QUESTIONS,
    "satisfaction": _SATISFACTION_QUESTIONS,
    "outcome": _OUTCOME_QUESTIONS,
    "exit": _EXIT_QUESTIONS,
}


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _heading1(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=1)
    if h.runs:
        h.runs[0].font.color.rgb = _FOREST


def _heading2(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=2)
    if h.runs:
        h.runs[0].font.color.rgb = _TEAL


def _para(doc: Document, text: str, italic: bool = False, size: Optional[int] = None) -> None:
    p = doc.add_paragraph(text)
    if p.runs:
        if italic:
            p.runs[0].italic = True
        if size:
            p.runs[0].font.size = Pt(size)


def _blank_lines(doc: Document, count: int = 3) -> None:
    """Add blank lines for handwritten open-ended responses."""
    for _ in range(count):
        p = doc.add_paragraph("_" * 80)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if p.runs:
            p.runs[0].font.color.rgb = _GRAY
            p.runs[0].font.size = Pt(9)


def _likert_row(doc: Document, anchors: tuple) -> None:
    """Render a 1-5 Likert scale as a simple table row."""
    table = doc.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    scale_width = Inches(1.3)

    labels = ["1", "2", "3", "4", "5"]
    for i, label in enumerate(labels):
        cell = table.rows[0].cells[i]
        cell.width = scale_width
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)

    # Anchor labels
    for i, label in enumerate(labels):
        cell = table.rows[1].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i == 0:
            run = p.add_run(anchors[0])
        elif i == 4:
            run = p.add_run(anchors[1])
        else:
            run = p.add_run("")
        run.font.size = Pt(8)
        run.font.color.rgb = _TEAL

    doc.add_paragraph()


def _mc_options(doc: Document, items: list) -> None:
    """Render multiple choice options with ☐ checkbox."""
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.add_run("  ☐  " + item).font.size = Pt(10)


def _rank_options(doc: Document, items: list) -> None:
    """Render rank-order options (1 = most important)."""
    note = doc.add_paragraph("Rank in order of importance (1 = most important):")
    if note.runs:
        note.runs[0].italic = True
        note.runs[0].font.size = Pt(9)
    for item in items:
        p = doc.add_paragraph()
        p.add_run(f"  ___ {item}").font.size = Pt(10)


def _hr(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run("─" * 72)
    run.font.color.rgb = _GRAY
    run.font.size = Pt(7)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _build_cover(
    doc: Document,
    survey_type: str,
    program_name: str,
    org_name: str,
    today_str: str,
) -> None:
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_para.add_run("[Organization Logo]")
    r.bold = True
    r.font.color.rgb = _GRAY
    r.font.size = Pt(10)

    doc.add_paragraph()

    survey_label = _SURVEY_TYPE_LABELS.get(survey_type, "Participant Survey")
    title_para2 = doc.add_paragraph()
    title_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para2.add_run(survey_label)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = _FOREST

    prog_para = doc.add_paragraph()
    prog_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prog_run = prog_para.add_run(program_name)
    prog_run.bold = True
    prog_run.font.size = Pt(13)
    prog_run.font.color.rgb = _TEAL

    org_para = doc.add_paragraph()
    org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_run = org_para.add_run(org_name)
    org_run.font.size = Pt(11)

    doc.add_paragraph()

    # Intro + privacy statement
    intro_texts = {
        "intake": (
            f"Welcome to {program_name}! Please take a few minutes to complete this form. "
            f"Your answers will help us get to know you and make sure we can best support you."
        ),
        "satisfaction": (
            f"Thank you for participating in {program_name}. "
            f"We want to hear about your experience so we can keep improving. "
            f"This survey takes about 5-10 minutes."
        ),
        "outcome": (
            f"Congratulations on your progress in {program_name}! "
            f"Please reflect on your experience and share how the program has impacted you. "
            f"Your answers help us demonstrate our impact to funders and improve future programming."
        ),
        "exit": (
            f"Thank you for being part of {program_name}. "
            f"As you leave, we'd love to hear about your experience — what helped, what didn't, "
            f"and what you're planning next. Your honest feedback matters."
        ),
    }
    intro_text = intro_texts.get(survey_type, f"Please complete this survey about {program_name}.")
    _para(doc, intro_text)

    doc.add_paragraph()

    privacy = doc.add_paragraph()
    pr = privacy.add_run(
        "This survey is voluntary. Your responses are confidential and will only be used "
        "in aggregate to improve our programs. You may skip any question you prefer not to answer."
    )
    pr.italic = True
    pr.font.size = Pt(9)
    pr.font.color.rgb = _TEAL


def _render_question(
    doc: Document,
    q_num: int,
    question: dict,
    language_level: str = "plain",
) -> None:
    """Render a single question with its appropriate format."""
    q_text = question["text"]
    q_type = question["q_type"]

    # Question label
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    label_run = p.add_run(f"{q_num}. ")
    label_run.bold = True
    label_run.font.size = Pt(10)
    text_run = p.add_run(q_text)
    text_run.font.size = Pt(10)

    if q_type == "likert5":
        anchors = question.get("anchors", ("Strongly disagree", "Strongly agree"))
        _likert_row(doc, anchors)

    elif q_type == "mc":
        items = question.get("items", [])
        _mc_options(doc, items)
        doc.add_paragraph()

    elif q_type == "open":
        _blank_lines(doc, 3)

    elif q_type == "rank":
        items = question.get("items", [])
        _rank_options(doc, items)
        doc.add_paragraph()

    elif q_type == "yesno":
        _mc_options(doc, ["Yes", "No"])
        doc.add_paragraph()


def _build_survey_sections(
    doc: Document,
    questions: list,
    additional_topics: list,
    language_level: str,
) -> None:
    """Group questions by section and render them."""
    # Group by section
    sections = {}
    section_order = []
    for q in questions:
        sec = q["section"]
        if sec not in sections:
            sections[sec] = []
            section_order.append(sec)
        sections[sec].append(q)

    q_num = 1
    for sec in section_order:
        _heading2(doc, sec)
        for q in sections[sec]:
            _render_question(doc, q_num, q, language_level)
            q_num += 1

    # Append additional topic questions as open-ended in a new section
    if additional_topics:
        _heading2(doc, "Additional Questions")
        for topic in additional_topics:
            q = {
                "section": "Additional Questions",
                "text": f"Please share any thoughts or experiences related to {topic}:",
                "q_type": "open",
            }
            _render_question(doc, q_num, q, language_level)
            q_num += 1


def _build_scoring_guide(doc: Document, survey_type: str) -> None:
    """Append a scoring guide for outcome surveys only."""
    if survey_type != "outcome":
        return

    doc.add_page_break()
    _heading1(doc, "Scoring Guide (Staff Use Only)")
    _para(doc, "This section is for program staff. Do not distribute to participants.")
    _hr(doc)

    _heading2(doc, "Likert Scale Interpretation")
    scoring = [
        ("4.5 – 5.0", "Very high agreement / readiness — strong positive outcome"),
        ("3.5 – 4.4", "High agreement — positive trend, note for follow-up"),
        ("2.5 – 3.4", "Moderate — mixed experience, may benefit from check-in"),
        ("1.5 – 2.4", "Low agreement — follow up with participant"),
        ("1.0 – 1.4", "Very low — flag for case manager review"),
    ]
    table = doc.add_table(rows=len(scoring) + 1, cols=2)
    table.style = "Table Grid"

    headers = ["Average Score", "Interpretation"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    for i, (score_range, meaning) in enumerate(scoring):
        row = table.rows[i + 1]
        row.cells[0].text = score_range
        row.cells[1].text = meaning
        for cell in row.cells:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()
    _para(
        doc,
        "For pre/post comparisons: calculate the average score for relevant questions at intake "
        "and at outcome survey administration. A positive difference indicates measurable gains. "
        "Document individual outliers for case management follow-up.",
        italic=True,
    )


# ---------------------------------------------------------------------------
# Core render function
# ---------------------------------------------------------------------------

def render(
    survey_type: str,
    program_name: str,
    org_name: str,
    target_population: str,
    program_focus: Optional[str] = None,
    additional_topics: Optional[list] = None,
    language_level: Optional[str] = "plain",
    output_dir: str = "/mnt/user-data/outputs",
) -> str:
    """
    Generate a participant survey instrument Word document.

    Parameters
    ----------
    survey_type : str
        "intake", "satisfaction", "outcome", or "exit"
    program_name : str
        Name of the program
    org_name : str
        Full name of the organization
    target_population : str
        Who will complete the survey: "youth", "adults", "parents", "general"
    program_focus : str, optional
        Program domain for tailored questions (e.g. "workforce development")
    additional_topics : list of str, optional
        Extra question areas to weave in
    language_level : str, optional
        "plain" (default) or "formal"
    output_dir : str
        Directory to save the output file

    Returns
    -------
    str
        Absolute path to the generated .docx file
    """
    # Validate
    valid_types = ("intake", "satisfaction", "outcome", "exit")
    if not survey_type or survey_type.lower() not in valid_types:
        raise ValueError(f"survey_type must be one of: {', '.join(valid_types)}")
    survey_type = survey_type.lower()

    for field, value in [("program_name", program_name), ("org_name", org_name), ("target_population", target_population)]:
        if not value or not str(value).strip():
            raise ValueError(f"{field} is required and cannot be empty")

    language_level = (language_level or "plain").lower()
    if language_level not in ("plain", "formal"):
        language_level = "plain"

    additional_topics = additional_topics or []
    today_str = datetime.date.today().strftime("%B %d, %Y")

    doc = Document()

    _build_cover(doc, survey_type, program_name, org_name, today_str)
    doc.add_page_break()

    questions = list(_QUESTION_BANKS[survey_type])
    _build_survey_sections(doc, questions, additional_topics, language_level)

    _build_scoring_guide(doc, survey_type)

    # Footer note
    _hr(doc)
    note = doc.add_paragraph()
    nr = note.add_run(
        f"Generated draft — {_SURVEY_TYPE_LABELS[survey_type]} for {program_name}. "
        f"Pilot with 2-3 participants before full rollout. Translate as needed."
    )
    nr.italic = True
    nr.font.color.rgb = _ORANGE_DRAFT
    nr.font.size = Pt(9)

    # Save
    os.makedirs(output_dir, exist_ok=True)
    timestamp = int(time.time())
    safe_name = re.sub(r"[^\w]+", "_", program_name)[:30].strip("_")
    filename = f"survey_{survey_type}_{safe_name}_{timestamp}.docx"
    out_path = os.path.join(output_dir, filename)
    doc.save(out_path)

    return out_path
